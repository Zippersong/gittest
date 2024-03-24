# -*- coding: utf-8 -*-
"""
Created on Fri Aug 18 15:07:33 2023

@author: zp457
"""
#%%-----------20240119涛哥要的银信金24年应还息费、服务费，考虑到sql没有unstack，还是放在python实现

import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 截至目前的逾期天数
sql = """ select order_no,MAX( DATEDIFF(CURRENT_DATE,receivable_time)) as od_day_now 
from loan_core.core_repayment_plan where is_del=0 and repayment_status=2 group by order_no   """
od_days=pd.read_sql(sql,cnx)
od_days['od30']=od_days.apply(lambda x: 1 if x.od_day_now>30 else 0,axis=1)


od_days_30=od_days[od_days.od_day_now>30]
od_days.od30.value_counts()



# 截至目前的放款订单及状态
sql = """ select a.order_no,a.withdraw_amount,DATE_FORMAT(a.loan_time,"%Y%m") as 放款月份,b.settle_status,b.settle_date from loan_core.core_loan_order as a
left join loan_core.core_loan_order_extend as b on a.order_no=b.order_no where a.loan_status=2  """
loan_order=pd.read_sql(sql,cnx)

loan_order_t=pd.merge(loan_order,od_days,how='left',on='order_no')
loan_order_t.od30.fillna(0,inplace=True)
loan_order_t['在库']=loan_order_t.apply(lambda x: 1 if pd.isna(x.settle_date) else 0 ,axis=1)


# a_test=pd.crosstab(loan_order_t.在库,loan_order_t.od30)#216906863839678492

loan_order=loan_order[~loan_order.order_no.isin(od_days_30.order_no)]
loan_order['在库']=loan_order.apply(lambda x: 1 if pd.isna(x.settle_date) else 0 ,axis=1)


# 截至目前的放款订单已还金额
sql = """ select a.order_no,a.receivable_time,b.subject_type,b.receivable_amount as 应还金额,b.repayment_amount as 已还金额,
b.remission_amount as 减免金额
from loan_core.core_repayment_plan as a 
left join loan_core.core_repayment_plan_detail as b  on a.id=b.repay_plan_id
where a.is_del=0 and b.is_del=0  """
repay_plan=pd.read_sql(sql,cnx)


repay_plan=repay_plan[~repay_plan.order_no.isin(od_days_30.order_no)]
repay_plan['yinghuan_month']=pd.to_datetime(repay_plan['receivable_time']).dt.strftime('%Y%m')

loan_order_temp=loan_order[['order_no','放款月份']]
repay_plan=pd.merge(repay_plan,loan_order_temp,how='left',on='order_no')
repay_plan['调整应还金额']=repay_plan['应还金额']-repay_plan['已还金额']-repay_plan['减免金额']

yinghuan_principal = repay_plan[repay_plan.subject_type=='principal'].groupby(['放款月份','yinghuan_month']).agg({'调整应还金额':'sum'}).unstack()
yinghuan_interest = repay_plan[repay_plan.subject_type=='interest'].groupby(['放款月份','yinghuan_month']).agg({'调整应还金额':'sum'}).unstack()
yinghuan_service = repay_plan[repay_plan.subject_type=='serviceFee'].groupby(['放款月份','yinghuan_month']).agg({'调整应还金额':'sum'}).unstack()


# yinghuan_principal = repay_plan[repay_plan.subject_type=='principal'].groupby(['放款月份','yinghuan_month']).agg({'应还金额':'sum'}).unstack()
# yinghuan_interest = repay_plan[repay_plan.subject_type=='interest'].groupby(['放款月份','yinghuan_month']).agg({'应还金额':'sum'}).unstack()
# yinghuan_service = repay_plan[repay_plan.subject_type=='serviceFee'].groupby(['放款月份','yinghuan_month']).agg({'应还金额':'sum'}).unstack()

# shihuan_principal = repay_plan[repay_plan.subject_type=='principal'].groupby(['放款月份','yinghuan_month']).agg({'已还金额':'sum'}).unstack()
# shihuan_interest = repay_plan[repay_plan.subject_type=='interest'].groupby(['放款月份','yinghuan_month']).agg({'已还金额':'sum'}).unstack()
# shihuan_service = repay_plan[repay_plan.subject_type=='serviceFee'].groupby(['放款月份','yinghuan_month']).agg({'已还金额':'sum'}).unstack()

# jianmian_principal = repay_plan[repay_plan.subject_type=='principal'].groupby(['放款月份','yinghuan_month']).agg({'减免金额':'sum'}).unstack()
# jianmian_interest = repay_plan[repay_plan.subject_type=='interest'].groupby(['放款月份','yinghuan_month']).agg({'减免金额':'sum'}).unstack()
# jianmian_service = repay_plan[repay_plan.subject_type=='serviceFee'].groupby(['放款月份','yinghuan_month']).agg({'减免金额':'sum'}).unstack()

for table_name in ['yinghuan_principal','yinghuan_interest','yinghuan_service']:
# for table_name in ['yinghuan_principal','yinghuan_interest','yinghuan_service','shihuan_principal','shihuan_interest','shihuan_service','jianmian_principal','jianmian_interest','jianmian_service']:
    current_table = globals()[table_name]
    current_table.columns = [f'amount_{col[1]}' for col in current_table.columns]
    current_table = current_table.rename_axis(columns=None)
    # print(table_name)
    current_table=current_table[['amount_202311','amount_202312','amount_202401','amount_202402','amount_202403','amount_202404','amount_202405','amount_202406','amount_202407','amount_202408','amount_202409','amount_202410','amount_202411','amount_202412','amount_202501']]
    globals()[table_name] = current_table

shihuan_principal_orderno = repay_plan[repay_plan.subject_type=='principal'].groupby(['order_no']).agg({'已还金额':'sum'}).reset_index().rename(columns={'已还金额':'已还本金'})
loan_order=pd.merge(loan_order,shihuan_principal_orderno,how='left',on='order_no')
loan_order['在库']=loan_order.apply(lambda x: 1 if pd.isna(x.settle_date) else 0 ,axis=1)
loan_order['在库本金余额']=loan_order['withdraw_amount']-loan_order['已还本金']
loan_month=loan_order.groupby('放款月份').agg({'withdraw_amount':'sum','order_no':'count','在库':'sum','在库本金余额':'sum'})

loan_month_yinghuan_principal=pd.merge(loan_month,yinghuan_principal,how='left',left_index=True,right_index=True)
loan_month_yinghuan_interest=pd.merge(loan_month,yinghuan_interest,how='left',left_index=True,right_index=True)
loan_month_yinghuan_service=pd.merge(loan_month,yinghuan_service,how='left',left_index=True,right_index=True)

# loan_month_shihuan_principal=pd.merge(loan_month,shihuan_principal,how='left',left_index=True,right_index=True)
# loan_month_shihuan_interest=pd.merge(loan_month,shihuan_interest,how='left',left_index=True,right_index=True)
# loan_month_shihuan_service=pd.merge(loan_month,shihuan_service,how='left',left_index=True,right_index=True)

# loan_month_jianmian_principal=pd.merge(loan_month,jianmian_principal,how='left',left_index=True,right_index=True)
# loan_month_jianmian_interest=pd.merge(loan_month,jianmian_interest,how='left',left_index=True,right_index=True)
# loan_month_jianmian_service=pd.merge(loan_month,jianmian_service,how='left',left_index=True,right_index=True)


for table_name in ['loan_month_yinghuan_principal','loan_month_yinghuan_interest','loan_month_yinghuan_service']:
     
    current_table = globals()[table_name]
    file_path = os.path.join(r'D:\Work\out_data\银信金', f'{table_name}.xlsx')
   # Save DataFrame to Excel
    current_table.to_excel(file_path)
#%%-----------快银付还款卡与pos结算卡的一致性检查；匹配注册手机号与pos端手机号
import pandas as pd
import numpy as np 
import pymysql
from datetime import datetime
import hashlib 

import os
import pymysql
import json
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


sql="""
with white as 
(
select u.id,l.mcht_code,l.brand,
case when u1.id_number is not null then "第一批"
     when u2.id_number is not null then "第二批"
		 else "其他渠道" end as 白名单批次
from juin_loan_core_prd.white_list l 
left join juin_loan_core_prd.user_info u on l.id_number = u.id_number_md5 
left join juin_risk_operate.white_list_kyf1 u1 on l.id_number = u1.id_number
left join juin_risk_operate.white_list_kyf2 u2 on l.id_number = u2.id_number

where l.channel_source = 'kyf' and u.id is not null
),
kkk AS
(
SELECT user_id, bank_no
FROM juin_loan_core_prd.repayment_trade_log
GROUP BY user_id, bank_no
),
oddays as 
(
select  order_id
      ,MAX( if( 
date(settle_time)>0 ,0,DATEDIFF(CURRENT_DATE,repayment_date))) as od_day
 
from  juin_loan_core_prd.repayment_plan_period   
group by 1
 
),
fk as 
(
 select 
    b.id as user_id,
		c.flow_id,
		c.id as order_id,
    CONVERT ( AES_DECRYPT( from_base64 ( b.phone ), '9342266da419cfb4' ) USING utf8 ) as 商户手机号,
    CONVERT ( AES_DECRYPT( from_base64 ( b.name ), '9342266da419cfb4' ) USING utf8 ) as 商户姓名,
    CONVERT ( AES_DECRYPT( from_base64 ( b.id_number ), '9342266da419cfb4' ) USING utf8 ) as 商户身份证,

    case when c.pay_status=20 then c.loan_date else 0 end as 放款日期,
    case when c.pay_status=20 then c.loan_amount else 0 end as 放款金额,
    case when c.pay_status=20 then c.period else 0 end as 合同期限
    
from 
    juin_loan_core_prd.order_record as c
    left join juin_loan_core_prd.user_info as b on c.user_id=b.id

where 
      b.is_delete=0 and c.is_delete=0 and c.user_id!=608608 and c.pay_status=20 and c.channel_id not in (4,6,7)
    and c.user_id in (select id from white )
 )
 select a1.*,b.account_number as 放款卡,b.account_name as 放款卡账户名,c.bank_no as 扣款卡号,d.od_day from
(select 
    a.order_id,
		a.user_id,
    a.商户姓名,
    a.商户手机号,
    a.商户身份证,
    a.放款日期,
    a.合同期限 as 放款期数,
    a.放款金额,
	case when count(a.user_id) over (PARTITION by a.user_id order by a.放款日期) =1 then "Y" else "N" end as 是否首笔贷款,
	c.mcht_code as 商户编号,
	c.brand as 品牌名称,
	c.白名单批次
from fk as a
left join white as c on a.user_id=c.id  ) as a1
left join juin_loan_core_prd.user_bank_account as b on a1.user_id=b.user_id and b.bind_status="BIND_SUCCESS"  and b.is_preferred=1
left join kkk as c on a1.user_id=c.user_id
left join oddays as d on a1.order_id=d.order_id

"""
kyf_k=pd.read_sql(sql,cnx)


jf_mx=pd.read_excel(r'D:\Work\out_data\快银付\快银付放款客户的pos结算卡号.xlsx',sheet_name=r'Sheet1',dtype={'商户编号':str,'结算卡号':str})
jf_mx=jf_mx[['商户编号','结算卡号']]
end=pd.merge(kyf_k,jf_mx,how='left',on='商户编号')

jf_phone=pd.read_excel(r'D:\Work\out_data\快银付\快银付客户的手机号码.xlsx',sheet_name=r'Sheet1',dtype={'mcht_code':str,'phone':str})
end=pd.merge(end,jf_phone,how='left',left_on='商户编号', right_on='mcht_code')
end.to_excel(r'D:\Work\out_data\快银付\放款客户信息.xlsx',index=False)
#%%-----------快银付联系人与客户的重叠度

import pandas as pd
import numpy as np 
import pymysql
from datetime import datetime
import hashlib 

import os
import pymysql
import json
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)



sql = """ 
select a.user_id,a.relationship,a.idx,a.phone,a.phone_md5,a.id_number,a.id_number_md5 from juin_loan_core_prd.user_contact_person as a
left join juin_loan_core_prd.user_info as b  on a.user_id=b.id
left join juin_loan_core_prd.white_list as c   on b.id_number_md5=c.id_number
where c.channel_source="kyf"  
"""
kyf_lxr=pd.read_sql(sql,cnx)


sql = """ 
select a.*  from juin_loan_core_prd.user_info as a
left join  juin_loan_core_prd.white_list as b on a.id_number_md5=b.id_number
where b.channel_source="kyf" 
"""
user_info_kyf=pd.read_sql(sql,cnx)

user_info_kyf_set=set(user_info_kyf.phone_md5)
kyf_lxr_phone=kyf_lxr[kyf_lxr.phone_md5.isin(user_info_kyf_set)]

user_info_kyf_merge=user_info_kyf[['id','phone_md5']].rename(columns={'id':'user_id'})
kyf_lxr_phone_c=pd.merge(kyf_lxr_phone,user_info_kyf_merge,on='phone_md5',how='left')

# select * from order_record where user_id in (810710,791399,810706,524491)  and pay_status=20

# select * from risk_credit_result where user_id in (810710,791399,810706,524491)    and process_node =0

user_info_kyf_set=set(user_info_kyf.id_number_md5)
kyf_lxr_id=kyf_lxr[kyf_lxr.id_number_md5.isin(user_info_kyf_set)]
#%%-----------提现环节验证内部逾期跟内部黑名单字段逻辑

import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
import json
from decimal import Decimal
import pickle


os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
    cnx = pymysql.connect(**cnx_args)


# 提现环节
sql = """ 	 select a.user_id,a.order_id, a.request_body,c.channel_source from 
                (
                 select user_id,order_id,request_body,
                 ROW_NUMBER() over (PARTITION by order_id order by create_time desc,update_time  desc)  as countid
                 FROM juin_loan_core_prd.risk_request_apply 
                 WHERE process_node IN (1,5) 
                 ) as a
										 left join juin_loan_core_prd.user_info as b on a.user_id=b.id
left join juin_loan_core_prd.white_list as c on b.id_number_md5=c.id_number
where   a.countid=1   """   #注意客户授信失败过了冻结期后的再授信问题，还有客户过了缓存期的提现调用三方问题
risk_request_apply=pd.read_sql(sql,cnx)

# 解析urule入参第二代
tables = {}
def for_jiexi(risk_request_apply):
    table_list1=[
    'inputAntiFraudParameter',
    'inputApplyParameter',
    'inputThirdParameter']

    table_list2=[
    'inputBaiRongParameter',
    'inputDianHuaParameter',
    'inputIcekreditParameter',
    'inputOcrParameter',
    'inputRong360Parameter',
    'inputTdParameter',
    'inputTongDunParameter',
    'inputUnionPayParameter',
    'inputWeiYanParameter']


    for i in range(len(risk_request_apply)):
        print(i)
        try:
            temp_json = json.loads(risk_request_apply['request_body'][i])

            for table_name in table_list1:
                table_data = pd.DataFrame(temp_json[table_name], index=[i])

                table_data['user_id'] = risk_request_apply['user_id'][i]
                table_data['order_id'] = risk_request_apply['order_id'][i]
                table_data['channel_source'] = risk_request_apply['channel_source'][i]
                    
                if table_name not in tables:
                    tables[table_name] = table_data
                else:
                    tables[table_name] = pd.concat([tables[table_name], table_data])
        except Exception as e:
            print(f"An error occurred in iteration {i}: {e}")
            continue

    return tables
for_jiexi(risk_request_apply)

with open('D:\\Work\\Information\\zipper\\raw\\tables_tx3.pickle', 'wb') as f:
    pickle.dump(tables, f)

inputApplyParameter=pd.DataFrame(tables['inputApplyParameter'])
js=inputApplyParameter[['channel','custMaxOvddays','custOvdDays','user_id','order_id','channel_source']]
a=js.describe()
js.custOvdDays.value_counts()
inputAntiFraudParameter=pd.DataFrame(tables['inputAntiFraudParameter'])
js1=inputAntiFraudParameter[['user_id','order_id','channel_source','link1HnOvdDays','link1HnMaxovdDays','link2HnMaxovdDays','link2HnOvdDays','linkHnMaxovdDays','linkHnOvdDays',
                            'blBankcard','blCname' ,'blIdno','blTel','blLinktel1','blLinktel2','blBankTel'
                             ]]
for var in ['link2HnMaxovdDays','link2HnOvdDays','linkHnMaxovdDays','linkHnOvdDays','blBankcard','blCname' ,'blIdno','blTel','blLinktel1','blLinktel2','blBankTel']:
    agroup=js1.groupby(var)['user_id'].count().reset_index().rename(columns={var:'分组'})
    agroup['变量名']=var
    if var=='link2HnMaxovdDays':
        temp=agroup
    else:
        temp=pd.concat([temp,agroup])

js1.blBankTel.value_counts()
agroup=js1.groupby('link2HnMaxovdDays')['user_id'].count
 








#%%-----------彩虹

import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
import json
from decimal import Decimal
import pickle

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


# 提现环节
sql = """ 	 select a.user_id,a.flow_id,a.order_id,a.request_body,c.channel_source from 
                (
                 select user_id,flow_id,request_body,order_id,
                 ROW_NUMBER() over (PARTITION by flow_id order by create_time desc,update_time  desc)  as countid
                 FROM juin_loan_core_prd.risk_request_apply 
                 WHERE process_node IN (1,5) 
                 ) as a
										 left join juin_loan_core_prd.user_info as b on a.user_id=b.id
left join juin_loan_core_prd.white_list as c on b.id_number_md5=c.id_number
where   a.countid=1   """  
risk_request_apply=pd.read_sql(sql,cnx)

# 贷后表现：放款月,fpd,order_id,flow_id

sql="""
select
		a.flow_id,
		a.user_id,
		a.id as order_id, 
		loan_date,
		a.period,
		DATE_FORMAT(loan_date,'%Y-%m') as loan_mon,
		loan_amount,
		case when max_od_day >=7 then 1 else 0 end as dpd7,
		case when f_od>=7  then 1 else 0 end as fpd7_fz,
		case when DATEDIFF(NOW(),fpd.repayment_date)>=7 then 1 else 0 end fpd7_fm

from juin_loan_core_prd.order_record as a 
left join (SELECT*,case when status not in (0) then  DATEDIFF(IFNULL(settle_time,NOW()),repayment_date) end as f_od
						from juin_loan_core_prd.repayment_plan_period where period_number=1 ) as fpd on a.id=fpd.order_id
															
left join (SELECT order_id,
							 max( DATEDIFF(IFNULL(settle_time,NOW()),repayment_date) )  as max_od_day           
							from juin_loan_core_prd.repayment_plan_period
							where is_delete=0 and status not in (0) group by 1 order by 1) as od on a.id=od.order_id
where a.pay_status in (20) 
"""
dh_behave=pd.read_sql(sql,cnx)


# 解析urule入参
tables = {}
def for_jiexi(risk_request_apply):
    table_list1=[
    'inputAntiFraudParameter',
    'inputApplyParameter',
    'inputThirdParameter']

    table_list2=[
    'inputBaiRongParameter',
    'inputDianHuaParameter',
    'inputIcekreditParameter',
    'inputOcrParameter',
    'inputRong360Parameter',
    'inputTdParameter',
    'inputTongDunParameter',
    'inputUnionPayParameter',
    'inputWeiYanParameter']


    for i in range(len(risk_request_apply)):
        print(i)
        try:
            temp_json = json.loads(risk_request_apply['request_body'][i])

            for table_list, table_prefix in [(table_list1, ''), (table_list2, 'n')]:
                for table_name in table_list:
                    if table_name in table_list1:
                        table_data = pd.DataFrame(temp_json[table_name], index=[i])
                    else:
                        if table_name in table_list2:
                            table_data = pd.DataFrame(temp_json['inputThirdParameter'][table_name], index=[i])

                    table_data['user_id'] = risk_request_apply['user_id'][i]
                    table_data['flow_id'] = risk_request_apply['flow_id'][i]
                    table_data['order_id'] = risk_request_apply['order_id'][i]
               

                    if table_name not in tables:
                        tables[table_name] = table_data
                    else:
                        tables[table_name] = pd.concat([tables[table_name], table_data])
        except Exception as e:
            print(f"An error occurred in iteration {i}: {e}")
            continue

    return tables
for_jiexi(risk_request_apply)

#%%-----------银信金催收

import pandas as pd
import numpy as np 
import pymysql
from datetime import datetime
import hashlib 
import datetime
import os
import pymysql
import json
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import pickle


os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# a_temp=pd.DataFrame(tables['inputIcekreditParameter'])
# a_temp_set=set(a_temp.user_id)

# 存在部分还款但是settle_time有值的，待优化
sql = """ 	select  user_id,order_no as order_id,receivable_amount as period_total_amount,receivable_time as repayment_date,repayment_time as settle_time from loan_core.core_repayment_plan     """
repayment_plan_period=pd.read_sql(sql,cnx)
repayment_plan_period['repayment_date'] = pd.to_datetime(repayment_plan_period['repayment_date']).dt.date
repayment_plan_period['settle_time'] = pd.to_datetime(repayment_plan_period['settle_time']).dt.date


repayment_cutdate = pd.DataFrame()
for (i,cut_date) in enumerate(pd.date_range(datetime.date(2023,10,1),
                                            datetime.date(2024,3,18),
                                            closed='left')):
    print(i)
    cut_date=cut_date.date()
    
    assign = repayment_plan_period.loc[repayment_plan_period['repayment_date']<=cut_date,:]
    
    assign['od_days']=assign.apply(lambda x :-999 if (pd.notnull(x.settle_time) and ( x['settle_time']<=x['repayment_date'] or  x['settle_time'] <cut_date ))
                                                    else (cut_date -  x['repayment_date']).days ,axis=1)
    assign_bad=assign.groupby('order_id').agg({'od_days':'max'}).reset_index()
    assign_bad=assign_bad[assign_bad.od_days>0]
    assign['assign_bad_amount']=assign.apply(lambda x: x.period_total_amount if x.od_days>=0 else 0,axis=1 )#细节，如果客户连续逾期，下一次账单应该要合并当前期数的
    assign_bad_set=set(assign_bad.order_id)
    assign_bad=assign[assign.order_id.isin(assign_bad_set)].groupby('order_id').agg({'od_days':'max','assign_bad_amount':'sum'}).reset_index().rename(columns={'assign_bad_amount':'period_total_amount'})


    assign_good=assign.loc[~assign.order_id.isin(assign_bad_set),['order_id','od_days','period_total_amount','repayment_date']]
    assign_good.sort_values(['order_id','repayment_date'],inplace=True)
    assign_good.drop_duplicates('order_id',keep='last',inplace=True)
    assign_temp=pd.concat([assign_good,assign_bad],axis=0)
    assign_temp['cut_date'] = cut_date
    if i==0:
        repayment_cutdate = assign_temp
    else:
        repayment_cutdate = pd.concat([repayment_cutdate,assign_temp])


repayment_cutdate_copy=repayment_cutdate.copy()

repayment_cutdate.drop(columns='repayment_date',inplace=True)

# 拼接客户的结清日期
repayment_plan_period_st=repayment_plan_period.loc[~pd.isnull(repayment_plan_period.settle_time),['order_id','settle_time']]
repayment_plan_period_st.drop_duplicates(['order_id','settle_time'],keep='last',inplace=True)
repayment_plan_period_st['还款']=1
repayment_cutdate['cut_date'] = pd.to_datetime(repayment_cutdate['cut_date']).dt.date
repayment_plan_period_st['settle_time'] = pd.to_datetime(repayment_plan_period_st['settle_time']).dt.date
repayment_cutdate=pd.merge(repayment_cutdate,repayment_plan_period_st,how='left',left_on=['order_id','cut_date'],right_on=['order_id','settle_time'])


# 拼接客户的账单日
repayment_plan_period_rdate=repayment_plan_period[['order_id','repayment_date']]
repayment_plan_period_rdate.drop_duplicates(['order_id','repayment_date'],keep='last',inplace=True)
repayment_plan_period_rdate['账单日']=1
repayment_cutdate['cut_date'] = pd.to_datetime(repayment_cutdate['cut_date']).dt.date
repayment_plan_period_rdate['repayment_date'] = pd.to_datetime(repayment_plan_period_rdate['repayment_date']).dt.date
repayment_cutdate=pd.merge(repayment_cutdate,repayment_plan_period_rdate,how='left',left_on=['order_id','cut_date'],right_on=['order_id','repayment_date'])


def date_status(x):
    if x==0:
        return 'C'
    elif x>0 and x<=30:
        return 'M1'
    elif x>30 and x<=60:
        return "M2"
    elif x>60 and x<=90:
        return 'M3'
    elif x>90  :
        return 'M3+'
    
repayment_cutdate['cut_status']=repayment_cutdate.od_days.apply(date_status)
repayment_cutdate['cut_month']= pd.to_datetime(repayment_cutdate.cut_date).dt.strftime('%Y%m')
repayment_cutdate['cut_date']= pd.to_datetime(repayment_cutdate.cut_date).dt.date


with open('D:\\Work\\out_data\\分析类\\日常\\repayment_cutdate_yxj.pickle', 'wb') as f:
    pickle.dump(repayment_cutdate, f)#20240319封板
    
# a=repayment_cutdate[(repayment_cutdate.cut_status=='M3') & (repayment_cutdate.cut_month=='202403') ] 
# a1=repayment_plan_period[repayment_plan_period.order_id=='DY20230320293043963651297336']
# a2=repayment_cutdate[repayment_cutdate.order_id=='DY20230320293043963651297336']
# 计算-全量
temp1=repayment_cutdate[repayment_cutdate.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate[repayment_cutdate.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate[repayment_cutdate.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate[(repayment_cutdate.cut_status!='C') & (repayment_cutdate.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})

temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_银信金系统.xlsx',index=False)


a=repayment_cutdate[(repayment_cutdate.还款==1)  &  (repayment_cutdate.cut_month=='202401') & (repayment_cutdate.cut_status=='M3')]
a.sort_values(by='od_days',ascending=False,inplace=True)
a1=repayment_cutdate[repayment_cutdate.order_id=="DY20230620326397104803156061"]
a2=repayment_plan_period[repayment_plan_period.order_id=="DY20230620326397104803156061"]


#%%-----------钜银催收

import pandas as pd
import numpy as np 
import pymysql
from datetime import datetime
import datetime
import hashlib 
import os
import pymysql
import json
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import pickle


os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# a_temp=pd.DataFrame(tables['inputIcekreditParameter'])
# a_temp_set=set(a_temp.user_id)


sql = """ 	select  user_id,order_id,period_total_amount,period_principal_amount,period_interest_amount,repayment_date,settle_time,is_overdue from juin_loan_core_prd.repayment_plan_period     """
repayment_plan_period=pd.read_sql(sql,cnx)
repayment_plan_period['repayment_date'] = pd.to_datetime(repayment_plan_period['repayment_date']).dt.date
repayment_plan_period['settle_time'] = pd.to_datetime(repayment_plan_period['settle_time']).dt.date


repayment_cutdate = pd.DataFrame()
assign = pd.DataFrame()
for (i,cut_date) in enumerate(pd.date_range(datetime.date(2023,12,16),
                                            datetime.date(2024,3,18),#这里日期要更新到跑批的昨天
                                            closed='left')):
    print(i)
    cut_date=cut_date.date()
 
    assign = repayment_plan_period.loc[repayment_plan_period['repayment_date']<=cut_date,:]
    assign['od_days']=assign.apply(lambda x :-999 if (pd.notnull(x.settle_time) and ( x['settle_time']<=x['repayment_date'] or  x['settle_time'] <cut_date ))
                                                    else (cut_date -  x['repayment_date']).days ,axis=1)
        
        
    assign_bad=assign.groupby('order_id').agg({'od_days':'max'}).reset_index()
    assign_bad=assign_bad[assign_bad.od_days>0]
    assign['assign_bad_amount']=assign.apply(lambda x: x.period_total_amount if x.od_days>=0 else 0,axis=1 )
    assign_bad_set=set(assign_bad.order_id)
    assign_bad=assign[assign.order_id.isin(assign_bad_set)].groupby('order_id').agg({'od_days':'max','assign_bad_amount':'sum'}).reset_index().rename(columns={'assign_bad_amount':'period_total_amount'})

    assign_good=assign.loc[~assign.order_id.isin(assign_bad_set),['order_id','od_days','period_total_amount','repayment_date']]
    assign_good.sort_values(['order_id','repayment_date'],inplace=True)
    assign_good.drop_duplicates('order_id',keep='last',inplace=True)
    assign_temp=pd.concat([assign_good,assign_bad],axis=0)
    assign_temp['cut_date'] = cut_date
    if i==0:
        repayment_cutdate = assign_temp
    else:
        repayment_cutdate = pd.concat([repayment_cutdate,assign_temp])


repayment_cutdate_copy=repayment_cutdate.copy()

repayment_cutdate.drop(columns='repayment_date',inplace=True)

# 拼接客户的结清日期
repayment_plan_period_st=repayment_plan_period.loc[~pd.isnull(repayment_plan_period.settle_time),['order_id','settle_time']]
repayment_plan_period_st.drop_duplicates(['order_id','settle_time'],keep='last',inplace=True)
repayment_plan_period_st['还款']=1
repayment_cutdate['cut_date'] = pd.to_datetime(repayment_cutdate['cut_date']).dt.date
repayment_plan_period_st['settle_time'] = pd.to_datetime(repayment_plan_period_st['settle_time']).dt.date
repayment_cutdate=pd.merge(repayment_cutdate,repayment_plan_period_st,how='left',left_on=['order_id','cut_date'],right_on=['order_id','settle_time'])


# 拼接客户的账单日
repayment_plan_period_rdate=repayment_plan_period[['order_id','repayment_date']]
repayment_plan_period_rdate.drop_duplicates(['order_id','repayment_date'],keep='last',inplace=True)
repayment_plan_period_rdate['账单日']=1
repayment_cutdate['cut_date'] = pd.to_datetime(repayment_cutdate['cut_date']).dt.date
repayment_plan_period_rdate['repayment_date'] = pd.to_datetime(repayment_plan_period_rdate['repayment_date']).dt.date
repayment_cutdate=pd.merge(repayment_cutdate,repayment_plan_period_rdate,how='left',left_on=['order_id','cut_date'],right_on=['order_id','repayment_date'])

# 拼接channel_id channel_source
sql = """ 	select  a.id as order_id,a.channel_id,a.funder_id,
c.channel_source 
from juin_loan_core_prd.order_record as a
left join juin_loan_core_prd.user_info as b on a.user_id=b.id
left join juin_loan_core_prd.white_list as c on b.id_number_md5=c.id_number
where a.pay_status=20  """
order_record=pd.read_sql(sql,cnx)

repayment_cutdate=pd.merge(repayment_cutdate,order_record,how='left',on='order_id')



def date_status(x):
    if x==0:
        return 'C'
    elif x>0 and x<=30:
        return 'M1'
    elif x>30 and x<=60:
        return "M2"
    elif x>60 and x<=90:
        return 'M3'
    elif x>90  :
        return 'M3+'
    
repayment_cutdate['cut_status']=repayment_cutdate.od_days.apply(date_status)
repayment_cutdate['cut_month']= pd.to_datetime(repayment_cutdate.cut_date).dt.strftime('%Y%m')
repayment_cutdate['cut_date']= pd.to_datetime(repayment_cutdate.cut_date).dt.date


with open('D:\\Work\\out_data\\分析类\\日常\\repayment_cutdate_jyyk.pickle', 'wb') as f:
    pickle.dump(repayment_cutdate, f)#20240319封板
    
a=repayment_cutdate[(repayment_cutdate.cut_status=='M2') & (repayment_cutdate.cut_month=='202401') ] 
a1=repayment_plan_period[repayment_plan_period.order_id==274]
a2=repayment_cutdate[repayment_cutdate.order_id==274]


# 计算-全量
temp1=repayment_cutdate[repayment_cutdate.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate[repayment_cutdate.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate[repayment_cutdate.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate[(repayment_cutdate.cut_status!='C') & (repayment_cutdate.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})

temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_全量.xlsx',index=False)


# 计算-玖富
repayment_cutdate_jf=repayment_cutdate[repayment_cutdate.channel_id==4]
temp1=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C' ].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate_jf[(repayment_cutdate_jf.cut_status!='C') & (repayment_cutdate_jf.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})


temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_玖富.xlsx',index=False)


# 计算-快银付
repayment_cutdate_kyf=repayment_cutdate[repayment_cutdate.channel_source=='kyf']
temp1=repayment_cutdate_kyf[repayment_cutdate_kyf.cut_status!='C' ].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate_kyf[repayment_cutdate_kyf.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate_kyf[repayment_cutdate_kyf.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate_kyf[(repayment_cutdate_kyf.cut_status!='C') & (repayment_cutdate_kyf.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})


temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_快银付.xlsx',index=False)


# 计算-银信金
repayment_cutdate_yxj=repayment_cutdate[repayment_cutdate.channel_source=='yxj']
temp1=repayment_cutdate_yxj[repayment_cutdate_yxj.cut_status!='C' ].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate_yxj[repayment_cutdate_yxj.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate_yxj[repayment_cutdate_yxj.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate_yxj[(repayment_cutdate_yxj.cut_status!='C') & (repayment_cutdate_yxj.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})


temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_银信金.xlsx',index=False)


# 计算-星光
repayment_cutdate_jf=repayment_cutdate[repayment_cutdate.funder_id==1]
temp1=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C' ].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate_jf[(repayment_cutdate_jf.cut_status!='C') & (repayment_cutdate_jf.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})


temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_星光.xlsx',index=False)

# 计算-民和
repayment_cutdate_jf=repayment_cutdate[repayment_cutdate.funder_id==2]
temp1=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C' ].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate_jf[(repayment_cutdate_jf.cut_status!='C') & (repayment_cutdate_jf.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})


temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_民和.xlsx',index=False)


# 计算-星光_2月22日
repayment_cutdate_jf=repayment_cutdate[(repayment_cutdate.funder_id==1 )& (repayment_cutdate.cut_date>=datetime.date(2024,2,22))]
temp1=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C' ].groupby(['cut_month','cut_status']).agg({'order_id':lambda x: x.nunique()}).reset_index().rename(columns={'order_id':'笔数'})
temp1_c=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].groupby(['cut_month','cut_status']).agg({'还款':'sum'}).reset_index().rename(columns={'还款':'回收笔数'})
temp2=repayment_cutdate_jf[repayment_cutdate_jf.cut_status!='C'].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'金额'})
temp2_c=repayment_cutdate_jf[(repayment_cutdate_jf.cut_status!='C') & (repayment_cutdate_jf.还款==1)].drop_duplicates(subset=['order_id','cut_month','cut_status']).groupby(['cut_month','cut_status']).agg({'period_total_amount':'sum'}).reset_index().rename(columns={'period_total_amount':'回收金额'})


temp=pd.merge(temp1,temp2,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp1_c,how='left',on=['cut_month','cut_status'])
temp=pd.merge(temp,temp2_c,how='left',on=['cut_month','cut_status'])   
temp['笔数回收率']=temp['回收笔数']/temp['笔数']
temp['金额回收率']=temp['回收金额']/temp['金额']
temp.to_excel(r'D:\Work\out_data\分析类\日常\催收效果_星光22.xlsx',index=False)



    
#%%-----------数据库验证

import pandas as pd
import numpy as np 
import pymysql
from datetime import datetime
import hashlib 

import os
import pymysql
import json
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# -------------------------------------------------------------------------基础表


sql = """ 
select a.*, JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信审批结果')) AS 审批结果,
JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信额度')) AS 审批额度 from 
(
 select *,count(user_id) over (PARTITION by user_id order by create_time  desc) as countid

 FROM juin_loan_core_prd.risk_request_apply 
 WHERE process_node IN (0,4) 
 ) as a
where a.countid=1
"""
risk_apply_sx_new=pd.read_sql(sql,cnx)


sql = """ 
select * from juin_loan_core_prd.risk_credit_result
"""
risk_credit_result=pd.read_sql(sql,cnx)

sql = """ 
select * from juin_loan_core_prd.user_info
"""
user_info=pd.read_sql(sql,cnx)

sql = """ 
select * from juin_loan_core_prd.white_list
"""
white_list=pd.read_sql(sql,cnx)


sql = """ 
select * from juin_loan_core_prd.manual_credit_audit_info
"""
manual_credit_audit_info=pd.read_sql(sql,cnx)



# --------------------------------------------------------------------------------------------验证快银付人审拒绝后result怎么存
# 结论：
# 1、remark有值，可能是星光拒绝，也可能是人审拒绝（manual_credit_audit_info.reject_reason）
# 2、如果remark有值，risk_credit_result会将credit_result改为拒绝，credit_limit清0
# 3、人审拒绝案例：user_Id=805787
white_list_kyf=white_list[white_list.channel_source=='kyf']
user_info_kyf=user_info.loc[user_info.id_number_md5.isin(white_list_kyf.id_number),['id_number_md5','id']].rename(columns={'id':'user_id'})
risk_apply_sx_new_kyf=risk_apply_sx_new[risk_apply_sx_new.user_id.isin(user_info_kyf.user_id)]
print(risk_apply_sx_new_kyf.shape)
risk_apply_sx_new_kyf=pd.merge(risk_apply_sx_new_kyf,user_info_kyf,on='user_id',how='left')
print(risk_apply_sx_new_kyf.shape)
manual_credit_audit_info_sx=manual_credit_audit_info.loc[manual_credit_audit_info.audit_type=="1",['user_id','audit_status']]
risk_apply_sx_new_kyf=pd.merge(risk_apply_sx_new_kyf,manual_credit_audit_info_sx,on='user_id',how='left')
print(risk_apply_sx_new_kyf.shape)
risk_credit_result_yanzheng=risk_credit_result.loc[risk_credit_result.process_node==0,['user_id','credit_result','remark','credit_limit']]
risk_apply_sx_new_kyf=pd.merge(risk_apply_sx_new_kyf,risk_credit_result_yanzheng,on='user_id',how='left')
print(risk_apply_sx_new_kyf.shape)





# a=manual_credit_audit_info[manual_credit_audit_info.user_id==805787]



risk_credit_result.columns







#%%-----------快银付存量白名单有授信记录

import pandas as pd
import numpy as np 
import pymysql
from datetime import datetime
import hashlib 

import os
import pymysql
import json
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# -------------------------------------------------------------------------基础表


sql = """ 
select a.*, JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信审批结果')) AS 审批结果,
JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信额度')) AS 审批额度 from 
(
 select *,count(user_id) over (PARTITION by user_id order by create_time  desc) as countid

 FROM juin_loan_core_prd.risk_request_apply 
 WHERE process_node IN (0,4) 
 ) as a
where a.countid=1
"""
risk_apply_sx_new=pd.read_sql(sql,cnx)


sql = """ 
select * from juin_loan_core_prd.risk_credit_result
"""
risk_credit_result=pd.read_sql(sql,cnx)

sql = """ 
select * from juin_loan_core_prd.user_info
"""
user_info=pd.read_sql(sql,cnx)

sql = """ 
select * from juin_loan_core_prd.white_list
"""
white_list=pd.read_sql(sql,cnx)


sql = """ 
select * from juin_loan_core_prd.manual_credit_audit_info
"""
manual_credit_audit_info=pd.read_sql(sql,cnx)

white_list_kyf=white_list[white_list.channel_source=='kyf']
user_info_kyf=user_info.loc[user_info.id_number_md5.isin(white_list_kyf.id_number),['id_number_md5','id']].rename(columns={'id':'user_id'})
risk_credit_result_kyf=risk_credit_result.loc[(risk_credit_result.user_id.isin(user_info_kyf.user_id)) & (risk_credit_result.process_node==0) & (risk_credit_result.credit_result=='通过'),['user_id','credit_result']]
print(risk_credit_result_kyf.shape)
risk_credit_result_kyf=pd.merge(risk_credit_result_kyf,user_info_kyf,on='user_id',how='left')
print(risk_credit_result_kyf.shape)

white_list_kyf=pd.merge(white_list_kyf,risk_credit_result_kyf,left_on='id_number',right_on='id_number_md5',how='left')
white_list_kyf_end=white_list_kyf[white_list_kyf.id_number_md5.isna()]

white_list_kyf_end['新期限']=white_list_kyf_end.apply(lambda x: "3" if x.credit_limit<6000 else x.loan_term,axis=1)
white_list_kyf_end['是否发生改动']=white_list_kyf_end.apply(lambda x: 1 if x.新期限!=x.loan_term  else 0,axis=1)

end=white_list_kyf_end.loc[white_list_kyf_end['是否发生改动']==1,['id_number','新期限','channel_source']].rename(columns={'新期限':'loan_term'})
end.to_excel(r'D:\Work\out_data\快银付\白名单\邮件存档\快银付存量白名单期限调整.xlsx',index=False)

# white_list_kyf_end.duplicated(subset='id_number').any()#判断pd是否有重复值


#%%-----------快银付7万匹配

import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


base7=pd.read_excel(r'D:\Work\out_data\快银付\快银付测试用户画像.xlsx',sheet_name=r'数据明细')
a=base7.columns
base7.duplicated(subset='id_card').any()#判断pd是否有重复值
base7
duplicated_rows = base7[base7.duplicated(subset='id_card', keep=False)]#取出重复值

import hashlib
def encrypt_to_md5(name):
    # 将姓名编码为UTF-8格式的字节串
    name_bytes = name.encode('utf-8')
    
    # 创建MD5对象
    md5_hash = hashlib.md5()
    
    # 更新MD5对象的内容
    md5_hash.update(name_bytes)
    
    # 获取MD5哈希值的十六进制表示
    md5_digest = md5_hash.hexdigest()
    
    return md5_digest

base7['name_md5']=base7.name.apply(encrypt_to_md5)
base7['id_card_md5']=base7.id_card.apply(encrypt_to_md5)

def phone_to_md5(phone_number):
    # 将手机号转换为字符串类型
    phone_str = str(phone_number)
    
    # 将手机号编码为UTF-8格式的字节串
    phone_bytes = phone_str.encode('utf-8')
    
    # 创建MD5对象
    md5_hash = hashlib.md5()
    
    # 更新MD5对象的内容
    md5_hash.update(phone_bytes)
    
    # 获取MD5哈希值的十六进制表示
    md5_digest = md5_hash.hexdigest()
    
    return md5_digest
base7['phone_md5']=base7.phone.apply(phone_to_md5)


  

sql = """ select
	   name_md5,id_number_md5,phone_md5
	from 
        juin_loan_core_prd.user_info where is_delete=0 and  id in (636722,636308)"""
user_kyf=pd.read_sql(sql,cnx)

base7_name=base7[base7.name_md5.isin(user_kyf.name_md5)]
base7_phone=base7[base7.phone_md5.isin(user_kyf.phone_md5)]
base7_id=base7[base7.id_card_md5.isin(user_kyf.id_number_md5)]


base7.drop_duplicates(subset='id_card',keep='last',inplace=True)

base7_out=base7[['name_md5','phone_md5','id_card_md5','one_active_months','three_active_months','six_active_months','twelve_active_months']]
base7_out.to_excel(r'D:\Work\out_data\快银付\白名单\7万\base7.xlsx',index=False)


base7_out=base7[['name_md5','phone_md5','id_card_md5','name','phone','id_card']]
base7_out.to_excel(r'D:\Work\out_data\快银付\白名单\7万\base7_验.xlsx',index=False)

base43=pd.read_excel(r'D:\Work\out_data\快银付\白名单\20240129白名单.xlsx',sheet_name=r'129邮件_数据库版本')#有一个客户的MD5码是大小写，但是py又区分大小写
base7_43=base7[['id_card_md5','name','phone','id_card','mcht_code','brand']]#存在客户对应多个brand，待0204再问问
base43a=pd.merge(base43,base7_43,how='left',left_on='id_number_md5',right_on='id_card_md5')
base43a.drop_duplicates(subset='id_number_md5',keep='last',inplace=True)
base7_out.to_excel(r'D:\Work\out_data\快银付\白名单\20240129白名单_真实3要素版.xlsx',index=False)



duplicated_rows = base43a[base43a.duplicated(subset='id_number_md5', keep=False)]#取出重复值
base43.duplicated(subset='id_number_md5').any()#判断pd是否有重复值

a1=base7.sample(5)

base7.duplicated(subset='id_card').any()#判断pd是否有重复值



# 20240204白名单额度调整
whitelist=pd.read_excel(r'D:\Work\out_data\快银付\白名单\20240126快银付第一批白名单_id_amount.xlsx',sheet_name=r'Sheet1')

import hashlib
def encrypt_to_md5(name):
    # 将姓名编码为UTF-8格式的字节串
    name_bytes = name.encode('utf-8')
    
    # 创建MD5对象
    md5_hash = hashlib.md5()
    
    # 更新MD5对象的内容
    md5_hash.update(name_bytes)
    
    # 获取MD5哈希值的十六进制表示
    md5_digest = md5_hash.hexdigest()
    
    return md5_digest

whitelist['id_card_md5']=whitelist.id_card.apply(encrypt_to_md5)
whitelist=whitelist[['id_card_md5','按宣传区间下线调整']]
base43=pd.read_excel(r'D:\Work\out_data\快银付\白名单\20240129白名单.xlsx',sheet_name=r'129邮件_数据库版本')#有一个客户的MD5码是大小写，但是py又区分大小写
base43=base43[base43.渠道来源=="kyf"]
base43=base43[base43.备注!="快银付内部测试名单"]
base43=pd.merge(base43,whitelist,how='left',left_on='id_number_md5',right_on='id_card_md5')

sql="""
with risk_apply_kyf_list as 
(
select distinct user_id  from 
(
select  
    user_id,
    JSON_UNQUOTE(JSON_EXTRACT(request_body, '$.inputApplyParameter.channel')) AS urule渠道,
	JSON_UNQUOTE(JSON_EXTRACT(request_body, '$.inputApplyParameter.yxLoanCustomer')) AS 是否银信放款客户,
    JSON_UNQUOTE(JSON_EXTRACT(request_body, '$.inputApplyParameter.whiteList')) AS 是否白名单
from 
   juin_loan_core_prd. risk_request_apply 
where 
    process_node=0 and is_delete=0 and channel_id not in (4,6,7)
) as a where a.是否白名单=1 and a.是否银信放款客户=0
),
sx as 
(
 select 
     b.id as user_id,
     CONVERT ( AES_DECRYPT( from_base64 ( b.id_number ), '9342266da419cfb4' ) USING utf8 ) as 商户身份证,

     a.submit_time as 授信日期,
     JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信审批结果')) AS 授信审批结果,
     CAST(JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信额度')) AS DECIMAL(10, 2)) AS 授信额度,
     
     case when c.remark is not null then "拒绝" else JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信审批结果')) end as 星光审批结果
     
 from 
     juin_loan_core_prd.risk_request_apply as a
     left join juin_loan_core_prd.user_info as b on a.user_id=b.id
     left join juin_loan_core_prd.risk_credit_result as c on a.user_id=c.user_id
 where 
     a.is_delete=0 and b.is_delete=0 and c.is_delete=0 and a.user_id!=608608
     and a.process_node=0 and a.user_id in (select user_id from risk_apply_kyf_list)
     and c.process_node=0

 )
 select * from sx
"""

sx=pd.read_sql(sql,cnx)

sx['id_card_md5']=sx.商户身份证.apply(encrypt_to_md5)
sx['有授信记录']=1
sx=[['id_card_md5','有授信记录']]



# 20240202涛哥想看进来的这些人的用户画像
sql="""
with risk_apply_kyf_list as 
(
select distinct user_id  from 
(
select  
    user_id,
    JSON_UNQUOTE(JSON_EXTRACT(request_body, '$.inputApplyParameter.channel')) AS urule渠道,
	JSON_UNQUOTE(JSON_EXTRACT(request_body, '$.inputApplyParameter.yxLoanCustomer')) AS 是否银信放款客户,
    JSON_UNQUOTE(JSON_EXTRACT(request_body, '$.inputApplyParameter.whiteList')) AS 是否白名单
from 
    juin_loan_core_prd.risk_request_apply 
where 
    process_node=0 and is_delete=0 and channel_id not in (4,6,7)
) as a where a.是否白名单=1 and a.是否银信放款客户=0
),
sx as 
(
 select 
     b.id as user_id,
     CONVERT ( AES_DECRYPT( from_base64 ( b.phone ), '9342266da419cfb4' ) USING utf8 ) as 商户手机号,
     CONVERT ( AES_DECRYPT( from_base64 ( b.name ), '9342266da419cfb4' ) USING utf8 ) as 商户姓名,
     CONVERT ( AES_DECRYPT( from_base64 ( b.id_number ), '9342266da419cfb4' ) USING utf8 ) as 商户身份证,
     a.submit_time as 授信日期,
     JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信审批结果')) AS 授信审批结果,
     CAST(JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信额度')) AS DECIMAL(10, 2)) AS 授信额度,
		 c.reason_code,
		 c.remark,
     
     case when c.remark is not null then "拒绝" else JSON_UNQUOTE(JSON_EXTRACT(a.response_body, '$.授信审批结果')) end as 星光审批结果
     
 from 
     juin_loan_core_prd.risk_request_apply as a
     left join juin_loan_core_prd.user_info as b on a.user_id=b.id
     left join juin_loan_core_prd.risk_credit_result as c on a.user_id=c.user_id
 where 
     a.is_delete=0 and b.is_delete=0 and c.is_delete=0 and a.user_id!=608608
     and a.process_node=0 and a.user_id in (select user_id from risk_apply_kyf_list)
     and c.process_node=0

 )
 select * from sx  

 """
sx=pd.read_sql(sql,cnx)

sx1=base7[base7.id_card.isin(sx.商户身份证)]
sx1.duplicated(subset='id_card').any()#判断pd是否有重复值


bins =[float('-inf'),20,25,30,35,40,45,50,55,float('inf')]
temp_intervals = pd.cut(sx1['age'], bins=bins, ordered=True)
sx1['age'+'_intervals'] = temp_intervals

sx1['age'+'_intervals'].value_counts().to_excel(r'D:\Work\out_data\快银付\授信客户年龄分布.xlsx')
sx1['city'].value_counts().to_excel(r'D:\Work\out_data\快银付\temp.xlsx')
sx1_amount=sx1[['one_amt','three_amt','six_amt','twelve_amt']]
sx1_amount.describe().to_excel(r'D:\Work\out_data\快银付\temp.xlsx')
#%%-----------20240229 给周总的放款跟已收权益
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ 
with fk_jiufu as 
(
SELECT
    order_no,
    DATE_FORMAT( loan_date, '%Y-%m-%d' )AS 放款时间,
     loan_amount  as 放款金额,
     period as 期限,include_benefit as 是否购买权益,actual_rate as 定价
from 
    juin_loan_core_prd.order_record  
where pay_status in (20) and channel_id=4
),
 quanyi as 
(
    select
	    DISTINCT order_no,
		collection_time as 收款时间,
	    equity_amount as 权益金额
	     
	from 
        juin_loan_core_prd.equity_record where is_delete=0
)
    select 
        a.*,
		c.收款时间,
		c.权益金额
	 
	from 
	    fk_jiufu as a
	left join 
	    quanyi as c on a.order_no=c.order_no	 """
fk_qy=pd.read_sql(sql,cnx)


fk=fk_qy[['order_no','放款时间','放款金额','期限','是否购买权益','定价']]
fk.sort_values(by='放款时间',inplace=True)
quanyi=fk_qy[fk_qy.权益金额.notna()]
quanyi.sort_values(by='收款时间',inplace=True)

fk.to_excel(r'D:\Work\out_data\权益\玖富对账\放款20240229.xlsx',index=False)
quanyi.to_excel(r'D:\Work\out_data\权益\玖富对账\已收权益20240229.xlsx',index=False)



#%%-----------20240202与玖富对账权益金额

import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


# 权益金额
sql = """ select
	    DISTINCT order_no,
		collection_time as 收款时间,
	    equity_amount as 权益金额,
	    equity_source as 权益来源,
	    amount_type as 收退款类型,
	    refund_amount as 退款金额,
	    refund_time as 退款时间
	from 
        juin_loan_core_prd.equity_record where is_delete=0 """
equity_record=pd.read_sql(sql,cnx)

equity_record=pd.merge(equity_record,dz,how='left',left_on='order_no',right_on='借据')




dz=pd.read_excel(r'D:\Work\out_data\权益\玖富对账\提供矩银_权益实还.xlsx',sheet_name=r'Sheet1')

dz=pd.merge(dz,equity_record,how='left',right_on='order_no',left_on='借据')
dz.to_excel(r'D:\Work\out_data\权益\玖富对账\temp.xlsx')

# 20240223给玖富之前对台账时发现缺少的订单
lista=pd.read_excel(r'D:\Work\out_data\权益\玖富对账\20240204对账差异.xlsx',sheet_name=r'Sheet1')


# 权益金额
sql = """ select
	    DISTINCT order_no,
		collection_time as 收款时间,
	    equity_amount as 权益金额,
	    equity_source as 权益来源,
	    amount_type as 收退款类型,
	    refund_amount as 退款金额,
	    refund_time as 退款时间
	from 
        juin_loan_core_prd.equity_record where is_delete=0 """
equity_record=pd.read_sql(sql,cnx)

lista_end=lista[lista.借据.isin(equity_record.order_no)]
lista_end1=lista[~lista.借据.isin(equity_record.order_no)]
lista_end1=lista_end1[lista_end1.钜银核实.isna()]
lista_end1.to_excel(r'D:\Work\out_data\权益\玖富对账\20240204对账差异_20240223.xlsx',index=False)

#%%-----------20231229涛哥要的对账数据

import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


# 权益金额
sql = """ select
	    DISTINCT order_no,
		collection_time as 收款时间,
	    equity_amount as 权益金额,
	    equity_source as 权益来源,
	    amount_type as 收退款类型,
	    refund_amount as 退款金额,
	    refund_time as 退款时间
	from 
        juin_loan_core_prd.equity_record where is_delete=0 """
equity_record=pd.read_sql(sql,cnx)

# 放款金额
sql = """ select  
        a.order_no,
    	a.id as order_id,
    	a.loan_date,
    	a.actual_rate,
        a.loan_amount,
    	a.period,
        case when a.funder_id=2 then '民和'
             when a.funder_id=1 then '星光' end as 放款通道,
        b.status
	from 
		juin_loan_core_prd.order_record as a
    left join 
        juin_loan_core_prd.repayment_plan as b  on a.id=b.order_id
    where a.pay_status=20"""
order_record=pd.read_sql(sql,cnx)

order_record=pd.merge(order_record,equity_record,how='left',on='order_no')
order_record['退款金额'].fillna(0,inplace=True)
order_record['权益收款']=order_record.apply(lambda x: x.权益金额 if x.收退款类型=='C' else 0,axis=1)
order_record['放款月份']=pd.to_datetime(order_record['loan_date']).dt.strftime('%Y%m')

# 在库、应收利息、应收本金,20231229发现repayment_period_id为2089的重复扣款
sql = """ 
select 
      a.order_id,a.period_number,a.period_principal_amount,a.period_interest_amount,a.repayment_date,a.settle_time,
      b.principal_amount,b.interest_amount
from juin_loan_core_prd.repayment_plan_period as a 
left join 
(select  repayment_period_id,sum(principal_amount) as principal_amount ,sum(interest_amount) as interest_amount   from juin_loan_core_prd.repayment_trade_divide_accounts where id!=224 group by 1) as b on a.id=b.repayment_period_id

"""
repayment_plan_period=pd.read_sql(sql,cnx)
repayment_plan_period['账单月份']=pd.to_datetime(repayment_plan_period['repayment_date']).dt.strftime('%Y%m')

order_record=pd.merge(order_record,repayment_plan_period,how='left',on='order_id')

end_order=order_record.groupby(['order_id','放款月份','放款通道','status']).agg({'loan_amount':'mean','principal_amount':'sum','权益收款':'mean','退款金额':'mean'}).reset_index()
end=end_order.groupby(['放款通道','放款月份']).agg({'loan_amount':'sum','order_id':'count','principal_amount':'sum','权益收款':'sum','退款金额':'sum','status': lambda x: (x != 2).sum()}).reset_index()
end['在库本金']=end['loan_amount']-end['principal_amount']
end['loan_amount'] = end['loan_amount'].apply(lambda x: '{:,.2f}'.format(x))
end['在库本金'] = end['在库本金'].apply(lambda x: '{:,.2f}'.format(x))
end['已收权益']=end['权益收款']-end['退款金额']
end.rename(columns={'loan_amount':'放款金额','order_id':'放款笔数','status':'在库笔数',},inplace=True)
end=end[['放款通道','放款月份','放款金额','放款笔数','在库本金','在库笔数','已收权益']]


# 使用 groupby 和 unstack 处理两个 index 和多个 columns
end_month_receiv_principle = order_record.groupby(['放款通道','放款月份', '账单月份']).agg({'period_principal_amount':'sum'}).unstack()
end_month_receiv_interest = order_record.groupby(['放款通道','放款月份', '账单月份']).agg({'period_interest_amount':'sum'}).unstack()
end_month_receip_principle = order_record.groupby(['放款通道','放款月份', '账单月份']).agg({'principal_amount':'sum'}).unstack()
end_month_receip_interest = order_record.groupby(['放款通道','放款月份', '账单月份']).agg({'interest_amount':'sum'}).unstack()

end_month_receiv_principle.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\账单月份应收本金.xlsx')
end_month_receiv_interest.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\账单月份应收利息.xlsx')
end_month_receip_principle.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\账单月份已还本金.xlsx')
end_month_receip_interest.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\账单月份已还利息.xlsx')

end.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\账单月份放款在库.xlsx')



#%%-----------涛哥关注的当天提现是否是之前授信与注册
# 20231229 新增 授信、提现额度分布、转化率
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ select user_id,flow_id,id as order_id,loan_date as 借款放款时间,funder_id,loan_amount from juin_loan_core_prd.order_record where pay_status=20 """
jk=pd.read_sql(sql,cnx)
jk['借款放款时间']=pd.to_datetime(jk['借款放款时间']).dt.date

# a=jk.groupby('借款放款时间').agg({'loan_amount':['sum','count']})

sql = """ select user_id,flow_id,submit_time as 授信申请时间,response_body from juin_loan_core_prd.risk_request_apply  where process_node=0  and flow_id not in ('jiufuDD_48d3a2ad3cfa40068e26','jiufuDD_672c6fd8708943ef81b5','jiufuDD_7639b2dd6f994fffb1fb')  """ 
sx=pd.read_sql(sql,cnx)
sx['授信申请时间']=pd.to_datetime(sx['授信申请时间']).dt.date
sx['授信结果']=sx.apply(lambda x:  json.loads(x.response_body)['授信审批结果'],axis=1)
sx['授信额度']=sx.apply(lambda x:  json.loads(x.response_body)['授信额度'],axis=1)
sx['授信是否通过']=sx.apply(lambda x: 1 if x.授信结果=='通过' else 0,axis=1)
sx=sx[sx.授信是否通过==1]


# 将每个flow_id（一个客户基本只有一个，后面会存在单客户多个flow_id)
sx_jk=pd.merge(sx,jk,how='left',on='flow_id')
sx_group=sx_jk.groupby('flow_id').agg({'loan_amount':'sum','借款放款时间':'max'}).reset_index()#max代表客户是在哪天超额度的
sx_group.rename(columns={"借款放款时间":'最早放款时间'},inplace=True)
sx_group['提现']=sx_group.apply(lambda x: 1 if x.loan_amount>0 else 0,axis=1)
sx=pd.merge(sx,sx_group,how='left',on='flow_id')

bins=[0,10000,20000,30000,40000,50000,60000,70000,80000,90000,100000,500000]
sx['授信额度区间']=pd.cut(sx['授信额度'], bins,duplicates='drop')
sx['提现额度区间']=pd.cut(sx['loan_amount'], bins,duplicates='drop')

a_end=pd.crosstab(index=sx.授信额度区间,
                columns=sx.授信申请时间,
                values=sx.授信是否通过,
                aggfunc=['sum'])
a_end.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\授信额度分布.xlsx')

tx_end=pd.crosstab(index=sx.提现额度区间,
                columns=sx.最早放款时间,
                values=sx.loan_amount,
                aggfunc=['count'])

tx_end.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\提信额度分布.xlsx')

a_sxtx_rate=sx.groupby('授信额度区间').agg({'授信是否通过':'sum','授信额度':'sum','loan_amount':'sum','提现':'sum'}).rename(columns={'授信是否通过':"授信通过客户量",'loan_amount':'放款金额'})
a_sxtx_rate['额度使用率']=a_sxtx_rate['放款金额']/a_sxtx_rate['授信额度']
a_sxtx_rate.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\授信后的提现使用率20231229.xlsx')


js=sx[sx.最早放款时间==datetime.date(2023,12,28)]

a=sx.describe()

a=sx.groupby('授信申请时间').agg({'授信是否通过':'sum','授信额度':'sum','loan_amount':'sum'}).rename(columns={'授信是否通过':"授信通过客户量",'loan_amount':'放款金额'})
a['额度使用率']=a['放款金额']/a['授信额度']
a.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\授信后的提现使用率.xlsx')



sql = """ select id as user_id,create_time as 注册时间 from juin_loan_core_prd.user_info  """#id 无重复
user_info=pd.read_sql(sql,cnx)
user_info['注册时间']=pd.to_datetime(user_info['注册时间']).dt.date

jk=jk[jk.借款放款时间==datetime.date(2023,12,28)]
jk=pd.merge(jk,sx,how='left',on='flow_id')
jk=pd.merge(jk,user_info,how='left',on='user_id')

a=jk.groupby('授信申请时间').size()
a1=jk.groupby('注册时间').size()

#%%-----------给玲智的客户明细+身份证正反面的地址
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ 
select 
a.flow_id,
a.id as order_id,
a.user_id,
a.loan_date,
a.loan_amount,
date(e.repayment_date) as 账单日,
e.period_total_amount as 本期应还金额,
e.period_principal_amount as 本期应还本金,
e.period_interest_amount as 本期应还利息,
e.settle_time as 本期结清日期,
e.period_number,d.brand,
CONVERT ( AES_DECRYPT( from_base64 ( c.phone ), '9342266da419cfb4' ) USING utf8 ) as 商户手机号,
CONVERT ( AES_DECRYPT( from_base64 ( c.name ), '9342266da419cfb4' ) USING utf8 ) as 商户姓名,
CONVERT ( AES_DECRYPT( from_base64 ( c.id_number ), '9342266da419cfb4' ) USING utf8 ) as 商户身份证
from juin_loan_core_prd.order_record as a 
left join juin_loan_core_prd.user_info as c  on a.user_id=c.id
left join juin_loan_core_prd.white_list as d on c.id_number_md5=d.id_number
left join juin_loan_core_prd.repayment_plan_period as e on a.id=e.order_id
where a.pay_status=20   and d.channel_source = 'kyf' and period_number=1 
"""
kyf_detail=pd.read_sql(sql,cnx)



sql = """ select *  from juin_loan_core_prd.third_tencent_request_info  """
third_tencent_request_info=pd.read_sql(sql,cnx)
third_tencent_request_info['url']=third_tencent_request_info.apply(lambda x:json.loads(x.request_param)['idcardStr'],axis=1)

third_tencent_request_info_font=third_tencent_request_info.loc[third_tencent_request_info.api_name=='OCR_FRONT',['url','flow_id']].rename(columns={'url':'url_f'})
third_tencent_request_info_back=third_tencent_request_info.loc[third_tencent_request_info.api_name=='OCR_BACK',['url','flow_id']].rename(columns={'url':'url_b'})

# third_tencent_request_info_back.duplicated(subset='user_id').any()#判断pd是否有重复值

kyf_detail_end=kyf_detail.copy()
kyf_detail_end=pd.merge(kyf_detail_end,third_tencent_request_info_font,how='left',on='flow_id')
kyf_detail_end=pd.merge(kyf_detail_end,third_tencent_request_info_back,how='left',on='flow_id')
kyf_detail_end=kyf_detail_end[kyf_detail_end.]
kyf_detail_end=kyf_detail_end[['user_id','loan_date','loan_amount','账单日','本期应还金额','本期应还本金','本期应还利息','本期结清日期','brand',
                               '商户手机号','商户姓名','商户身份证','url_f','url_b']]
kyf_detail_end.to_excel(r'D:\Work\out_data\分析类\日常\快银付首期应还客户明细20240301.xlsx',index=False)


kyf_detail_end.drop_duplicates(subset=['user_id','order_id','url_f'] ,keep='last' ,inplace=True)


#%%-----------xg对数，给对方异常客户的身份证正反面
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ select user_id,flow_id  from juin_loan_core_prd.risk_credit_result where  remark="身份证头像识别失败！"  """
xg1=pd.read_sql(sql,cnx)


sql = """ select *  from juin_loan_core_prd.xingguang_request_log  where api_type='IN_NET_RESULT'  """
xg_base=pd.read_sql(sql,cnx)
xg_base.drop_duplicates(subset='user_id',keep='last',inplace=True)
xg_base=xg_base[xg_base.user_id.isin(xg1.user_id)]
xg_base['通道拒绝原因']=xg_base.apply(lambda x:json.loads(x.response_body)['data']['remark'],axis=1)
xg_base1=xg_base[['user_id','星光拒绝原因']]


# 临时加表
xg=pd.read_excel(r'D:\Work\out_data\分析类\身份证照片_需求1\to银信金.xlsx')

sql = """ select id as user_id, id_number_md5 from juin_loan_core_prd.user_info    """
user_info=pd.read_sql(sql,cnx)


xg1=pd.merge(xg,user_info,how='left',left_on='MD5',right_on='id_number_md5')



sql = """ select *  from juin_loan_core_prd.third_tencent_request_info   """
third_tencent_request_info=pd.read_sql(sql,cnx)
third_tencent_request_info=third_tencent_request_info[third_tencent_request_info.user_id.isin(xg1.user_id)]
third_tencent_request_info['url']=third_tencent_request_info.apply(lambda x:json.loads(x.request_param)['idcardStr'],axis=1)
third_tencent_request_info_font=third_tencent_request_info.loc[third_tencent_request_info.api_name=='OCR_FRONT',['user_id','url']].rename(columns={'url':'url_f'})
third_tencent_request_info_back=third_tencent_request_info.loc[third_tencent_request_info.api_name=='OCR_BACK',['user_id','url']].rename(columns={'url':'url_b'})

xg1=pd.merge(xg1,xg_base1,how='left',on='user_id')

xg1=pd.merge(xg1,third_tencent_request_info_font,how='left',on='user_id')
xg1=pd.merge(xg1,third_tencent_request_info_back,how='left',on='user_id')
xg1=xg1[xg1.url_f.notna()]
xg1.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\星光拒绝原因客户照片\通道拒绝明细.xlsx',index=False)

# xg1=xg1[xg1.星光拒绝原因.isin(['身份证头像识别失败！','ocr识别失败，请重试 '])]

os.chdir(r"D:\Work\out_data\分析类\身份证照片_需求1")


import requests
def download_file(url, local_filename):
    with requests.get(url, stream=True) as response:
        with open(local_filename, 'wb') as file:
            for chunk in response.iter_content(chunk_size=8192):
                file.write(chunk)
# 正面
for index, row in xg1.iterrows():
    url = row['url_f']
    deal_desc = row['deal_desc']
    local_filename = f'D:\\Work\out_data\\分析类\\身份证照片_需求1\\{deal_desc}_正面_{index}.jpg'
    download_file(url, local_filename)

# 反面
for index, row in xg1.iterrows():
    url = row['url_b']
    flow_id = row['flow_id']
    local_filename = f'D:\\Work\out_data\\分析类\\星光身份证正反面\\{flow_id}_反面_{index}.jpg'
    download_file(url, local_filename)

#%%-----------星光漏斗2.0 含T系列
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 1、risk_request_apply 有user_id flow_id order_id
# 2、三方接口表没有order_id
# 3、三方缓存：30天，根据当前订单(授信or提现)的时间与另一个时间(？)比较
# 4、人脸缓存：180天，也是根据当前订单(提现)的时间与另一个时间(？)比较

# 20231222 注册时间会因为客户再次授信跟着更新，这个bug 技术部会修正
sql = """ select id as user_id,create_time as 注册时间 from juin_loan_core_prd.user_info where  date(create_time) >='2023-11-20' and date(create_time) <='2023-12-10'  """#id 无重复
# sql = """ select id as user_id,create_time as 注册时间 from juin_loan_core_prd.user_info where  date(create_time) >='2023-12-13' and date(create_time) <='2023-12-21'  """#id 无重复
user_info=pd.read_sql(sql,cnx)
user_info['注册时间']=pd.to_datetime(user_info['注册时间']).dt.date
user_info['注册']=1


# 存在user_id重复，基本都是间隔30天以后，符合拒绝锁定期，但不清楚是客户行为还是玖富行为；
sql = """ select user_id,flow_id,submit_time as 授信申请时间,response_body from juin_loan_core_prd.risk_request_apply  where process_node=0   """ 
sx=pd.read_sql(sql,cnx)
sx['发起授信']=1
sx['授信申请时间']=pd.to_datetime(sx['授信申请时间']).dt.date
sx['授信结果']=sx.apply(lambda x:  json.loads(x.response_body)['授信审批结果'],axis=1)
sx['授信是否通过']=sx.apply(lambda x: 1 if x.授信结果=='通过' else 0,axis=1)
sx.drop(columns=['response_body','授信结果'],inplace=True)

# 第一次授信发起时间
f_sx=sx[['user_id', '授信申请时间']].rename(columns={'授信申请时间':'首次授信申请时间'})
f_sx = f_sx.sort_values(by=['user_id', '首次授信申请时间'], ascending=[True, True])
f_sx.drop_duplicates(subset='user_id',keep='first',inplace=True)

# 第一次授信通过时间
tg_sx=sx.loc[sx.授信是否通过==1,['user_id', '授信申请时间']].rename(columns={'授信申请时间':'首次授信通过时间'})
tg_sx = tg_sx.sort_values(by=['user_id', '首次授信通过时间'], ascending=[True, True])
tg_sx.drop_duplicates(subset='user_id',keep='first',inplace=True)

#授信后进入星光入网
sql = """ select user_id,1 as 星光入网拒绝,create_time as 星光入网时间 from juin_loan_core_prd.risk_credit_result  where remark!=''  and credit_result='拒绝' """ 
xgrw=pd.read_sql(sql,cnx)
xgrw['星光入网时间']=pd.to_datetime(xgrw['星光入网时间']).dt.date


# 存在user_id重复，flow_id也重复，符合业务特点，基本是一个user_id对应一个flow_id(授信环节),授信通过之后允许多次提现申请(order_id)
sql = """ select user_id,flow_id,order_id,submit_time as 提现申请时间,response_body from juin_loan_core_prd.risk_request_apply  where process_node=1 """ 
tx=pd.read_sql(sql,cnx)
tx['发起提现']=1
tx['提现申请时间']=pd.to_datetime(tx['提现申请时间']).dt.date
tx['提现结果']=tx.apply(lambda x:  json.loads(x.response_body)['授信审批结果'],axis=1)
tx['提现是否通过']=tx.apply(lambda x: 1 if x.提现结果=='通过' else 0,axis=1)
tx.drop(columns=['response_body','提现结果'],inplace=True)

# 第一次提现发起时间
f_tx=tx[['user_id', '提现申请时间']].rename(columns={'提现申请时间':'首次提现申请时间'})
f_tx = f_tx.sort_values(by=['user_id', '首次提现申请时间'], ascending=[True, True])
f_tx.drop_duplicates(subset='user_id',keep='first',inplace=True)

# 第一次提现通过时间
tg_tx=tx.loc[tx.提现是否通过==1,['user_id', '提现申请时间']].rename(columns={'提现申请时间':'首次提现通过时间'})
tg_tx = tg_tx.sort_values(by=['user_id', '首次提现通过时间'], ascending=[True, True])
tg_tx.drop_duplicates(subset='user_id',keep='first',inplace=True)




# sql = """ select user_id,flow_id,bind_status,create_time as 绑卡申请时间 from juin_loan_core_prd.user_bank_account  where funder_id=1  """ #换绑卡同一个接口，20231222没有重复user_id,flow_id,事后问一下是不是数据库设计导致只存一条
sql = """ select user_id,flow_id,bind_status,create_time as 绑卡申请时间 from juin_loan_core_prd.user_bank_account   """ #换绑卡同一个接口，20231222没有重复user_id,flow_id,事后问一下是不是数据库设计导致只存一条

bk=pd.read_sql(sql,cnx)
bk['绑卡申请时间']=pd.to_datetime(bk['绑卡申请时间']).dt.date
bk['发起绑卡']=1
bk['绑卡是否通过']=bk.apply(lambda x: 1 if x.bind_status=='BIND_SUCCESS' else 0,axis=1)
bk=bk[['user_id','绑卡申请时间','发起绑卡','绑卡是否通过']]


sql = """ select user_id,flow_id,detect_status,create_time as 人脸申请时间 from juin_loan_core_prd.face_result_info    """#人脸是有缓存的，具体天数后面再咨询
rl=pd.read_sql(sql,cnx)
rl['人脸申请时间']=pd.to_datetime(rl['人脸申请时间']).dt.date
rl['发起人脸']=1
rl['人脸是否通过']=rl.apply(lambda x: 1 if x.detect_status=='1' else 0,axis=1)


# 第一次发起人脸时间
f_rl=rl[['user_id', '人脸申请时间']].rename(columns={'人脸申请时间':'首次人脸申请时间'})
f_rl = f_rl.sort_values(by=['user_id', '首次人脸申请时间'], ascending=[True, True])
f_rl.drop_duplicates(subset='user_id',keep='first',inplace=True)

# 第一次人脸通过时间
tg_rl=rl.loc[rl.人脸是否通过==1,['user_id', '人脸申请时间']].rename(columns={'人脸申请时间':'首次人脸通过时间'})
tg_rl = tg_rl.sort_values(by=['user_id', '首次人脸通过时间'], ascending=[True, True])
tg_rl.drop_duplicates(subset='user_id',keep='first',inplace=True)



# sql = """ select user_id,flow_id,id as order_id,pay_status,submit_time as 借款申请时间,loan_date as 借款放款时间,device_type,funder_id from juin_loan_core_prd.order_record  where funder_id=1  """

sql = """ select user_id,flow_id,id as order_id,pay_status,submit_time as 借款申请时间,loan_date as 借款放款时间,device_type,funder_id from juin_loan_core_prd.order_record  """
jk=pd.read_sql(sql,cnx)
jk['借款申请时间']=pd.to_datetime(jk['借款申请时间']).dt.date
jk['借款放款时间']=pd.to_datetime(jk['借款放款时间']).dt.date
jk['发起借款']=1
jk['借款是否通过']=jk.apply(lambda x: 1 if x.pay_status==20 else 0,axis=1)

# 第一次发起借款时间
f_jk=jk[['user_id', '借款申请时间']].rename(columns={'借款申请时间':'首次借款申请时间'})
f_jk = f_jk.sort_values(by=['user_id', '首次借款申请时间'], ascending=[True, True])
f_jk.drop_duplicates(subset='user_id',keep='first',inplace=True)

# 第一次放款通过时间
tg_jk=jk.loc[jk.借款是否通过==1,['user_id', '借款放款时间']].rename(columns={'借款放款时间':'首次借款通过时间'})
tg_jk = tg_jk.sort_values(by=['user_id', '首次借款通过时间'], ascending=[True, True])
tg_jk.drop_duplicates(subset='user_id',keep='first',inplace=True)


# 20231222 T系列版本

user_info=pd.merge(user_info,f_sx,on='user_id',how='left')
user_info=pd.merge(user_info,tg_sx,on='user_id',how='left')

user_info=pd.merge(user_info,xgrw,on='user_id',how='left')

user_info=pd.merge(user_info,f_tx,on='user_id',how='left') 
user_info=pd.merge(user_info,tg_tx,on='user_id',how='left') 

user_info=pd.merge(user_info,bk,on='user_id',how='left') 

user_info=pd.merge(user_info,f_rl,on='user_id',how='left') 
user_info=pd.merge(user_info,tg_rl,on='user_id',how='left') 

user_info=pd.merge(user_info,f_jk,on='user_id',how='left') 
user_info=pd.merge(user_info,tg_jk,on='user_id',how='left') 

time_list=[
'注册时间',
'首次授信申请时间',
'首次授信通过时间',
'星光入网时间',
'首次提现申请时间',
'首次提现通过时间',
'绑卡申请时间',
'首次人脸申请时间',
'首次人脸通过时间',
'首次借款申请时间',
'首次借款通过时间'
    ]

for var in time_list:
    user_info[var]=pd.to_datetime(user_info[var]).dt.date

user_info['发起绑卡'].fillna('0',inplace=True)
user_info['发起绑卡']=user_info['发起绑卡'].astype('int')
user_info['绑卡是否通过'].fillna('0',inplace=True)
user_info['绑卡是否通过']=user_info['绑卡是否通过'].astype('int')
user_info['星光入网通过'] = user_info.apply(lambda x: 1 if x['星光入网拒绝'] != 1 and pd.notna(x['首次授信通过时间']) else 0, axis=1)

wj=pd.read_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\误拒名单20231220.xlsx')
user_info=pd.merge(user_info,wj,how='left',on='user_id')
user_info=user_info[user_info.误拒!=1]


# 20231222版_常规版
end=pd.DataFrame({ '注册':user_info.注册.sum(),
                  '授信申请': user_info.首次授信申请时间.notna().sum(),
                  '授信通过': user_info.首次授信通过时间.notna().sum(),
                  '星光入网通过': user_info.星光入网通过.sum(),
                  '提现申请': user_info.首次提现申请时间.notna().sum(),
                  '提现通过': user_info.首次提现通过时间.notna().sum(),
                  '绑卡申请': user_info.发起绑卡.sum(),
                  '绑卡通过': user_info.绑卡是否通过.sum(),
                  '人脸申请': user_info.首次人脸申请时间.notna().sum(),
                  '人脸通过': user_info.首次人脸通过时间.notna().sum(),
                  '放款成功': user_info.首次借款通过时间.notna().sum()
                  },index=['星光渠道'])
end.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\星光_user_注册漏斗_25_星光入网.xlsx')


# 20231222版_T0版
end_t0 = pd.DataFrame({
    '注册': [user_info['注册'].sum()],
    '授信申请': sum([(user_info['首次授信申请时间'] - user_info['注册时间']).dt.days<=0 & user_info['首次授信申请时间'].notna()]).sum(),
    '授信通过': sum([(user_info['首次授信通过时间'] - user_info['注册时间']).dt.days<=0 & user_info['首次授信通过时间'].notna()]).sum(),
    '提现申请': sum([(user_info['首次提现申请时间'] - user_info['注册时间']).dt.days<=0 & user_info['首次提现申请时间'].notna()]).sum(),
    '提现通过': sum([(user_info['首次提现通过时间'] - user_info['注册时间']).dt.days<=0 & user_info['首次提现通过时间'].notna()]).sum(),
    '绑卡申请': sum([(user_info['绑卡申请时间'] - user_info['注册时间']).dt.days<=0 & user_info['绑卡申请时间'].notna()]).sum(),
    '绑卡通过': sum([(user_info['绑卡申请时间'] - user_info['注册时间']).dt.days<=0 & user_info['绑卡申请时间'].notna() & (user_info['绑卡是否通过']==1) ]).sum(),
    '人脸申请': sum([(user_info['首次人脸申请时间'] - user_info['注册时间']).dt.days<=0 & user_info['首次人脸申请时间'].notna()]).sum(),
    '人脸通过': sum([(user_info['首次人脸通过时间'] - user_info['注册时间']).dt.days<=0 & user_info['首次人脸通过时间'].notna()]).sum(),
    '放款成功': sum([(user_info['首次借款通过时间'] - user_info['注册时间']).dt.days<=0 & user_info['首次借款通过时间'].notna()]).sum()
}, index=['星光渠道'])

end_t0.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\星光_user_注册漏斗_22_t0.xlsx')


# 20231222版_TN版
end_t0 = pd.DataFrame({
    '注册': [user_info['注册'].sum()],
    '授信申请': sum([(user_info['首次授信申请时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次授信申请时间'].notna()]).sum(),
    '授信通过': sum([(user_info['首次授信通过时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次授信通过时间'].notna()]).sum(),
    '提现申请': sum([(user_info['首次提现申请时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次提现申请时间'].notna()]).sum(),
    '提现通过': sum([(user_info['首次提现通过时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次提现通过时间'].notna()]).sum(),
    '绑卡申请': sum([(user_info['绑卡申请时间'] - user_info['注册时间']).dt.days<=3 & user_info['绑卡申请时间'].notna()]).sum(),
    '绑卡通过': sum([(user_info['绑卡申请时间'] - user_info['注册时间']).dt.days<=3 & user_info['绑卡申请时间'].notna() & (user_info['绑卡是否通过']==1) ]).sum(),
    '人脸申请': sum([(user_info['首次人脸申请时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次人脸申请时间'].notna()]).sum(),
    '人脸通过': sum([(user_info['首次人脸通过时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次人脸通过时间'].notna()]).sum(),
    '放款成功': sum([(user_info['首次借款通过时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次借款通过时间'].notna()]).sum()
}, index=['星光渠道'])

end_t0.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\星光_user_注册漏斗_22_t3.xlsx')



# 20231222版 :11.20-12.10
time_list=[
'注册时间',
'首次授信申请时间',
'首次授信通过时间',
'首次提现申请时间',
'首次提现通过时间',
'首次借款申请时间',
'首次借款通过时间'
    ]

for var in time_list:
    user_info[var]=pd.to_datetime(user_info[var]).dt.date
    
# 20231222版_TN版
end_t0 = pd.DataFrame({
    '注册': [user_info['注册'].sum()],
    '授信申请': sum([(user_info['首次授信申请时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次授信申请时间'].notna()]).sum(),
    '授信通过': sum([(user_info['首次授信通过时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次授信通过时间'].notna()]).sum(),
    '提现申请': sum([(user_info['首次提现申请时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次提现申请时间'].notna()]).sum(),
    '提现通过': sum([(user_info['首次提现通过时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次提现通过时间'].notna()]).sum(),
    '放款成功': sum([(user_info['首次借款通过时间'] - user_info['注册时间']).dt.days<=3 & user_info['首次借款通过时间'].notna()]).sum()
}, index=['玖富11201210'])

end_t0.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\玖富user_11201210.xlsx')

# js=user_info.loc[((user_info['首次授信通过时间'] - user_info['注册时间']).dt.days<=3) & (user_info['首次授信通过时间'].notna()) & (user_info['首次借款通过时间'].isna()),['user_id','注册时间','首次借款通过时间']]
js=user_info.loc[((user_info['首次授信通过时间'] - user_info['注册时间']).dt.days<=3) & (user_info['首次授信通过时间'].notna()) ,['user_id','注册时间','首次借款通过时间']]



js=pd.merge(js,jk,how='left',on='user_id')
js.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\js.xlsx')

# 20231221版
user_info=pd.merge(user_info,sx,on='user_id',how='left')#右表按通过降序 user_id 是否通过排序，取第一个user_id
user_info=pd.merge(user_info,u_tx,on=['user_id','flow_id'],how='left')#右表按通过降序 user_id 是否通过排序，取第一个user_id
user_info=pd.merge(user_info,bk,on=['user_id','flow_id'],how='left')#右表按通过降序 user_id 是否通过排序，取第一个user_id
user_info=pd.merge(user_info,rl,on=['user_id','flow_id'],how='left')#右表按通过降序 user_id 是否通过排序，取第一个user_id
user_info=pd.merge(user_info,jk,on=['user_id','flow_id'],how='left')#右表按通过降序 user_id 是否通过排序，取第一个user_id



end=pd.DataFrame({ '注册':user_info.注册.sum(),
                  '授信申请':user_info.发起授信.sum(),
                  '授信通过':user_info.授信是否通过.sum(),
                 '提现申请':user_info.发起提现.sum(),
                  '提现通过':user_info.提现是否通过.sum(),
                  '发起人脸':user_info.发起人脸.sum(),
                  '人脸通过':user_info.人脸是否通过.sum(),
                  '发起绑卡':user_info.发起绑卡.sum(),
                  '绑卡通过':user_info.绑卡是否通过.sum()
                  
                  },index=['星光渠道'])

user_info.device_type.fillna('3',inplace=True)
end1=user_info.groupby(['device_type']).agg({'注册':'sum',
                                           '发起授信':'sum',
                                           '授信是否通过':'sum',
                                           '发起提现':'sum',
                                           '提现是否通过':'sum',
                                           '发起人脸':'sum',
                                           '人脸是否通过':'sum',
                                           '发起绑卡':'sum',
                                           '绑卡是否通过':'sum',
                                           '发起借款':'sum',
                                           '借款是否通过':'sum'
                                           }).rename(columns={'发起授信':'授信申请',
                                                              '授信是否通过':'授信通过',
                                                              '发起提现':'提现申请',
                                                              '提现是否通过':'提现通过',
                                                              '人脸是否通过':'人脸通过',
                                                              '绑卡是否通过':'绑卡通过',
                                                              '借款是否通过':'借款通过'})
end1.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\星光_user_注册漏斗.xlsx')


# 明细
end1_detail=user_info[user_info.发起授信==1]
wj=pd.read_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\误拒名单20231220.xlsx')
end1_detail=pd.merge(end1_detail,wj,how='left',on='user_id')
end1_detail_out=end1_detail[[
    'flow_id',
    '授信申请时间',
    '发起授信',
    '授信是否通过',
    '提现申请时间',
    '发起提现',
    '提现是否通过',
    '绑卡申请时间',
    '发起绑卡',
    '绑卡是否通过',
    '人脸申请时间',
    '发起人脸',
    '人脸是否通过',
    '借款是否通过',
    '误拒'
    ]]

a=end1_detail.columns
end1_detail_out.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\授信客户明细_附误拒标签.xlsx')

# 提现漏斗
tx_lo=pd.merge(u_tx,bk,on=['user_id','flow_id'],how='left')
tx_lo=pd.merge(tx_lo,rl,on=['user_id','flow_id'],how='left')
tx_lo=pd.merge(tx_lo,jk,on=['user_id','flow_id'],how='left')
tx_lo=tx_lo[tx_lo.funder_id==1]




end_tx=tx_lo.groupby('device_type').agg({
                                           '发起提现':'sum',
                                           '提现是否通过':'sum',
                                           '发起人脸':'sum',
                                           '人脸是否通过':'sum',
                                           '发起绑卡':'sum',
                                           '绑卡是否通过':'sum',
                                           '发起借款':'sum',
                                           '借款是否通过':'sum'
                                           }).rename(columns={
                                                              '发起提现':'提现申请',
                                                              '提现是否通过':'提现通过',
                                                              '人脸是否通过':'人脸通过',
                                                              '绑卡是否通过':'绑卡通过',
                                                              '借款是否通过':'借款是过'})

end_tx.to_excel(r'C:\Users\zp457\Desktop\新建文件夹\漏斗\星光_user_提现漏斗.xlsx',index=False)






#%%-----------全量放款客户匹配玖富权益解析（已经复制转移到单独的py文件方便脚本运行)
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
# 生产库
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 指定文件夹路径
folder_path = r'D:\Work\risk_rule\权益\权益'  # 将 'your_folder_path' 替换为你的文件夹路径

# 初始化一个空的DataFrame
end = pd.DataFrame()

# 遍历文件夹中的所有txt文件
for file_index, filename in enumerate(os.listdir(folder_path), 1):
    file_path = os.path.join(folder_path, filename)

    # 打开当前txt文件
    with open(file_path, 'r') as file:
        for line_number, line in enumerate(file, 1):
            # 使用 "|" 分隔符拆分每一行
            fields = line.strip().split('|')

            # 创建一个DataFrame表示当前行
            df = pd.DataFrame([fields], index=[line_number])

            # 将当前行的DataFrame追加到end表
            if line_number==1:
                end=df.copy() 
            else :
                end=pd.concat([end,df],axis=0)
    if file_index==1:
        end_df=end.copy()
    else:
        end_df=pd.concat([end_df,end],axis=0)

col_name=['order_no','权益','玖富计算时间','loan_amount','period','actual_rate','loan_date','玖富id1','玖富id2','玖富id3','玖富提交时间']
end_df.columns=col_name

sql = """ select order_no,id,loan_date as 内部放款时间,actual_rate as 内部放款利率,period as 内部放款期限 from juin_loan_core_prd.order_record    """
order_record=pd.read_sql(sql,cnx)
end_df=pd.merge(order_record,end_df,how='left',on='order_no')

sql = """ select order_id,principal_amount,total_interest_amount from juin_loan_core_prd.repayment_plan    """
repayment_plan_period=pd.read_sql(sql,cnx)
end_df=pd.merge(end_df,repayment_plan_period,how='left',right_on='order_id',left_on='id')
end_df.drop(columns=['id'],inplace=True)
end_df['是否购买权益']=end_df.apply(lambda x: 1 if pd.notna(x.权益)  else 0  ,axis=1)
end_df.loc[end_df['order_no'] == 'JYL1001-20231115-jiufu-fxfGMap', '内部放款利率'] = 36
end_df.to_excel(r'D:\Work\risk_rule\权益\权益总名单.xlsx',index=False)


from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate?charset=utf8mb4'
# db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate?charset=utf8mb4&collation=utf8mb4_general_ci'
engine = create_engine(db_url)
end_df.to_sql(name='equity', con=engine, if_exists='replace', index=False)
#%%-----------pmt

def pmt(r,N,pv,fv,typec):
    pmt=r/((1+r)**N-1)*-(pv*(1+r)**N+fv)
    if typec==1:
        pmt=pmt/(1+r)
    return pmt
ab=pmt(0.02,12,3101400,0,0)

def fv(r,N,c,pv,typec):
    if typec==1:
        c=c*(1+r)
    fv=-(c*((1+r)**N-1)/r+pv*(1+r)**N)

    return fv
a=fv(0.02,12,ab,3101400)

def ipmt(r,per,nper,pv,fv,typec):
    ipmt=fv(r,per-1.pmt(r,nper,pv,fv,typec),pv,typec)*r
    
    if typec==1:
        ipmt=ipmt/(1+r)
    return ipmt

def ppmt(r,per,nper,pv,fv,typec):
    
    return pmt(r,nper,pv,fv,typec) -ipmt(r,per,nper,pv,fv,typec)


#%%-----------flask
from flask import Flask

app = Flask(__name__)

@app.route('/')
def index():
    return {
                                                            "msg": "success",
                                                    "data": "welcome to use flask.",
                                                    "大点声":"我爱你",
                                                    "测试":"啊手动阀打发"
                                        
    }

@app.route('/user/<u_id>')
def user_info(u_id):
    return {
        "msg": "success",
        "data": {
            "id": u_id,
            "username": 'yuz',
            "age": 18
        }
    }

help(Flask(__name__))

app.run()


import flask
help(flask)
#%%-----------基于内存的数据库，比如 SQLite
如果你的场景中没有固定的数据库连接对象（cnx），而是希望使用 Python 的本地对象（比如字典、列表、DataFrame 等）来构建查询，你可以考虑使用基于内存的数据库，比如 SQLite，以及相关的 Python 库。SQLite 允许你在内存中创建和操作数据库，而不需要连接到外部数据库服务器。

以下是一个使用 SQLite 内存数据库的示例，可以在其中动态构建 SQL 查询语句并执行：


import sqlite3
import pandas as pd

# 创建内存数据库连接
conn = sqlite3.connect(':memory:')

# 假设你有一个 Python 字典，其中包含表名和查询条件
query_data = {
    "table_name": "white_list",
    "condition": "age > 30"
}

# 创建表格并插入示例数据
cursor = conn.cursor()
cursor.execute("CREATE TABLE white_list (id INT, name TEXT, age INT)")
cursor.execute("INSERT INTO white_list (id, name, age) VALUES (1, 'John', 25)")
cursor.execute("INSERT INTO white_list (id, name, age) VALUES (2, 'Jane', 35)")
conn.commit()

# 构建 SQL 查询语句
sql = "SELECT * FROM {table_name} WHERE {condition}".format(**query_data)

# 执行 SQL 查询
white_list = pd.read_sql(sql, conn)

# white_list 包含了从数据库中检索到的数据
print(white_list)

# 关闭连接
conn.close()






#%%-----------三方对账
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
# 生产库
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# TD
sql = """ 
select lower(api_name) as api_name,DATE_FORMAT( create_time, '%Y-%m' ) AS month,
DATE_FORMAT( create_time, '%Y-%m-%d' ) AS date,
sum(case when status="0" then 1 else 0 end) as 累计成功调用量,
sum(case when status="1" then 1 else 0 end) as 累计失败调用量,
sum(case when MONTH(create_time) = 11 AND YEAR(create_time) = 2023 and status="0" then 1 else 0 end) as 11月成功调用量,
sum(case when MONTH(create_time) = 11 AND YEAR(create_time) = 2023 and status="1" then 1 else 0 end) as 11月失败调用量
 from juin_loan_core_prd.third_td_request_info where lower(api_name)!='get_token' group by 1,2,3

   """
TD=pd.read_sql(sql,cnx)
TD['来源']='生产'


# 备份库
with open(r"JYLOAN.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ 
 select lower(api_name) as api_name,DATE_FORMAT( create_time, '%Y-%m' ) AS month,
 DATE_FORMAT( create_time, '%Y-%m-%d' ) AS date,
 sum(case when status="0" then 1 else 0 end) as 累计成功调用量,
 sum(case when status="1" then 1 else 0 end) as 累计失败调用量,
 sum(case when MONTH(create_time) = 11 AND YEAR(create_time) = 2023 and status="0" then 1 else 0 end) as 11月成功调用量,
 sum(case when MONTH(create_time) = 11 AND YEAR(create_time) = 2023 and status="1" then 1 else 0 end) as 11月失败调用量
  from juin_cord_prd_backup_1113.third_td_request_info where lower(api_name)!='get_token' group by 1,2,3

    """
TD_bf=pd.read_sql(sql,cnx)
TD_bf['来源']='备份'

# 备份库

sql = """ 
 select lower(api_name) as api_name,DATE_FORMAT( create_time, '%Y-%m' ) AS month,
 DATE_FORMAT( create_time, '%Y-%m-%d' ) AS date,
 sum(case when status="0" then 1 else 0 end) as 累计成功调用量,
 sum(case when status="1" then 1 else 0 end) as 累计失败调用量,
 sum(case when MONTH(create_time) = 11 AND YEAR(create_time) = 2023 and status="0" then 1 else 0 end) as 11月成功调用量,
 sum(case when MONTH(create_time) = 11 AND YEAR(create_time) = 2023 and status="1" then 1 else 0 end) as 11月失败调用量
  from ky_loan.third_td_request_info where lower(api_name)!='get_token' group by 1,2,3

    """

TD_cs=pd.read_sql(sql,cnx)
TD_cs['来源']='测试'  
    
    


td_11=pd.concat([TD,TD_bf],axis=0)
td_11.to_excel(r'D:\Work\Information\三方对接\talkingdata\月度对账单\TD11.xlsx',index=False)
TD_cs.to_excel(r'D:\Work\Information\三方对接\talkingdata\月度对账单\TD11_cs.xlsx',index=False)

#%%-----------线下统计各三方得调用数据 , 将三方计费标准配置表导入到生产库-风控库
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
# 生产库
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text
table_pz = pd.read_excel(r"D:\Work\Information\三方对接\三方梳理及报表\配置表.xlsx",sheet_name=r'Sheet1')
# 生产库
db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@fk-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate?charset=utf8mb4'
engine = create_engine(db_url)
table_pz.to_sql(name='supplier_config_new', con=engine, if_exists='replace', index=False)




from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text
table_pz = pd.read_excel(r"D:\Work\Information\三方对接\三方梳理及报表\配置表.xlsx",sheet_name=r'Sheet1')
# 生产库
db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate?charset=utf8mb4'
engine = create_engine(db_url)
table_pz.to_sql(name='supplier_config', con=engine, if_exists='replace', index=False)

# 备份库
db_url = 'mysql+pymysql://test:kytest@192.168.20.213:3306/juin_risk_operate?charset=utf8mb4'
engine = create_engine(db_url)
table_pz.to_sql(name='supplier_config', con=engine, if_exists='replace', index=False)



# 1、融360

sql = """ select * from juin_loan_core_prd.third_rong360_request_info    """
var_df=pd.read_sql(sql,cnx)
print(" userid数量：",var_df['user_id'].nunique()) 
columns_m=['api_name','status','response_code','process_node']
result360=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    temp = var_df[col].value_counts().to_frame('计数').reset_index().rename(columns={'index':'取值'})
    temp.sort_values(by='取值',inplace=True)
    temp['ColName']=col
    temp['计数占比']=temp['计数']/var_df.shape[0]
    if i==0:
        result360=temp
    else:
        result360=pd.concat([result360,temp])
        
def jj_360(x):
    if x.api_name=='ANTIFRAUDFORRISKREPORTV1':
        return '查中'
    elif x.api_name=='ANTIFRAUDFORRISKV11':
        return '查中'
    elif x.api_name=='AVGCALLTIME':
        return '查中'
    elif x.api_name=='AVGFLOW':
        return '查中'
    elif x.api_name=='AVGPHONEBILL':
        return '查中'
    elif x.api_name=='BASICRISKFACTORV2':
        return '查中'
    elif x.api_name=='DETAILBLACKLIST':
        return '查中'
    elif x.api_name=='LAWGREYLIST':
        return '查询'
    elif x.api_name=='MULTILOANULTIMATEV2REPORT':
        return '查询'
    elif x.api_name=='NATURALPERSONSPCLLIST':
        return '查询'
    elif x.api_name=='NEWBLACKLISTDETAIL':
        return '查询'
    elif x.api_name=='QIANSAIV1':
        return '查中'
    elif x.api_name=='STANDARDV51':
        return '查中'
    elif x.api_name=='ZHANXINXWV1':
        return '查中'
out_rong360=pd.crosstab(var_df['api_name'],var_df['status']).reset_index().rename(columns={'0':'查询成功','1':'查询失败'})        
out_rong360['收费方式']=out_rong360.apply(jj_360,axis=1)
out_rong360['三方']='融360'
        
# 查看法律诉讼详情
# ssxq=var_df.loc[var_df.api_name=='NEWBLACKLISTDETAIL',['user_id']] 
#    ssxq=var_df[var_df.api_name=='NEWBLACKLISTDETAIL']   
# ssxq.to_excel(r'D:\Work\Information\三方对接\三方梳理及报表\ssxq.xlsx') 

# a=list(dysb.user_id)
# dysb=var_df[var_df.status=='1']
# dycg1=var_df[var_df.user_id.isin(a)]
# dycg2=var_df[var_df.user_id==26908]

# 2、td
sql = """ select * from juin_loan_core_prd.third_td_request_info    """
var_df=pd.read_sql(sql,cnx)
print(" userid数量：",var_df['user_id'].nunique()) 
columns_m=['api_name','status','response_code','process_node']
result_td=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    temp = var_df[col].value_counts().to_frame('计数').reset_index().rename(columns={'index':'取值'})
    temp.sort_values(by='取值',inplace=True)
    temp['ColName']=col
    temp['计数占比']=temp['计数']/var_df.shape[0]
    if i==0:
        result_td=temp
    else:
        result_td=pd.concat([result_td,temp])
    
out_td=pd.crosstab(var_df['api_name'],var_df['status']).reset_index().rename(columns={'0':'查询成功','1':'查询失败'})        
out_td['收费方式']='查中'
out_td['三方']='TD'

# 3、微言
sql = """ select * from juin_loan_core_prd.third_weiyan_request_info    """
var_df=pd.read_sql(sql,cnx)
print(" userid数量：",var_df['user_id'].nunique()) 
columns_m=['api_name','status','response_code','process_node']
result_wy=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    temp = var_df[col].value_counts().to_frame('计数').reset_index().rename(columns={'index':'取值'})
    temp.sort_values(by='取值',inplace=True)
    temp['ColName']=col
    temp['计数占比']=temp['计数']/var_df.shape[0]
    if i==0:
        result_wy=temp
    else:
        result_wy=pd.concat([result_wy,temp])


def jj_wy(x):
    if x.api_name=='MULTIPOINT_LOAN':
        return '查询'
    else:
        return '查中'
    
    
out_wy=pd.crosstab(var_df['api_name'],var_df['status']).reset_index().rename(columns={'0':'查询成功','1':'查询失败'})        
out_wy['收费方式']=out_wy.apply(jj_wy,axis=1)
out_wy['三方']='微言'

# 4、冰鉴
sql = """ select * from juin_loan_core_prd.third_icekredit_request_info    """
var_df=pd.read_sql(sql,cnx)
print(" userid数量：",var_df['user_id'].nunique()) 
columns_m=['api_name','status','response_code','process_node']
result_bj=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    temp = var_df[col].value_counts().to_frame('计数').reset_index().rename(columns={'index':'取值'})
    temp.sort_values(by='取值',inplace=True)
    temp['ColName']=col
    temp['计数占比']=temp['计数']/var_df.shape[0]
    if i==0:
        result_bj=temp
    else:
        result_bj=pd.concat([result_bj,temp])
# out_bj=pd.crosstab(var_df['api_name'],var_df['status']).reset_index().rename(columns={'0':'查询成功','1':'查询失败'})   

def jj_bj(x):
    if x.api_name=='huomou22':
        return '查中'
    else:
        return '查询'

        
out_bj=pd.crosstab(var_df['api_name'],var_df['response_code']).reset_index().rename(columns={'00':'查询成功'})        
out_bj['收费方式']=out_bj.apply(jj_bj,axis=1)
out_bj['三方']='冰鉴'


# 5、百融
sql = """ select * from juin_loan_core_prd.third_bairong_request_info    """
var_df=pd.read_sql(sql,cnx)
print(" userid数量：",var_df['user_id'].nunique()) 
columns_m=['api_name','status','response_code','process_node']
result_br=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    temp = var_df[col].value_counts().to_frame('计数').reset_index().rename(columns={'index':'取值'})
    temp.sort_values(by='取值',inplace=True)
    temp['ColName']=col
    temp['计数占比']=temp['计数']/var_df.shape[0]
    if i==0:
        result_br=temp
    else:
        result_br=pd.concat([result_br,temp])
        

out_br=pd.crosstab(var_df['api_name'],var_df['status']).reset_index().rename(columns={'0':'查询成功','1':'查询失败'})              
out_br['收费方式']='查询'
out_br['三方']='百融'


# 周舟想评估的每天调用量,近期每天的放款任务都是15W
var_df['create_date']=pd.to_datetime(var_df['create_time']).dt.date
current_date = datetime.now().date()

for i in range(7):
    day = current_date - timedelta(days=i)
    temp_df=var_df[var_df.create_date==day]
    out_br_pd=pd.crosstab(temp_df['api_name'],temp_df['status']).reset_index().rename(columns={'0':'查询成功','1':'查询失败'}) 
    out_br_pd['date']=day
    if i==0:
        end=out_br_pd.copy()
    else :
        end=pd.concat([end,out_br_pd],axis=0)

end['api_name']=end['api_name'].apply(lambda x:x.lower())
end=pd.merge(end,table_pz,how='left',left_on='api_name',right_on='接口名称')
end=end[~end.api_name.isin(['login','get_token'])]
end['收费方式']='查询'
end['费用']=end.apply(lambda x:  x.价格*x.查询成功 if x.收费方式=='查中' else x.价格*(x.查询成功+x.查询失败),axis=1)
a=end.groupby('date').费用.agg(['sum'])

# 6、电话邦,同盾,银联智策,共用朴道接口
sql = """ select * from juin_loan_core_prd.third_pudao_request_info    """
var_df=pd.read_sql(sql,cnx)
print(" userid数量：",var_df['user_id'].nunique()) 
columns_m=['api_name','status','response_code','process_node']
result_pd=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    temp = var_df[col].value_counts().to_frame('计数').reset_index().rename(columns={'index':'取值'})
    temp.sort_values(by='取值',inplace=True)
    temp['ColName']=col
    temp['计数占比']=temp['计数']/var_df.shape[0]
    if i==0:
        result_pd=temp
    else:
        result_pd=pd.concat([result_pd,temp])

def jj_pd(x):
    if x.api_name=='NUMBER_LABEL':
        return '查询'
    else:
        return '查中'
def jj_channel(x):
    if x.api_name=='NUMBER_LABEL':
        return '朴道-电话邦'
    elif x.api_name=='TONG_DUN':
        return '朴道-同盾'
    elif x.api_name=='UNIONPAY':
        return '朴道-银联智策'
    
out_pd=pd.crosstab(var_df['api_name'],var_df['status']).reset_index().rename(columns={'0':'查询成功','1':'查询失败'})              
out_pd['收费方式']=out_pd.apply(jj_pd,axis=1)
out_pd['三方']=out_pd.apply(jj_channel,axis=1)

list_channel=['out_pd','out_br','out_rong360','out_bj','out_td','out_wy']
for table_name in list_channel:
    # printP(table_name)
    if table_name=='out_pd':
        exec(f"channel_cost = {table_name}")
    else:
        exec(f"channel_cost = pd.concat([channel_cost, {table_name}],axis=0)")
        # channel_cost=pd.concat([channel_cost,i],axis=0)
channel_cost['查询成功'].fillna(0,inplace=True)
channel_cost['查询失败'].fillna(0,inplace=True)
channel_cost['api_name']=channel_cost['api_name'].apply(lambda x:x.lower())
table_pz = pd.read_excel(r"D:\Work\Information\三方对接\三方梳理及报表\配置表.xlsx",sheet_name=r'Sheet1')
channel_cost=pd.merge(channel_cost,table_pz,how='left',left_on='api_name',right_on='接口名称')
channel_cost=channel_cost[~channel_cost.api_name.isin(['login','get_token'])]
channel_cost=channel_cost[['三方','接口名称中文名','收费方式','查询成功','查询失败','价格']]
channel_cost['费用']=channel_cost.apply(lambda x:  x.价格*x.查询成功 if x.收费方式=='查中' else x.价格*(x.查询成功+x.查询失败),axis=1)

sql = """ 	select  user_id  from juin_loan_core_prd.risk_request_apply where process_node=0   """
risk_request_apply=pd.read_sql(sql,cnx)

total_count=pd.DataFrame({'三方':'总计','接口名称中文名':'','收费方式':'','查询成功':risk_request_apply.shape[0],'查询失败':0,'价格':channel_cost['价格'].mean(),'费用':channel_cost['费用'].sum()},index=[100])
channel_cost = pd.concat([channel_cost, total_count],axis=0)

channel_cost.to_excel(r'D:\Work\Information\三方对接\三方梳理及报表\累计三方调用量.xlsx',index=False)
#%%-----------钜银客户之多方评分的相关性，原逻辑在（火山、腾讯、蓝象产品测试分析）
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
import numpy as np
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)



jf_all_score=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试结果\内部调用三方评分.xlsx')
col=jf_all_score.var_urule
jf_inside=big_table[col]#去（生产）钜银贷urule入参，出参解析（决策模拟、验证、优化）

# 查看相关性
for var in col:
    jf_inside[var]=jf_inside[var].fillna('-1').astype(float)

corr_data=jf_inside.copy()
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)

corr_data=jf_inside.copy()
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)
rc = {'font.sans-serif': 'SimHei',
      'axes.unicode_minus': False}
sns.set(context='notebook', style='ticks', rc=rc)
# 创建相关系数矩阵的热图，并设置轴标签
plt.figure(figsize=(13, 13))
# sns.set(font_scale=1)
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", vmin=-1, vmax=1,square=True,
            xticklabels=col,
            yticklabels=col,
            cbar=False)
plt.title("相关系数矩阵热图")
plt.show()

# 分箱然后看分布

list850=[
'huiyan23Bj', 'huomou22Bj', 'qingyun22Bj', 'xingyu22Bj',
'zxjcR360', 'zxqsScoreR360', 'zxsjwlScoreR360',
'zxxdfxyzR360', 'zxxwfxR360',  'tengyun108Td',
'tdDeScore', 'tdDtScore', 'tdGlScore',
'tdXwScore', 'fxpCjHjfv4Wy', 'fxptyWy',   'zzpfLspgV2Wy'
    ]
list100=[
'wddtScoreR360','zyScoreR360','tengyun239Td',
'czqdtgzScoreWy', 'dtsqScoreWy', 'dzqdtgzScoreWy','fyhdtgzScoreWy','yhdtgzScoreWy', 'qt3NdfV1Wy', 'qtNdfV1Wy'
   ]

jf_inside['id']=range(1, len(jf_inside) + 1)

# 850系列
jf_inside850=jf_inside[list850+['id']]

for var in [col for col in jf_inside850.columns if col!='id' ]:
    bins =[float('-inf'),350,450,550,650,750,850,float('inf')]
    temp_intervals = pd.cut(jf_inside850[var], bins=bins, ordered=True)
    jf_inside850[var+'_intervals_pf'] = temp_intervals
    # 添加 'QS' 到类别中
    jf_inside850[var+'_intervals_pf'] = jf_inside850[var+'_intervals_pf'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    jf_inside850[var+'_intervals_pf'].fillna('DEFAULT', inplace=True)

for index,var in enumerate([col for col in jf_inside850.columns if '_intervals_pf' in col]):
    temp = jf_inside850.groupby(var)[var].count().rename_axis("bins").reset_index()
    temp['百分比'+var] = temp[var] / temp[var].sum() 
    if index==0:
        temp_df=temp
    else:
        temp_df=pd.merge(temp_df,temp,how='outer',on='bins')
temp_df.to_excel(r'D:\Work\out_data\分析类\juin客户分布\已接评分850_20240104.xlsx')      
        
# 100系列
jf_inside100=jf_inside[list100+['id']]

for var in [col for col in jf_inside100.columns if col!='id' ]:
    bins =[float('-inf'),20,30,40,50,60,70,80,float('inf')]
    temp_intervals = pd.cut(jf_inside100[var], bins=bins, ordered=True)
    jf_inside100[var+'_intervals_pf'] = temp_intervals
    # 添加 'QS' 到类别中
    jf_inside100[var+'_intervals_pf'] = jf_inside100[var+'_intervals_pf'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    jf_inside100[var+'_intervals_pf'].fillna('DEFAULT', inplace=True)
    
for index,var in enumerate([col for col in jf_inside100.columns if '_intervals_pf' in col]):
    temp = jf_inside100.groupby(var)[var].count().rename_axis("bins").reset_index()
    temp['百分比'+var] = temp[var] / temp[var].sum() 
    if index==0:
        temp_df=temp
    else:
        temp_df=pd.merge(temp_df,temp,how='outer',on='bins')
temp_df.to_excel(r'D:\Work\out_data\分析类\juin客户分布\已接评分100_20240104.xlsx')      


 
#%%-----------火山、腾讯、蓝象产品测试分析
import seaborn as sns
import matplotlib.pyplot as plt

def ks_calc_cross_NB(data,score_col,class_col):

    temp_nodefault=data[data[score_col]!='DEFAULT']
    temp_default=data[data[score_col]=='DEFAULT']
    
    crossfreq=pd.crosstab(temp_nodefault[score_col],temp_nodefault[class_col])
    crossfreq.rename(columns={0:'good',1:'bad'},inplace=True)
    
    crossdens=crossfreq.cumsum(axis=0)/crossfreq.sum()
    crossdens.rename(columns={'good':'good_cum_rate','bad':'bad_cum_rate'},inplace=True)
    crossdens['diff_rate']=abs(crossdens['good_cum_rate']-crossdens['bad_cum_rate'])
    max_value = crossdens['diff_rate'].max()
    crossdens['KS'] = [1 if x == max_value else 0 for x in crossdens['diff_rate']]

    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    crossfreq.reset_index(inplace=True)
    crossfreq.rename(columns={score_col:'Bins'},inplace=True)
    
    crossfreq_default=pd.crosstab(temp_default[score_col],temp_default[class_col])
    crossfreq_default.rename(columns={0:'good',1:'bad'},inplace=True)
    crossfreq_default.reset_index(inplace=True)
    crossfreq_default.rename(columns={score_col:'Bins'},inplace=True)
    
    crossfreq=pd.concat([crossfreq,crossfreq_default],axis=0)
    
    crossfreq['var_name']=score_col
    crossfreq['bad_rate_bins']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    
    return crossfreq

# -------------------------------------------------------------------------------------------------------------聚合吧样本1122
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text
import seaborn as sns
import matplotlib.pyplot as plt


os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ 	select  user_id2 as user_id,loan_status
from juin_risk_operate.yx_fk_test1113_v2 where type=1  """
list_sr=pd.read_sql(sql,cnx)

list_sr.user_id=list_sr.user_id.astype('int64')

list_sr['source']='jhb'

# -------------------------------------------------------------------------------------------------------------银信金全体放款样本5666
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
import warnings
import numpy as np


# 设置警告过滤器，忽略特定类型的警告
warnings.filterwarnings("ignore", category=FutureWarning)

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql="""
select
		a.user_id,
		a.order_no,
	
		case when max_od_day>=15  then 1 else 0 end as dpd15_fz,
		case when max_od_day>=30  then 1 else 0 end as dpd30_fz,
		
		case when f_od>=1  then 1 else 0 end as fpd1_fz,
		case when f_od>=7  then 1 else 0 end as fpd7_fz,
		case when f_od>=15  then 1 else 0 end as fpd15_fz,
		case when f_od>=30  then 1 else 0 end as fpd30_fz,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=1 then 1 else 0 end fpd1_fm,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=7 then 1 else 0 end fpd7_fm,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=15 then 1 else 0 end fpd15_fm,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=30 then 1 else 0 end fpd30_fm

		from loan_core.core_loan_order as a
		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as f_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=1 ) as fpd on a.order_no=fpd.order_no
										
		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as s_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=2 ) as spd on a.order_no=spd.order_no

		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as t_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=3 ) as tpd on a.order_no=tpd.order_no

		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as q_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=4 ) as qpd on a.order_no=qpd.order_no
							
		left join (SELECT order_no,
							 max( DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time) )  as max_od_day           
							from loan_core.core_repayment_plan
							where is_del=0 and repayment_status in (3,2) group by 1 order by 1) as od on a.order_no=od.order_no

		where loan_status=2 and a.is_del=0 and a.loan_channel_code not in ('ceshi') and product_name not in ('灰度测试')  
"""
user_max_od=pd.read_sql(sql,cnx)




# -------------------------------------------------------------------------------------------------------------分析模板
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text
import seaborn as sns
import matplotlib.pyplot as plt

# 样本标签
base15=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试样本\测试样本_raw.xlsx',sheet_name='Sheet1')
lx_result=pd.read_excel(r'D:\Work\Information\三方对接\蓝象\蓝象结果.xlsx')#蓝象


# 先在这里看一下蓝象的经济于融360的话费
sql = """ select user_id,pjhf_j3m_r360  from juin_loan_core_prd.rong360_result_info    """
rong360_result_info=pd.read_sql(sql,cnx)


lx_fx=base15[base15.source=='jf']
lx_fx=pd.merge(lx_fx,lx_result,how='left',on='new_column')
lx_fx=pd.merge(lx_fx,rong360_result_info,how='left',on='user_id')
lx_fx['手机话费']=lx_fx.apply(lambda x: "NA" if x.pjhf_j3m_r360=='NA' else x['pjhf_j3m_r360'].split('-')[-1],axis=1)


freq_df = pd.crosstab(index=lx_fx['手机话费'], columns=lx_fx['lx_level'])
freq_df.to_excel(r'D:\Work\Information\三方对接\蓝象\经济能力vs手机话费.xlsx')
a=lx_fx.groupby(['lx_level','手机话费']).size()

python 怎么实现sas里的 proc freq data=a;

lx_fx.pjhf_j3m_r360.value_counts()

# base15.info()
base15=pd.merge(base15,user_max_od,how='left',on=['user_id','order_no'])
base15=pd.merge(base15,list_sr,how='left',on=['user_id','source'])

# 样本返回结果
base15_huoshan=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试结果\测试样本15000_huoshan_raw.xlsx',sheet_name='Sheet1')
base15_huoshan.drop(columns=['回溯日期','identity_no_md5','phone_md5'],inplace=True)
base15_tengxun_pf=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试结果\测试样本15000_tengxun_raw.xlsx',sheet_name='反欺诈')
base15_tengxun_pf.drop(columns=['回溯日期','identity_no_md5','phone_md5'],inplace=True)
base15_result_pf=pd.merge(base15_huoshan,base15_tengxun_pf,how='left',on='new_column')

base15=pd.merge(base15,base15_result_pf,how='left',on='new_column')
base15=pd.merge(base15,lx_result,how='left',on='new_column')

score_var=[
'Z1221',
'Z1211',
'Z17138',
'Z17140',
'Z17141',
'Z17146',
'Z17176',
'Z17178',
'Z17181',
'Z1112',
'Z1113',
'Z17214',
'Z17254',
'Z17256',
'Z17286',
'Z17289',
'Z17290',
'Z17291',
'Z17212',
'Z17213',
'Z1231',
'Z1241',
'Z1115',
'Z1421'
]

score_var_tx=[
    'fx_score_v03',
    'model_risk_v7_shighirr_score',
    'model_risk_v7_slowirr_score',
    'model_risk_v7_stongyong_score'
    ]
score_var_all=score_var+score_var_tx
# 查看银信金放款客户的标签分布
# dpd15_fz
# dpd30_fz
# fpd15_fz
# fpd30_fz
yxj=base15[base15.source=='yxjfk']
# yxj.fpd30_fz.value_counts()

# 查看聚合吧客户的标签分布
jhb=base15[base15.source=='jhb']
# jhb.loan_status.value_counts()

#------------------------------------------------------------------------------ 1.0、查看玖富客群所有三方评分相关系数
jf_all_score=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试结果\内部调用三方评分.xlsx')
col=['user_id']+list(jf_all_score.var_urule)
jf_inside=big_table[col]#内部解析字段


lx_result_pf=lx_result[['new_column','lx_score']]


base_jf=base15.copy()
base_jf=base_jf[base_jf.source=='jf']
base_jf=pd.merge(base_jf,jf_inside,how='left',on='user_id')
base_jf=pd.merge(base_jf,lx_result_pf,how='left',on='new_column')

col_var=score_var_all+list(jf_all_score.var_urule)+['lx_score']

for var in col_var:
    base_jf[var]=base_jf[var].astype(float)

summary_jf = pd.DataFrame(columns=['Column Name', 'Missing Values', 'Min', 'Mean', 'Median', 'Max'])
for column in [col for col in col_var]:
     missing_values = base_jf[column].isnull().sum()
     min_value = base_jf[column].min()
     mean_value = base_jf[column].mean()
     median_value = base_jf[column].median()
     max_value = base_jf[column].max()

     summary_jf = summary_jf.append({'Column Name': column, 
                                 'Missing Values': missing_values, 
                                 'Min': min_value, 
                                 'Mean': mean_value, 
                                 'Median': median_value, 
                                 'Max': max_value}, ignore_index=True)
     
# 看一下相关系数
corr_data=base_jf[col_var]

for col in col_var:
    corr_data[col].fillna(0.001,inplace=True)
    
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)

correlation_df = pd.DataFrame(correlation_matrix, columns=corr_data.columns, index=corr_data.columns)

correlation_df.to_excel(r'D:\Work\Information\三方对接\银豆\相关系数矩阵1.xlsx')



# 微言pk腾讯
base15_tengxun_dt=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试结果\测试样本15000_tengxun_raw.xlsx',sheet_name='行业风险-多头申请v3')

base15_tengxun_dt1=base15_tengxun_dt[['new_column','多头申请通用分(41001)','长周期多头共债子分(41003)',
                                      '短周期多头共债子分(41002)','非银行多头共债子分(41004)','银行多头共债子分(41005)']]

base15_tengxun_dt=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试结果\测试样本15000_tengxun_raw.xlsx',sheet_name='行业风险-圈团&可疑&电诈')

base15_tengxun_dt2=base15_tengxun_dt[['new_column','圈团1迭代浓度分V2(21007)','圈团3浓度分V1（23006）']]
			
base15_tengxun_dt1=pd.merge(base15_tengxun_dt1,base15_tengxun_dt2,how='left',on='new_column')

base_jf=pd.merge(base_jf,base15_tengxun_dt1,how='left',on='new_column')

col_tx=score_var_tx+[col for col in list(base15_tengxun_dt1.columns) if col !='new_column']
col_wy=list(jf_all_score.loc[jf_all_score.供应商=='微言','var_urule'])

col_var1=col_tx+col_wy
corr_data=base_jf[col_var1]

for col in col_var1:
    corr_data[col].fillna(0.001,inplace=True)
    
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)

correlation_df = pd.DataFrame(correlation_matrix, columns=corr_data.columns, index=corr_data.columns)

correlation_df.to_excel(r'D:\Work\Information\三方对接\银豆\相关系数矩阵_wytx.xlsx')


# 看火山分

for var in [col for col in base_jf.columns if col  in col_var ]:
    num_bins = 10
    bin_edges = np.linspace(base_jf[var].min(), base_jf[var].max(), num_bins + 1)
    bin_edges_series = pd.Series(bin_edges)
    bins = list(pd.concat([pd.Series([float('-inf')]), bin_edges_series, pd.Series([float('inf')])]))
    temp_intervals = pd.cut(base_jf[var], bins=bins, ordered=True)
    base_jf[var+'_intervals_pf'] = temp_intervals
    # 添加 'QS' 到类别中
    base_jf[var+'_intervals_pf'] = base_jf[var+'_intervals_pf'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    base_jf[var+'_intervals_pf'].fillna('DEFAULT', inplace=True)
    
    

for index,var in enumerate([col for col in base_jf.columns if '_intervals_pf' in col]):
    temp = base_jf.groupby(var)['user_id'].count().reset_index().rename(columns={"user_id":'计数',var:'Bins'})
    temp['百分比'+var] = temp['计数'] / temp['计数'].sum() 
    temp=temp[['百分比'+var]]
    if index==0:
        temp_df=temp
    else:
        temp_df=pd.concat([temp_df,temp],axis=1)
temp_df.to_excel(r'D:\Work\Information\三方对接\银豆\玖富所有评分分布.xlsx')

a=base_jf.columns

var='lx_score_intervals_pf'
temp = base_jf.groupby(var)['user_id'].count().reset_index().rename(columns={"user_id":'计数',var:'Bins'})
temp['百分比'+var] = temp['计数'] / temp['计数'].sum() 


#------------------------------------------------------------------------------ 1.1、查看各个分数的相关系数(火山系)
corr_data=base15[score_var]

for col in score_var:
    corr_data[col].fillna(0.001,inplace=True)
    
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)
rc = {'font.sans-serif': 'SimHei',
      'axes.unicode_minus': False}
sns.set(context='notebook', style='ticks', rc=rc)
# 创建相关系数矩阵的热图，并设置轴标签
plt.figure(figsize=(13, 13))
# sns.set(font_scale=1)
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", vmin=-1, vmax=1,square=True,
            xticklabels=score_var,
            yticklabels=score_var,
            cbar=False)
plt.title("相关系数矩阵热图")
plt.show()

plt.figure(figsize=(8, 6))
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", square=True, vmin=-1, vmax=1)
plt.title("Correlation Matrix Heatmap")
plt.show()

#------------------------------------------------------------------------------ 1.2、查看各个分数的相关系数(腾讯系)
corr_data=base15[score_var_tx]

for col in score_var_tx:
    corr_data[col].fillna(1,inplace=True)
    
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)
rc = {'font.sans-serif': 'SimHei',
      'axes.unicode_minus': False}
sns.set(context='notebook', style='ticks', rc=rc)
# 创建相关系数矩阵的热图，并设置轴标签
plt.figure(figsize=(13, 13))
# sns.set(font_scale=1)
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", vmin=-1, vmax=1,square=True,
            xticklabels=score_var_tx,
            yticklabels=score_var_tx,
            cbar=False)
plt.title("相关系数矩阵热图")
plt.show()

plt.figure(figsize=(8, 6))
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", square=True, vmin=-1, vmax=1)
plt.title("Correlation Matrix Heatmap")
plt.show()

#------------------------------------------------------------------------------ 1.3、查看各个分数的ks（yxj总体）

# 银信金放款客户
summary_yxj = pd.DataFrame(columns=['Column Name', 'Missing Values', 'Min', 'Mean', 'Median', 'Max'])
for column in [col for col in yxj.columns if col  in score_var_all+['lx_score'] ]:
     missing_values = yxj[column].isnull().sum()
     min_value = yxj[column].min()
     mean_value = yxj[column].mean()
     median_value = yxj[column].median()
     max_value = yxj[column].max()

     summary_yxj = summary_yxj.append({'Column Name': column, 
                                 'Missing Values': missing_values, 
                                 'Min': min_value, 
                                 'Mean': mean_value, 
                                 'Median': median_value, 
                                 'Max': max_value}, ignore_index=True)
     
summary_yxj.to_excel(r'D:\Work\Information\三方对接\银豆\yxj_summary.xlsx')

for var in [col for col in yxj.columns if col  in score_var_all+['lx_score']  ]:
    num_bins = 10
    bin_edges = np.linspace(yxj[var].min(), yxj[var].max(), num_bins + 1)
    bin_edges_series = pd.Series(bin_edges)
    bins = list(pd.concat([pd.Series([float('-inf')]), bin_edges_series, pd.Series([float('inf')])]))
    temp_intervals = pd.cut(yxj[var], bins=bins, ordered=True)
    yxj[var+'_intervals'] = temp_intervals
    # 添加 'QS' 到类别中
    yxj[var+'_intervals'] = yxj[var+'_intervals'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    yxj[var+'_intervals'].fillna('DEFAULT', inplace=True)

hs_ks=pd.DataFrame()
for class_col in ['dpd15_fz','dpd30_fz','fpd15_fz','fpd30_fz']:
    for index,var in enumerate([col for col in yxj.columns if '_intervals' in col]):
        temp=ks_calc_cross_NB(yxj,var,class_col)
        if index==0:
            hs_ks=temp
        else:
            hs_ks=pd.concat([hs_ks,temp])
        hs_ks.to_excel(r'D:\Work\Information\三方对接\银豆\yxj_'+class_col+'.xlsx')
        
#------------------------------------------------------------------------------ 1.4、查看各个分数的ks（jhb总体）
# 聚合吧客户 
summary_jhb = pd.DataFrame(columns=['Column Name', 'Missing Values', 'Min', 'Mean', 'Median', 'Max'])
for column in [col for col in jhb.columns if col  in score_var_all ]:
     missing_values = jhb[column].isnull().sum()
     min_value = jhb[column].min()
     mean_value = jhb[column].mean()
     median_value = jhb[column].median()
     max_value = jhb[column].max()

     summary_jhb = summary_jhb.append({'Column Name': column, 
                                 'Missing Values': missing_values, 
                                 'Min': min_value, 
                                 'Mean': mean_value, 
                                 'Median': median_value, 
                                 'Max': max_value}, ignore_index=True)
     
summary_jhb.to_excel(r'D:\Work\Information\三方对接\银豆\jhb_summary.xlsx')

for var in [col for col in jhb.columns if col  in score_var_all ]:
    num_bins = 10
    bin_edges = np.linspace(jhb[var].min(), jhb[var].max(), num_bins + 1)
    bin_edges_series = pd.Series(bin_edges)
    bins = list(pd.concat([pd.Series([float('-inf')]), bin_edges_series, pd.Series([float('inf')])]))
    temp_intervals = pd.cut(jhb[var], bins=bins, ordered=True)
    jhb[var+'_intervals'] = temp_intervals
    # 添加 'QS' 到类别中
    jhb[var+'_intervals'] = jhb[var+'_intervals'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    jhb[var+'_intervals'].fillna('DEFAULT', inplace=True)

jhb['target']=jhb.apply(lambda x: 1 if x.loan_status=='当前逾期' else 0  ,axis=1)

ks_target=pd.DataFrame()
for index,var in enumerate([col for col in jhb.columns if '_intervals' in col]):
    temp=ks_calc_cross_NB(jhb,var,'target')
    if index==0:
        ks_target=temp
    else:
        ks_target=pd.concat([ks_target,temp])
    ks_target.to_excel(r'D:\Work\Information\三方对接\银豆\jhb_ks.xlsx')     
        

#------------------------------------------------------------------------------ 1.4a、多头模式
base15_tengxun_dt=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试结果\测试样本15000_tengxun_raw.xlsx',sheet_name='行业风险-多头申请v3')
base15_tengxun_dt.drop(columns=['回溯日期','identity_no_md5','phone_md5'],inplace=True)

base15_dt=pd.read_excel(r'D:\Work\Information\三方对接\银豆\测试样本\测试样本_raw.xlsx',sheet_name='Sheet1')
a_dt=list(base15_dt.columns)
base15_dt=pd.merge(base15_dt,base15_tengxun_dt,how='left',on='new_column')

# 玖富微言检视
base15_dt_jf=base15_dt[base15_dt.source=='jf']
a_wy=base15_dt_jf.columns
cc_list=['最早一次查询距今的天数',
         '最近一次查询距今的天数',
'最早一次在银行机构距今的天数',
'最近一次在银行机构距今的天数',
'最早一次在非银机构距今的天数',
'最近一次在非银机构距今的天数']



cc_list_end=a_dt+cc_list
base15_dt_jf_js=base15_dt_jf[cc_list_end]

a_wy_dt=list(dt_wy.columns)
dt_wy_js=dt_wy[['user_id',
                'WY_最早一次查询距今的天数',
                'WY_最近一次查询距今的天数',
       'WY_最早一次在银行机构距今的天数',
       'WY_最近一次在银行机构距今的天数',
       'WY_最早一次在非银机构距今的天数',
       'WY_最近一次在非银机构距今的天数'
       ]]

base15_dt_jf_js_temp=pd.merge(base15_dt_jf_js,dt_wy_js,how='left',on='user_id')

base15_dt_jf_js_temp.to_excel(r'D:\Work\Information\三方对接\银豆\多头微言排查.xlsx',index=False)  

temp_js=base15_dt_jf_js_temp[['最近一次在非银机构距今的天数','WY_最早一次在非银机构距今的天数'
]]
temp=base15_dt_jf_js_temp[base15_dt_jf_js_temp.user_id==145705]

#------------------------------------------------------------------------------ 1.5、行业准入评估
base15['行业准入']=base15.apply(lambda x: 1 if pd.isna(x.v1_2) else 0,axis=1)
zr_df=base15['行业准入'].value_counts().reset_index().drop(columns='行业准入').rename(columns={'index':'行业准入'})

for df_source in ['jf','yxjfk','yxjjj','jhb']:
    temp_source=base15[base15.source==df_source]
    temp = temp_source.groupby('行业准入')['user_id'].count().reset_index().rename(columns={"user_id":'计数_'+df_source})
    temp['百分比_'+df_source] = temp['计数_'+df_source] / temp['计数_'+df_source].sum() 
    zr_df=pd.merge(zr_df,temp,how='left',on=['行业准入'])
zr_df.to_excel(r'D:\Work\Information\三方对接\银豆\行业准入.xlsx',index=False)  

lx_df=base15['lx_level'].value_counts().reset_index().drop(columns='lx_level').rename(columns={'index':'lx_level'})
for df_source in ['jf','yxjfk','yxjjj','jhb']:
    temp_source=base15[base15.source==df_source]
    temp = temp_source.groupby('lx_level')['user_id'].count().reset_index().rename(columns={"user_id":'计数_'+df_source})
    temp['百分比_'+df_source] = temp['计数_'+df_source] / temp['计数_'+df_source].sum() 
    lx_df=pd.merge(lx_df,temp,how='left',on=['lx_level'])
lx_df.to_excel(r'D:\Work\Information\三方对接\银豆\蓝象_level.xlsx',index=False)  


# 评分

for var in [col for col in base15.columns if col  in score_var_all+['lx_score'] ]:
    num_bins = 10
    bin_edges = np.linspace(base15[var].min(), base15[var].max(), num_bins + 1)
    bin_edges_series = pd.Series(bin_edges)
    bins = list(pd.concat([pd.Series([float('-inf')]), bin_edges_series, pd.Series([float('inf')])]))
    temp_intervals = pd.cut(base15[var], bins=bins, ordered=True)
    base15[var+'_intervals_pf'] = temp_intervals
    # 添加 'QS' 到类别中
    base15[var+'_intervals_pf'] = base15[var+'_intervals_pf'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    base15[var+'_intervals_pf'].fillna('DEFAULT', inplace=True)

# 拿表头
col_base_df=pd.DataFrame()
for index,var in enumerate([col for col in base15.columns if '_intervals_pf' in col]):
    temp=base15[var].value_counts().reset_index().drop(columns=var).rename(columns={'index':'Bins'})
    temp['var_name']=var
    if index==0:
        col_base_df=temp
    else:
        col_base_df=pd.concat([col_base_df,temp])
col_base_df.sort_values(by=['var_name','Bins'],inplace=True)  

# 拿完表头后分4个来源
for df_source in ['jf','yxjfk','yxjjj','jhb']:
    temp_source=base15[base15.source==df_source]
    for index,var in enumerate([col for col in base15.columns if '_intervals_pf' in col]):
        temp = temp_source.groupby(var)['user_id'].count().reset_index().rename(columns={"user_id":'计数_'+df_source,var:'Bins'})
        temp['百分比_'+df_source] = temp['计数_'+df_source] / temp['计数_'+df_source].sum() 
        temp['var_name']=var
        if index==0:
            temp_df=temp
        else:
            temp_df=pd.concat([temp_df,temp])
    col_base_df=pd.merge(col_base_df,temp_df,how='left',on=['var_name','Bins'])
    
col_base_df.to_excel(r'D:\Work\Information\三方对接\银豆\duotou_pf1.xlsx',index=False)  

        


#%%-----------talkingdata测试分析
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt

talkd_raw=pd.read_excel(r'D:\Work\Information\三方对接\测试数据\TD\样本2000_TD.xlsx')
talkd_raw['target15']=talkd_raw.apply(lambda x: 1 if x.od_days>=15  else 0 ,axis=1)#因为坏样本的逾期天数是用start_date往前推30天
talkd_raw['target30']=talkd_raw.apply(lambda x: 1 if x.od_days>=30 else 0 ,axis=1)

fpd15=fstq[['order_no','fpd15_fz']].rename(columns={'fpd15_fz':'target15_fpd'})
fpd30=fstq[['order_no','fpd30_fz']].rename(columns={'fpd30_fz':'target30_fpd'})
talkd_raw=pd.merge(talkd_raw,fpd15,on='order_no',how='left')
talkd_raw=pd.merge(talkd_raw,fpd30,on='order_no',how='left')
talkd_raw.drop(columns=['user_id','loan_channel_code',
                        'cycle','loan_time','od_days','id','real_name_md5',
                        'identity_no_md5','phone_md5'],inplace=True)

talkd_raw.info()

a_target30_fpd=talkd_raw.groupby('收入等级').agg({'target30_fpd':['count','sum']})
a_target15_fpd=talkd_raw.groupby('收入等级').agg({'target15_fpd':['count','sum']})
a_target15=talkd_raw.groupby('收入等级').agg({'target15':['count','sum']})
a_target30=talkd_raw.groupby('收入等级').agg({'target30':['count','sum']})

# 查看各个分数的相关系数
score_var=[col for col in talkd_raw.columns if col not in ['order_no', 'target15','target30','target15_fpd','target30_fpd']]
# score_var=['腾云分21'	,'腾云分87'	,'腾云分239',	'腾云分242','腾云分108','腾云分111']
# score_var=['腾云分111'	,'腾云分108'	]
corr_data=talkd_raw[score_var]
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)
rc = {'font.sans-serif': 'SimHei',
      'axes.unicode_minus': False}
sns.set(context='notebook', style='ticks', rc=rc)
# 创建相关系数矩阵的热图，并设置轴标签
plt.figure(figsize=(13, 13))
# sns.set(font_scale=1)
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", vmin=-1, vmax=1,square=True,
            xticklabels=score_var,
            yticklabels=score_var,
            cbar=False)
plt.title("相关系数矩阵热图")
plt.show()

plt.figure(figsize=(8, 6))
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", square=True, vmin=-1, vmax=1)
plt.title("Correlation Matrix Heatmap")
plt.show()


为什么相关系数矩阵热图有些地方是空白无值的

# 创建一个空的 DataFrame 用于存储结果
summary_df = pd.DataFrame(columns=['Column Name', 'Missing Values', 'Min', 'Mean', 'Median', 'Max'])


for column in [col for col in talkd_raw.columns if col not in ['order_no', 'target15','target30','target15_fpd','target30_fpd']]:
     missing_values = talkd_raw[column].isnull().sum()
     min_value = talkd_raw[column].min()
     mean_value = talkd_raw[column].mean()
     median_value = talkd_raw[column].median()
     max_value = talkd_raw[column].max()
 
 # 将结果添加到 summary_df 中
     summary_df = summary_df.append({'Column Name': column, 
                                 'Missing Values': missing_values, 
                                 'Min': min_value, 
                                 'Mean': mean_value, 
                                 'Median': median_value, 
                                 'Max': max_value}, ignore_index=True)
    
for var in [col for col in talkd_raw.columns if col not in ['order_no', 'target15','target30','target15_fpd','target30_fpd']]:
    print(var)
    num_bins = 10
    bin_edges = np.linspace(talkd_raw[var].min(), talkd_raw[var].max(), num_bins + 1)
    bin_edges_series = pd.Series(bin_edges)
    bins = list(pd.concat([pd.Series([float('-inf')]), bin_edges_series, pd.Series([float('inf')])]))
    temp_intervals = pd.cut(talkd_raw[var], bins=bins)
    talkd_raw[var+'_intervals'] = temp_intervals
    
def ks_calc_cross(data,score_col,class_col):
    '''
1、计算每个评分区间的好坏账户数（计算的是特征的KS的话，是每个特征对应的好坏账户数）。
2、计算每个评分区间的累计好账户数占总好账户数比率(good%)和累计坏账户数占总坏账户数比率(bad%)。
3、计算每个评分区间累计坏账户占比与累计好账户占比差的绝对值（累计good%-累计bad%），然后对这些绝对值取最大值即得此评分卡的KS值
    
功能：计算KS值，输出对应分割点和累计分布函数曲线图
    ----------
    输入值
    data : 二维数组或dataframe，包括模型得分和真实的标签
           
    score_col : 一维数组或series,代表模型得分(一般为预测正类的概率)
        
    class_col : 一维数组或series,代表真实的标签{{0,1}或{-1,1}}
       
    输出值：
    -------
    'ks':KS值,'crossdens':好坏人累积概率分布以及其差值gap

    '''
    crossfreq=pd.crosstab(data[score_col],data[class_col])
    crossfreq.rename(columns={0:'good',1:'bad'},inplace=True)
    crossdens=crossfreq.cumsum(axis=0)/crossfreq.sum()
    crossdens.rename(columns={'good':'good_cum_rate','bad':'bad_cum_rate'},inplace=True)
    crossdens['diff_rate']=abs(crossdens['good_cum_rate']-crossdens['bad_cum_rate'])
    max_value = crossdens['diff_rate'].max()
    crossdens['KS'] = [1 if x == max_value else 0 for x in crossdens['diff_rate']]

    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    crossfreq.reset_index(inplace=True)
    crossfreq.rename(columns={score_col:'Bins'},inplace=True)
    crossfreq['var_name']=score_col
    crossfreq['bad_rate_bins']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    # ks=crosssdens[crosssdens['gap']==crosssdens['gap'].max()]
    return crossfreq

ks_target30=pd.DataFrame()
for index,var in enumerate([col for col in talkd_raw.columns if '_intervals' in col]):
    temp=ks_calc_cross(talkd_raw,var,'target30')
    if index==0:
        ks_target30=temp
    else:
        ks_target30=pd.concat([ks_target30,temp])
        
ks_target30_fpd.to_excel(r'D:\Work\Information\三方对接\测试数据\TD\ks_target30_fpd.xlsx')
ks_target15_fpd.to_excel(r'D:\Work\Information\三方对接\测试数据\TD\ks_target15_fpd.xlsx')
ks_target30.to_excel(r'D:\Work\Information\三方对接\测试数据\TD\ks_target30.xlsx')
ks_target15.to_excel(r'D:\Work\Information\三方对接\测试数据\TD\ks_target15.xlsx')

#%%-----------冰鉴测试分析
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import seaborn as sns
bj_raw=pd.read_excel(r'D:\Work\Information\三方对接\测试数据\冰鉴\冰鉴-聚合吧科技有限公司-测试报告-2023-09-25\冰鉴数据源.xlsx')
bj_raw.info()
bj_raw['target7']=bj_raw.apply(lambda x: 1 if x.od_days>=7  else 0 ,axis=1)#因为坏样本的逾期天数是用start_date往前推30天
bj_raw['target1']=bj_raw.apply(lambda x: 1 if x.od_days>=1 else 0 ,axis=1)#因为坏样本的逾期天数是用start_date往前推30天
bj_raw['target15']=bj_raw.apply(lambda x: 1 if x.od_days>=15  else 0 ,axis=1)#因为坏样本的逾期天数是用start_date往前推30天
bj_raw['target30']=bj_raw.apply(lambda x: 1 if x.od_days>=30 else 0 ,axis=1)


# 去下面拿fpd15 跟 fpd30
fpd15=fstq[['order_no','fpd15_fz']].rename(columns={'fpd15_fz':'target15_fpd'})
fpd30=fstq[['order_no','fpd30_fz']].rename(columns={'fpd30_fz':'target30_fpd'})
bj_raw=pd.merge(bj_raw,fpd15,on='order_no',how='left')
bj_raw=pd.merge(bj_raw,fpd30,on='order_no',how='left')



# 定义bins，冰鉴的分数特定就是300-850，而且他们的分析也是按50拆分开，所以下面bins代码简单一点
bins = [i for i in range(300, 900, 50)]
bins[-1] = bins[-1]+1
bins.insert(0, float('-inf'))

def cut_scores(scores, interval=50):

    # 使用cut函数将评分分组到不同的区间中
    score_bins = pd.cut(scores, bins=bins, include_lowest=True)
    return score_bins

# 查看各个分数的相关系数
score_var=[col for col in bj_raw.columns if '版' in col]
corr_data=bj_raw[score_var]
corr_data_t=corr_data.T
correlation_matrix = np.corrcoef(corr_data.T)
rc = {'font.sans-serif': 'SimHei',
      'axes.unicode_minus': False}
sns.set(context='notebook', style='ticks', rc=rc)
# 创建相关系数矩阵的热图，并设置轴标签
plt.figure(figsize=(13, 13))
# sns.set(font_scale=1)
sns.heatmap(correlation_matrix, annot=True, cmap="coolwarm", fmt=".2f", 
            xticklabels=score_var,
            yticklabels=score_var,
            cbar=False)
plt.title("相关系数矩阵热图")
plt.show()





#bins
for var in [col for col in bj_raw.columns if '版' in col]:
    bj_raw[var+'_bins']=pd.cut(bj_raw[var], bins=bins, include_lowest=True,right=False)

#查看分布
# for index,var in enumerate([col for col in bj_raw.columns if 'bins' in col]):
#     temp=bj_raw[var].value_counts().to_frame().reset_index()
#     temp['varname']=var
#     temp.rename(columns={var:'values'},inplace=True)
#     if index==0:
#         temp_all=temp
#     else:
#         temp_all= pd.concat([temp_all,temp])

# 计算分箱后的ks
def ks_calc_cross(data,score_col,class_col):
    '''
1、计算每个评分区间的好坏账户数（计算的是特征的KS的话，是每个特征对应的好坏账户数）。
2、计算每个评分区间的累计好账户数占总好账户数比率(good%)和累计坏账户数占总坏账户数比率(bad%)。
3、计算每个评分区间累计坏账户占比与累计好账户占比差的绝对值（累计good%-累计bad%），然后对这些绝对值取最大值即得此评分卡的KS值
    
功能：计算KS值，输出对应分割点和累计分布函数曲线图
    ----------
    输入值
    data : 二维数组或dataframe，包括模型得分和真实的标签
           
    score_col : 一维数组或series,代表模型得分(一般为预测正类的概率)
        
    class_col : 一维数组或series,代表真实的标签{{0,1}或{-1,1}}
       
    输出值：
    -------
    'ks':KS值,'crossdens':好坏人累积概率分布以及其差值gap

    '''
    ks_dict={}
    crossfreq=pd.crosstab(data[score_col],data[class_col])
    crossfreq.rename(columns={0:'good',1:'bad'},inplace=True)
    crossdens=crossfreq.cumsum(axis=0)/crossfreq.sum()
    crossdens.rename(columns={'good':'good_cum_rate','bad':'bad_cum_rate'},inplace=True)
    crossdens['diff_rate']=abs(crossdens['good_cum_rate']-crossdens['bad_cum_rate'])
    max_value = crossdens['diff_rate'].max()
    crossdens['KS'] = [1 if x == max_value else 0 for x in crossdens['diff_rate']]

    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    crossfreq.reset_index(inplace=True)
    crossfreq.rename(columns={score_col:'Bins'},inplace=True)
    crossfreq['var_name']=score_col
    crossfreq['bad_rate_bins']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    # ks=crosssdens[crosssdens['gap']==crosssdens['gap'].max()]
    return crossfreq

ks_df=pd.DataFrame()
ks_df1=pd.DataFrame()
for cyc in [1,3,6,12]:
    bj_raw_temp=bj_raw[bj_raw.cycle==cyc].copy()
    for index,var in enumerate([col for col in bj_raw_temp.columns if 'bins' in col]):
        temp=ks_calc_cross(bj_raw_temp,var,'target7')
        temp['cycle']=cyc
        if index==0:
            ks_df1=temp
        else:
            ks_df1=pd.concat([ks_df1,temp])
    if cyc==1:
        ks_df=ks_df1
    else:
        ks_df=pd.concat([ks_df,ks_df1])
        


ks_df.to_excel(r'D:\Work\Information\三方对接\测试数据\冰鉴\冰鉴-聚合吧科技有限公司-测试报告-2023-09-25\ks_测试.xlsx')


ks_df.to_excel(r'D:\Work\Information\三方对接\测试数据\冰鉴\冰鉴-聚合吧科技有限公司-测试报告-2023-09-25\ks_测试_cycle.xlsx')
#%%-----------拒绝原因配置表导入到生产库-风控库
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text


os.chdir(r"D:\Work\out_data\PythonCode")
# 生产库 
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


table_pz = pd.read_excel(r"D:\Work\out_data\metabase\拒绝原因配置表.xlsx",sheet_name=r'Sheet1')

from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate?charset=utf8mb4'
# db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate?charset=utf8mb4&collation=utf8mb4_general_ci'
engine = create_engine(db_url)
table_pz.to_sql(name='reject_code', con=engine, if_exists='replace', index=False)




#%%-----------钜银贷三方报文分布-生产，基于各个request表取数，后续可以限定名单列表来看
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
import numpy as np
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 字段的中文解释
df_base = pd.read_excel(r"D:\Work\out_data\分析类\juin客户分布\df.xlsx",sheet_name=r'Sheet1')
df_base=df_base[df_base.COL2.notna()]
df_base['COL2']=df_base.COL2.apply(lambda x:x.lower())

# risk_credit_result只有通过才有额度，干脆直接解析报文来获取授信额度与提现额度
sql = """ select user_id,request_body  from juin_loan_core_prd.risk_request_apply  where process_node=0  """
sx=pd.read_sql(sql,cnx)

for i in range(len(sx)):
    temp_json=json.loads(sx['request_body'][i])
    temp = pd.DataFrame(temp_json['inputApplyParameter'],index=[i])
    temp['user_id']=sx['user_id'][i]
    if i==0:
        sx_df=temp
    else:
        sx_df=pd.concat([sx_df,temp])
sx_df=sx_df[['user_id','jfCreditLimit']]
sx_df['sx']=1

sql = """ select user_id,request_body  from juin_loan_core_prd.risk_request_apply  where process_node=1  """
tx=pd.read_sql(sql,cnx)

for i in range(len(tx)):
    try:
        temp_json=json.loads(tx['request_body'][i])
        temp = pd.DataFrame(temp_json['inputApplyParameter'],index=[i])
        temp['user_id']=tx['user_id'][i]
        if i==0:
            tx_df=temp
        else:
            tx_df=pd.concat([tx_df,temp])
    except Exception as e:
        pass
    
tx_df=tx_df[['user_id','withdrawAmount']]
tx_df=tx_df.groupby('user_id').agg({'withdrawAmount':'sum'}).reset_index()
tx_df['tx']=1

stx=pd.merge(sx_df,tx_df,how='left',on='user_id')

sql = """ select user_id,case when credit_result='通过' then 1 else 0 end as credit_result  from juin_loan_core_prd.risk_credit_result  where process_node=0  """
sxjg=pd.read_sql(sql,cnx)

stx=pd.merge(stx,sxjg,how='left',on='user_id')
stx['credit_amount']=stx.apply(lambda x:x.jfCreditLimit if x.credit_result==1 else 0,axis=1 )

# 后面的表现,等下个月再用
# sql = """ 	select  user_id2 as user_id,case when loan_status='当前逾期' then 1 else 0 end as target
# from juin_risk_operate.yx_fk_test1113_v2 where type=1  """
# list_sr=pd.read_sql(sql,cnx)
# list_sr.user_id=list_sr.user_id.astype('int64')


#-------------------------------------------------------------------------------------------------------------------微言
sql = """ 	select  user_id,flow_id,create_time,response_param from juin_loan_core_prd.third_weiyan_request_info where process_node='CREDIT' and status='0' and api_name='MULTIPOINT_LOAN' \
    and user_id in (45220,59756,106611,136506,148325,262885,303871)

"""
third_weiyan_request_info=pd.read_sql(sql,cnx)


# 初始化一个空的DataFrame

# 遍历 data_list 中的每个元素，将其转换为 DataFrame，并合并到 df 中
for i in range(third_weiyan_request_info.shape[0]):
    temp_json=json.loads(third_weiyan_request_info['response_param'][i])
    temp_json1=json.loads(temp_json['rule_detail']['factor_list'][0]['rule_list'][0]['variable_list'][0]['result'])
    temp_json2=temp_json1['MultipleLoansInfo']
    for k, data_dict in enumerate(temp_json2):
        temp_df = pd.DataFrame(data_dict, index=[0])
        temp_df1 = pd.DataFrame({'code'+str(temp_df.at[0, "riskCode"]):temp_df.at[0, "riskCodeValue"]}, index=[i])
        if k==0:
            dt_temp=temp_df1.copy()
        else:
            dt_temp= pd.concat([dt_temp, temp_df1], axis=1)
    dt_temp['user_id']=third_weiyan_request_info['user_id'][i]
    if i==0:
        dt_wy_raw=dt_temp.copy()
    else:
        dt_wy_raw= pd.concat([dt_wy_raw, dt_temp], axis=0)



dt_wy=dt_wy_raw.copy()
col_temp=pd.DataFrame({'col_name':dt_wy.columns})
col_temp['col_name']=col_temp['col_name'].apply(lambda x:x.replace("code",""))
dt_wy_pzb = pd.read_excel(r"D:\Work\out_data\分析类\juin客户分布\多头配置表_wy.xlsx",sheet_name=r'Sheet1')
dt_wy_pzb['code']=dt_wy_pzb['code'].astype('str')
col_temp=pd.merge(col_temp,dt_wy_pzb,how='left',left_on='col_name',right_on='code')
col_temp['value']=col_temp.apply(lambda x: 'user_id' if x.col_name=='user_id' else x.value,axis=1)
dt_wy.columns=col_temp.value
dt_wy_raw.columns=col_temp.value

dt_wy=pd.merge(dt_wy,stx,how='inner',on='user_id')

col_model=[attr for attr in list(dt_wy.columns) if  attr not in list(stx.columns)]

for var in col_model:
    dt_wy[var]=dt_wy[var].astype(float)

a_wy=dt_wy.describe().T


def bin_variable( var_df,var_name, n_bins):#根据当前最大最小值来等距分箱
    # 计算变量的最大值和最小值
    var_min = var_df[var_name].min()
    var_max = var_df[var_name].max()
    # 根据最大值和最小值以及分箱数量计算分箱区间
    bins = np.linspace(var_min, var_max, n_bins + 1)
    bins=np.insert(bins,0, float('-inf'))
    # 对变量进行分箱，并将分箱后的新变量添加到 DataFrame 中
    var_df[var_name] = pd.cut(var_df[var_name], bins,duplicates='drop')
   
for var_name in col_model:
    bin_variable( dt_wy,var_name, n_bins=10)

# 参照数据统计数据
result_dt=pd.DataFrame()
for (i,col) in enumerate(col_model):
    crossfreq=dt_wy.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result_dt=crossfreq
    else:
        result_dt=pd.concat([result_dt,crossfreq])
         
result_dt['三方']='微言'
result_dt_wy=result_dt.copy()

#-------------------------------------------------------------------------------------------------------------------TD
#----------------------------------------------------------------全量标签，全量标签挺有意思，命中率比金融属性高
sql = """ 	select  user_id,flow_id,create_time,response_param from juin_loan_core_prd.third_td_request_info 
where process_node='CREDIT' and status='0' and api_name='pid_full_tags_query'  \
    and user_id in (45220,59756,106611,136506,148325,262885,303871)
"""
third_td_request_info_pf=pd.read_sql(sql,cnx)

for i in range(third_td_request_info_pf.shape[0]):
    temp_json=json.loads(third_td_request_info_pf['response_param'][i])
    temp_json1=temp_json['data']['tags']
    print(i)
    # try:
    for key, value in temp_json1.items():
        for k in range(len(value)):
            temp_df = pd.DataFrame(value[k], index=[i])
            temp_df1 = pd.DataFrame({f"{key}_{temp_df.iloc[0, 0]}_{temp_df.iloc[0, 2]}":temp_df.iloc[0, 1]}, index=[i])
            if k==0:
                dt_temp=temp_df1.copy()
            else:
                dt_temp= pd.concat([dt_temp, temp_df1], axis=1)
        if key==next(iter(temp_json1)):
            dt_temp1=dt_temp.copy()
        else:
            dt_temp1= pd.concat([dt_temp1, dt_temp], axis=1)
        dt_temp1['user_id']=third_td_request_info_pf['user_id'][i]
    if i==0:
        dt_tdpf_raw=dt_temp1
    else:
        # dt_tdpf_raw= pd.concat([dt_tdpf_raw.reset_index(drop=True), dt_temp1], axis=0)
        dt_tdpf_raw= pd.concat([dt_tdpf_raw, dt_temp1], axis=0)
    # except Exception as e:
    #     continue

dt_tdpf=dt_tdpf_raw.copy()

a=dt_tdpf_raw.isnull().sum(axis=0)/dt_tdpf_raw.shape[0]#判断pd内部各列缺失值占比
a_df=pd.DataFrame(a)
a_df.rename(columns={0:'缺失率'},inplace=True)
a_df.sort_values(by='缺失率',inplace=True)

dt_tdpf=pd.merge(dt_tdpf,stx,how='inner',on='user_id')
col_model=[attr for attr in list(dt_tdpf.columns) if  attr not in list(stx.columns)]

for var in col_model:
    dt_tdpf[var].fillna("-1",inplace=True)
    dt_tdpf[var]=dt_tdpf[var].astype(float)


a=dt_tdff.isnull().sum(axis=0)/dt_tdff.shape[0]#判断pd内部各列缺失值占比

#--------------------------------------------- 金融属性:很多变量的缺失率高，剩下20个以内的变量看起来有点用
sql = """ 	select  user_id,flow_id,create_time,response_param from juin_loan_core_prd.third_td_request_info 
where process_node='CREDIT' and status='0' and api_name='financial-features-p' \
    and user_id in (45220,59756,106611,136506,148325,262885,303871) """
third_td_request_info_ff=pd.read_sql(sql,cnx)

for i in range(third_td_request_info_ff.shape[0]):
    temp_json=json.loads(third_td_request_info_ff['response_param'][i])
    dt_temp=pd.DataFrame(temp_json['data'],index=[i])
    dt_temp['user_id']=third_td_request_info_ff['user_id'][i]
    if i==0:
        dt_tdff_raw=dt_temp
    else:
        dt_tdff_raw= pd.concat([dt_tdff_raw, dt_temp], axis=0)

dt_tdff=dt_tdff_raw.copy()
col_temp=pd.DataFrame({'col_name':dt_tdff.columns})
dt_td_pzb = pd.read_excel(r"D:\Work\out_data\分析类\juin客户分布\多头配置表_td.xlsx",sheet_name=r'Sheet1')

col_temp=pd.merge(col_temp,dt_td_pzb,how='left',left_on='col_name',right_on='code')
col_temp['value']=col_temp.apply(lambda x: 'user_id' if x.col_name=='user_id' else x.value,axis=1)
col_temp['value']=col_temp.apply(lambda x: x.col_name if pd.isna(x.value) else x.value,axis=1)

# 肉眼检视缺失字段的缺失情况，TD的变量衍生还是挺快的
# col_temp['缺失']=col_temp.apply(lambda x: 1 if pd.isna(x.value) else 0 ,axis=1)
# col_none_df=col_temp[col_temp.缺失==1]
# col_none_df1=dt_tdff[col_none_df.col_name]

dt_tdff.columns=col_temp.value
dt_tdff_raw.columns=col_temp.value

a=dt_tdff.isnull().sum(axis=0)/dt_tdff.shape[0]#判断pd内部各列缺失值占比
#-------------------------------------------------------------------------------------------------------------------百融

sql = """ 	select  user_id,flow_id,create_time,response_param from juin_loan_core_prd.third_bairong_request_info 
where process_node='CREDIT' and status='0' and api_name='ApplyEvaluate' \
    and user_id in (45220,59756,106611,136506,148325,262885,303871) """
third_bairong_request_info=pd.read_sql(sql,cnx)

# temp_json=json.loads(third_bairong_request_info['response_param'][0])
# dt_temp=pd.DataFrame(temp_json,index=[0])

for i in range(third_bairong_request_info.shape[0]):
    temp_json=json.loads(third_bairong_request_info['response_param'][i])
    dt_temp=pd.DataFrame(temp_json,index=[i])
    dt_temp['user_id']=third_bairong_request_info['user_id'][i]
    if i==0:
        dt_bairong_raw=dt_temp.copy()
    else:
        dt_bairong_raw= pd.concat([dt_bairong_raw, dt_temp], axis=0)

attr=[attr for attr in list(dt_bairong_raw.columns) if 'Rule' not in attr]
attr=[attr for attr in attr if 'rule' not in attr]
attr=[attr for attr in attr if 'rs_' not in attr]
attr=[attr for attr in attr if attr not in ['flag_riskstrategy','code','swift_number']]
dt_bairong_raw=dt_bairong_raw[attr]


dt_bairong=dt_bairong_raw.copy()
col_temp=pd.DataFrame({'col_name':dt_bairong.columns})
col_temp['col_name']=col_temp['col_name'].apply(lambda x:x.lower())
dt_br_pzb = pd.read_excel(r"D:\Work\out_data\分析类\juin客户分布\多头配置表_百融.xlsx",sheet_name=r'Sheet1')
col_temp=pd.merge(col_temp,dt_br_pzb,how='left',left_on='col_name',right_on='code')
col_temp['value']=col_temp.apply(lambda x: 'user_id' if x.col_name=='user_id' else x.value,axis=1)

dt_bairong.columns=col_temp.value
dt_bairong_raw.columns=col_temp.value



dt_bairong=pd.merge(dt_bairong,stx,how='inner',on='user_id')





# -------------------------------------------------------------------------------------------------融360
# 多头指标有异常,没一个是超过1的
sql = """ 	select  user_id,flow_id,create_time,response_param from juin_loan_core_prd.third_rong360_request_info where process_node='CREDIT' and status='0' and api_name='MULTILOANULTIMATEV2REPORT' """
third_rong360_request_info=pd.read_sql(sql,cnx)

for i in range(third_rong360_request_info.shape[0]):
    try:
        temp_json=json.loads(third_rong360_request_info['response_param'][i])
        dt_temp=pd.DataFrame(temp_json['tianji_api_agenttj_multiloanultimatev2report_response']['feature_list'],index=[i])
        dt_temp['user_id']=third_rong360_request_info['user_id'][i]
        if i==0:
            dt_rong360_raw=dt_temp
        else:
            dt_rong360_raw= pd.concat([dt_rong360_raw, dt_temp], axis=0)
    except Exception as e:
        pass

dt_rong360=dt_rong360_raw.copy()
col_temp=pd.DataFrame({'col_name':dt_rong360.columns})
dt_rong360_pzb = pd.read_excel(r"D:\Work\out_data\分析类\juin客户分布\多头配置表_rong360.xlsx",sheet_name=r'Sheet1')

col_temp=pd.merge(col_temp,dt_rong360_pzb,how='left',left_on='col_name',right_on='code')
col_temp['value']=col_temp.apply(lambda x: 'user_id' if x.col_name=='user_id' else x.value,axis=1)
dt_rong360.columns=col_temp.value
dt_rong360_raw.columns=col_temp.value

drop_list=['三个月内Z类需求大于A等级订单的平均Z类需求',
           '三个月内申请期限标准差',
              '一年内出现过待生效状态订单的A等级订单的平均Z类需求',
              '一年内待生效订单总成功Z类需求度',
              '一年内待生效状态订单的平均Z类需求',
              '一年内审批中状态订单中小额订单平均Z类需求',
              '一年内待生效订单C等级订单占比',
              '一年内审批中状态订单A类订单占比',
              '三个月内Z类需求标准差',
              '三个月内B等级平均Z类需求',
              '历史最新资质能力类型',
              '历史A类资质能力是否为Ⅰ 等级',
                '一年内个人情况种类量级',
                '一年平均C类资质能力',
                '一年B类资质能力频率',
                '一年资质发生变动的种类',
                '历史B类资质能力等级',
                '历史资质发生变动的种类',
                '一年平均B类资质能力',
                '一年是否命中A类保障A类',
                '一个月内B类平均Z类能力',
                '一个月内平均C类资质能力',
                '一个月内是否命中D类资质能力',
                '半年内平均B类资质能力',
                '半年内平均C类资质能力',
                '半年内个人情况种类等级',
                '半年内资质发生变动种类',
                '三个月内最大异常期限',
                '三个月内申请活跃度',
                '三个月内负面订单平均确认Z类需求等级',
                '历史平均C类资质能力',
                '历史是否有B类Z类能力',
                '一年内已生效订单的总成功需求度',
                '历史D类资质能力变动频率',
                '历史A类资质能力类型变动频率',
                '历史C类资质能力是否发生变动',
                '历史C类资质能力获取频率',
                '历史E类资质情况种类数等级',
                'E类资质情况最新等级',
                '历史平均B类资质能力',
                '历史C类资质能力平均等级',
                'A类保障等级',
                '历史A类保障类A能力变动频率',
                '社会参与度最新评估',
                '历史A类保障C类获取频率',
                '历史A类保障C类变动频率',
                'B类保障等级',
                'B类保障变动频率',
                '半个月内待生效订单占比',
                '半个月内出现过待生效状态订单的A等级订单平均Z类需求',
                '半个月内拒绝订单小额订单量',
                '半个月内创建A类订单占比',
                '半年内创建订单中平均申请期限',
                '半年内异常订单数量异常等级A',
                '半年内创建订单中Z类需求标准差',
                '半年内Z类需求为B等级订单的平均情况',
                '半年内成功订单中C等级订单需求频率',
                '半年内负面订单占比',
                '一年内已生效订单机构占比',
'一年内成功订单总Z类需求',
'一年内创建订单中最大异常时长',
'一年内待生效状态订单中涉及总Z类需求',
'一年内待生效A类订单占比',
'一年内待生效订单的总期限',
'一年内负面订单机构占比',
'首次申请订单距今时长期限',
'首次成功订单距今期限',
'历史首个已生效订单距今时长期限',
'一年内最近订单距今期限',
'一年内最近异常订单距今期限',
'一年内失败订单的Z类需求为C等级订单占比',
'一年内正常关闭订单比率',
'一年内失败订单占比',
'一年内拒绝订单的Z类需求标准差',
'一年内已生效订单的总期限',
'一年内已生效订单涉及总Z类需求',
'一年内最近成功订单距今期限',
'一年申请活跃度',
'一个月内正常关闭订单比率',
'一个月内申请活跃度',
'一年内平均Z类需求',
'一年内B等级Z类需求订单占比',
'半年内待生效订单最大需求额度',
'半年内审批通过的订单中小额订单量等级',
'半年内待生效订单总Z类需求',
'半年内出现过审批中状态订单占比',
'一年内创建订单中Z类需求大于A等级订单的平均Z类需求',
'一年内创建订单中最小期限',
'一年内完整信息订单最小期限',
'一年内完整信息订单Z类需求',
'一年内负面订单平均Z类需求',
'一年内负面订单涉及Z类需求标准差',
'一年内负面订单总Z类需求',
'一年内已生效订单的Z类需求总等级',
'一年内负面订单的申请期限标准差',
'一年内待生效状态订单的B等级的平均订单Z类需求',
'一年内待生效状态订单的Z类需求最大量',
'一年内已生效订单最大Z类需求',
'一年内已生效订单平均Z类需求',
'一年内负面订单的最高Z类需求',
'半年内异常时长总等级',
'半年内异常期限',
'半年内负面订单机构占比',
'半年内申请活跃度',
'半年内正常关闭订单比率',
'半年内失败订单的最小期限',
'半年内已生效订单的最大Z类需求',
'一年内成功订单中C等级订单占比',
'一年内申请成功的平均Z类需求',
'一年内负面总Z类需求',
'一年内负面订单占比',
'一年内待生效B类订单量等级',
'一年内异常订单比率',
'一年内成功订单的最大Z类需求最',
'一年内最早异常订单距今期限',
'一年内负面订单中C等级订单占比',
'一年内C等级负面订单量等级',
'半年内已生效的订单占比',
'半年内负面订单确认E类资质等级标准差',
'半年内负面订单平均Z类需求等级'
              ]
dt_rong360.drop(columns=drop_list,inplace=True)
dt_rong360_raw.drop(columns=drop_list,inplace=True)


dt_rong360=pd.merge(dt_rong360,stx,how='inner',on='user_id')
col_model=[attr for attr in list(dt_rong360.columns) if  attr not in list(stx.columns)]

for var in col_model:
    dt_rong360.replace("","-1",inplace=True)
    dt_rong360[var]=dt_rong360[var].astype(float)
    
a=dt_rong360.describe().T
a.to_excel(r'D:\Work\out_data\分析类\juin客户分布\融360_多头describe.xlsx')

a_test=dt_rong360[['按手机号查询中短期在非银机构申请机构数']]
a_test1=dt_rong360[dt_rong360.按手机号查询中短期在非银机构申请机构数==0.1275]
#-------------------------------------------------------------------------------------------------------------------电话邦
又多了标签 catnames，与张韧沟通后表示他们的产品特性就是一直更新标签及对应字段




#-------------------------------------------------------------------------------------------------------------------银联智策
跟同盾都是朴道接口，因为前期沟通问题，没有报文积累
#-------------------------------------------------------------------------------------------------------------------同盾
tactics_score




#-------------------------------------------------------------------------------------------------------------------华瑞人行(还没调)







#%%-----------用户画像，基于urule入参中文名取各个result表的变量来看分布

# 这里预留了个代码，方便后续有还款target再优化，但优化工作量不小
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
import numpy as np
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 字段的中文解释
df_base = pd.read_excel(r"D:\Work\out_data\分析类\juin客户分布\df.xlsx",sheet_name=r'Sheet1')
df_base=df_base[df_base.COL2.notna()]
df_base['COL2']=df_base.COL2.apply(lambda x:x.lower())

# risk_credit_result只有通过才有额度，干脆直接解析报文来获取授信额度与提现额度
sql = """ select user_id,request_body  from juin_loan_core_prd.risk_request_apply  where process_node=0  """
sx=pd.read_sql(sql,cnx)

for i in range(len(sx)):
    temp_json=json.loads(sx['request_body'][i])
    temp = pd.DataFrame(temp_json['inputApplyParameter'],index=[i])
    temp['user_id']=sx['user_id'][i]
    if i==0:
        sx_df=temp
    else:
        sx_df=pd.concat([sx_df,temp])
sx_df=sx_df[['user_id','jfCreditLimit']]
sx_df['sx']=1

sql = """ select user_id,request_body  from juin_loan_core_prd.risk_request_apply  where process_node=1  """
tx=pd.read_sql(sql,cnx)

for i in range(len(tx)):
    temp_json=json.loads(tx['request_body'][i])
    temp = pd.DataFrame(temp_json['inputApplyParameter'],index=[i])
    temp['user_id']=tx['user_id'][i]
    if i==0:
        tx_df=temp
    else:
        tx_df=pd.concat([tx_df,temp])
tx_df=tx_df[['user_id','withdrawAmount']]
tx_df=tx_df.groupby('user_id').agg({'withdrawAmount':'sum'}).reset_index()
tx_df['tx']=1

stx=pd.merge(sx_df,tx_df,how='left',on='user_id')

sql = """ select user_id,case when credit_result='通过' then 1 else 0 end as credit_result  from juin_loan_core_prd.risk_credit_result  where process_node=0  """
sxjg=pd.read_sql(sql,cnx)

stx=pd.merge(stx,sxjg,how='left',on='user_id')
stx['credit_amount']=stx.apply(lambda x:x.jfCreditLimit if x.credit_result==1 else 0,axis=1 )

# 后面的表现,等下个月再用
# sql = """ 	select  user_id2 as user_id,case when loan_status='当前逾期' then 1 else 0 end as target
# from juin_risk_operate.yx_fk_test1113_v2 where type=1  """
# list_sr=pd.read_sql(sql,cnx)
# list_sr.user_id=list_sr.user_id.astype('int64')
#----------------------------------------------------------------------------------------------------------------------------------------------------申请信息


# user_work_info
sql = """ 	select user_id,income from juin_loan_core_prd.user_work_info     """
user_work_info=pd.read_sql(sql,cnx)

user_work_info['income']=user_work_info.income.map({2: '3000-5000', 
                                                    5: '12000以上',
                                                    4: '8000-120000', 
                                                    1: '3000以内'
                                                    })


# user_contact_person,要么(3,9,10) 要么(9,10),暂时不能用
# sql = """ 	select user_id, "1" as peiou from juin_loan_core_prd.user_contact_person  where relationship=3    """
# user_contact_person_peiou=pd.read_sql(sql,cnx)
# user_contact_person_peiou.drop_duplicates(subset='user_id',inplace=True)

# sql = """ 	select user_id, "1" as mm from juin_loan_core_prd.user_contact_person  where relationship=9   """
# user_contact_person_mm=pd.read_sql(sql,cnx)
# user_contact_person_mm.drop_duplicates(subset='user_id',inplace=True)

# user_personal_info,有些用户的邮箱是乱填的
sql = """ 	select user_id,education_level, marital_status, stirps from juin_loan_core_prd.user_personal_info     """
user_personal_info=pd.read_sql(sql,cnx)

user_personal_info['education_level']=user_personal_info.education_level.map({2: '初中', 
                                                                              3: '高中',
                                                                              4: '职高/中专/技术学校', 
                                                                              5: '大专', 
                                                                              6: '本科', 
                                                                              7: '硕士'
                                                                              })

user_personal_info['marital_status']=user_personal_info.marital_status.map({2: '已婚', 
                                                                              3: '离异',
                                                                              4: '其他', 
                                                                              1: '未婚'
                                                                              })
user_personal_info['stirps']=user_personal_info.stirps.map({"6": '苗族', 
                                                            "3": '回族',
                                                            "7": '彝族', 
                                                            "1": '汉族',
                                                            "10": '朝鲜族',
                                                            "11": '满族',
                                                            "15": '土家族',
                                                            "22": '畲族',
                                                            "27": '纳西族'
                                                            })

# order_record,这个表是提现表，所以暂时不需要了
# sql = """ 	select user_id,device_type from juin_loan_core_prd.order_record   """
# order_record=pd.read_sql(sql,cnx)
# order_record.drop_duplicates(subset='user_id',inplace=True)

# apply解析 
sql = """ 	select * from juin_loan_core_prd.risk_request_apply where process_node=0   """
risk_request_apply=pd.read_sql(sql,cnx)


for i in range(len(risk_request_apply)):
    temp_json=json.loads(risk_request_apply['request_body'][i])
    inputApplyParameter_temp = pd.DataFrame(temp_json['inputApplyParameter'],index=[i])
    inputApplyParameter_temp['user_id'] = risk_request_apply['user_id'][i]
    if i==0:
        inputApplyParameter = inputApplyParameter_temp.copy()
    else:
        inputApplyParameter= pd.concat([inputApplyParameter,inputApplyParameter_temp])
inputApplyParameter=inputApplyParameter[['user_id', 'age', 'sex',  'ruralRigister', 'nativeOrigin', 'nativeRegister']]


inputApplyParameter['sex']=inputApplyParameter.sex.map({"1": '男', "2": '女'})
inputApplyParameter['nativeOrigin']=inputApplyParameter.nativeOrigin.map({"0": '不一致', "1": '一致'})
inputApplyParameter['nativeRegister']=inputApplyParameter.nativeRegister.map({"0": '不一致', "1": '一致'})



apply_info=pd.merge(inputApplyParameter,user_personal_info,on='user_id',how='left')
apply_info=pd.merge(apply_info,user_work_info,on='user_id',how='left')
apply_info=pd.merge(apply_info,stx,how='inner',on='user_id')

# apply_info=pd.merge(apply_info,user_contact_person_mm,on='user_id',how='left')
# apply_info=pd.merge(apply_info,user_contact_person_peiou,on='user_id',how='left')


bins_temp = [-float('inf'), 0,21.1,25.1,30.1,35.1,40.1,45.1,50.1,55.1, float('inf')]
apply_info['age'] = pd.cut(apply_info['age'], bins=bins_temp, right=False)



col_model=[attr for attr in list(apply_info.columns) if attr  not in ['user_id']]
col_model=[attr for attr in col_model if attr  not in list(stx.columns)]
# 参照数据统计数据
result=pd.DataFrame()
for (i,col) in enumerate(col_model):
    crossfreq=apply_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.map({"ruralRigister": '是否农户_户籍', 
                                    "nativeOrigin": '身份证前六位地址\家庭地址\单位地址是否任一一致(精确到市)',
                                    "nativeRegister": '户籍地址\家庭地址\单位地址是否任一一致(精确到市)', 
                                    "stirps": '客户民族',
                                    "income": '客户收入',
                                    "age": '客户年龄',
                                    'education_level':'教育程度',
                                    'marital_status':'婚姻状况',
                                    'sex':'性别'
                                    })
apply_result=result.copy()
apply_result['三方']='申请变量'

#-----------------------------------------------------------------------------------------------------------------------------------------------------微言
sql = """ 	select  * from juin_loan_core_prd.weiyan_result_info where process_node='CREDIT'  """
weiyan_result_info=pd.read_sql(sql,cnx)
weiyan_result_info.drop(columns=['id','flow_id','create_time',
                                 'update_time','is_delete','failure_time',
                                 'process_node'],inplace=True)
weiyan_result_info=pd.merge(weiyan_result_info,stx,how='inner',on='user_id')
# weiyan_result_info=pd.merge(weiyan_result_info,list_sr,how='inner',on='user_id')#有逾期表现时再用

# score类
socre_list=['fxpty_wy',
'fxp_cj_hjfv4_wy',
'zzpf_lspg_v2_wy',
'qt_ndf_v1_wy',
'qt3_ndf_v1_wy',
'dtsq_score_wy',
'dzqdtgz_score_wy',
'czqdtgz_score_wy',
'fyhdtgz_score_wy',
'yhdtgz_score_wy']

duotou_list1=[
'apply_7d_count_wy',
'apply_7d_plat_count_wy',
'apply_15d_count_wy',
'apply_15d_plat_count_wy',
'apply_30d_count_wy',
'apply_30d_plat_count_wy'
    ]

# Step1 好像数据库都是默认var类型，所以先剔除score类去看变量的枚举值，枚举值大的再修改为int类型去跑describe
col_model=[attr for attr in list(weiyan_result_info.columns) if  attr not in socre_list]
col_model=[attr for attr in col_model if  attr not in duotou_list1]
col_model=[attr for attr in col_model if  attr not in list(stx.columns)]

# 参照数据统计数据
result=pd.DataFrame()
for (i,col) in enumerate(col_model):
    crossfreq=weiyan_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    # 后续只需要在上面的agg里定义good bad就可以跑下面代码，同理再去score那里类似参考且类似修改
    # crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    # crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    # crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    # crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
   
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)
# Step2 数值型 基本都是X天内申请次数，但最大值都在20以内

for var in duotou_list1:
    weiyan_result_info[var]=weiyan_result_info[var].astype(int)

def bin_variable( var_df,var_name, n_bins):#根据当前最大最小值来等距分箱
    # 计算变量的最大值和最小值
    var_min = var_df[var_name].min()
    var_max = var_df[var_name].max()
    # 根据最大值和最小值以及分箱数量计算分箱区间
    bins = np.linspace(var_min, var_max, n_bins + 1)
    bins=np.insert(bins,0, float('-inf'))
    # 对变量进行分箱，并将分箱后的新变量添加到 DataFrame 中
    var_df[var_name] = pd.cut(var_df[var_name], bins)
   
for var_name in duotou_list1:
    bin_variable( weiyan_result_info,var_name, n_bins=10)
    
# col_binned=[attr for attr in list(weiyan_result_info.columns) if  attr.endswith('_binned')]

# 参照数据统计数据
result_dt=pd.DataFrame()
for (i,col) in enumerate(duotou_list1):
    crossfreq=weiyan_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})

    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.sort_values(by='取值',inplace=True)
    if i==0:
        result_dt=crossfreq
    else:
        result_dt=pd.concat([result_dt,crossfreq])
        
# result_dt['变量名']=result_dt.变量名.apply(lambda x:x.replace("_binned",""))       
result_dt['变量名']=result_dt.变量名.apply(lambda x:x.replace("_",""))
result_dt=pd.merge(result_dt,df_base,how='left',left_on='变量名',right_on='COL2')
result_dt.COL1=result_dt.apply(lambda x: '15天内总申请平台数' if x.变量名=='apply15dplatcountwy' else x.COL1,axis=1)
result_dt.drop(columns=['变量名','COL2'],inplace=True)
result_dt.rename(columns={'COL1':'变量名'},inplace=True)

# Step3 score类
score_df=weiyan_result_info[socre_list]
# score_df=pd.concat([score_df,weiyan_result_info.target],axis=1)

for var in socre_list:
    weiyan_result_info[var]=weiyan_result_info[var].astype(int)
    
bins1 = [i for i in range(300, 950, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 101, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=['fxpty_wy','fxp_cj_hjfv4_wy','zzpf_lspg_v2_wy']

#bins1
for var in socre_list1:
    weiyan_result_info[var]=pd.cut(weiyan_result_info[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    weiyan_result_info[var]=pd.cut(weiyan_result_info[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=weiyan_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    
    # 后续只需要在上面的agg里定义good bad就可以跑下面代码，同理再去score那里类似参考且类似修改
    # crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    # crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    # crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    # crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
        
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)
wy_result=pd.concat([result_score,result],axis=0)
wy_result=pd.concat([wy_result,result_dt],axis=0)
wy_result['三方']='微言'


# wy_result.to_excel(r'D:\Work\out_data\分析类\特殊名单\微言.xlsx',index=False)

#--------------------------------------------------------------------------------------------------------------------------------------------------talkingdata
sql = """ 	select  user_id,tengyun108_td,tengyun239_td from juin_loan_core_prd.td_result_info where process_node='CREDIT'  """
td_result_info=pd.read_sql(sql,cnx)
td_result_info=pd.merge(td_result_info,stx,how='inner',on='user_id')
# td_result_info=pd.merge(td_result_info,list_sr,how='inner',on='user_id')
socre_list=['tengyun108_td','tengyun239_td']

# Step3 score类
for var in socre_list:
    td_result_info[var]=td_result_info[var].astype(float)

# a=score_df.describe()
bins1 = [i for i in range(300, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 101, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=['tengyun108_td']

#bins1
for var in socre_list1:
    td_result_info[var]=pd.cut(td_result_info[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    td_result_info[var]=pd.cut(td_result_info[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=td_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
        
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)

td_result=result_score.copy()
td_result['三方']='TD'


# result_score.to_excel(r'D:\Work\out_data\分析类\特殊名单\TD.xlsx',index=False)

#-------------------------------------------------------------------------------------------------------------------------------------------------------冰鉴
sql = """ 	select * from juin_loan_core_prd.icekredit_result_info where process_node='CREDIT'  """
icekredit_result_info=pd.read_sql(sql,cnx)
icekredit_result_info=icekredit_result_info[['user_id','hui_yan_score','qing_yun_score','huo_mou_score','xing_yu_score','hao_yue_score']]
icekredit_result_info=pd.merge(icekredit_result_info,stx,how='inner',on='user_id')

socre_list=['hui_yan_score',
'qing_yun_score',
'huo_mou_score',
'xing_yu_score']

# Step3 score类
# 填补缺失项：-1,转换数据类型为int或者float
for var in socre_list:
    icekredit_result_info[var].fillna('-1',inplace=True)
    icekredit_result_info[var]=icekredit_result_info[var].astype(int)

bins1 = [i for i in range(300, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

#bins1
for var in socre_list:
    icekredit_result_info[var]=pd.cut(icekredit_result_info[var], bins=bins1, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=icekredit_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])

result_score.变量名.replace({'qing_yun_score':'青云分',
                             'hui_yan_score':'慧眼分',
                             'huo_mou_score':'火眸分',
                             'xing_yu_score':'星宇分'},inplace=True)
bj_result=result_score.copy()
bj_result['三方']='冰鉴'


# result_score.to_excel(r'D:\Work\out_data\分析类\特殊名单\冰鉴.xlsx',index=False)

#--------------------------------------------------------------------------------------------------------------------------------------------------------融360
# 有时间去排雷，360的涉诉貌似在这个表里不重复
sql = """ 	select  * from juin_loan_core_prd.rong360_result_info  where process_node='CREDIT' """
rong360_result_info=pd.read_sql(sql,cnx)
rong360_result_info['ssxq']=rong360_result_info.apply(lambda x: "1" if len(x.flssxq_r360)>5 else "0",axis=1)
rong360_result_info.drop(columns=['id','flow_id','create_time','update_time','is_delete',
                                  'failure_time','process_node','flsshmd_r360','flssxq_r360',
                                  
                                  'pjhf_j3m_dw_r360',
                                 
                                  'pjll_j3m_dw_r360',
                                
                                  'pjthsc_j3m_dw_r360','fmddzb_6m_r360','fmddjgzb_r360'],inplace=True)
rong360_result_info=pd.merge(rong360_result_info,stx,how='inner',on='user_id')



# score类
socre_list=['zy_score_r360',
'zxqs_score_r360',
'zxsjwl_score_r360',
'zxjc_r360',
'zxxdfxyz_r360',
'zxxwfx_r360',
'wddt_score_r360']

# Step1 好像数据库都是默认var类型，所以先剔除score类去看变量的枚举值，枚举值大的再修改为int类型去跑describe
col_model=[attr for attr in list(rong360_result_info.columns) if  attr not in socre_list]
col_model=[attr for attr in col_model if  attr not in list(stx.columns)]


# var_df=rong360_result_info[col_model]
# var_df.drop(columns=['user_id','target'],inplace=True)

# 参照数据统计数据
result=pd.DataFrame()
for (i,col) in enumerate(col_model):
    crossfreq=rong360_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.COL1=result.apply(lambda x: '涉诉有内容' if x.变量名=='ssxq' else x.COL1,axis=1)
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)

# Step2 数值型 基本都是X天内申请次数，但最大值都在20以内

# Step3 score类

for var in socre_list:
    rong360_result_info[var].fillna("-1",inplace=True)
    rong360_result_info[var]=rong360_result_info[var].astype(int)

bins1 = [i for i in range(300, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 101, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=[
'zxqs_score_r360',
'zxsjwl_score_r360',
'zxjc_r360',
'zxxdfxyz_r360',
'zxxwfx_r360']

#bins1
for var in socre_list1:
    rong360_result_info[var]=pd.cut(rong360_result_info[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    rong360_result_info[var]=pd.cut(rong360_result_info[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=rong360_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
        
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)
r360_result=pd.concat([result_score,result],axis=0)
r360_result['三方']='融360'

# r360_result.to_excel(r'D:\Work\out_data\分析类\特殊名单\r360.xlsx',index=False)

#------------------------------------------------------------------------------------------------------------------------------------------------------百融
sql = """ 	select  * from juin_loan_core_prd.bairong_result_info  where process_node='CREDIT' """
bairong_result_info=pd.read_sql(sql,cnx)
bairong_result_info=pd.merge(bairong_result_info,stx,how='inner',on='user_id')

totallist=[
'br_telperiod',
'br_telstatus',
'br_link1_telstatus',
'br_link2_telstatus',
'br_sl_id_court_bad_allnum',
'br_sl_id_court_executed_allnum',
'br_sl_id_bank_bad_allnum',
'br_sl_id_bank_overdue_allnum',
'br_sl_id_bank_lost_allnum',
'br_sl_id_nbank_bad_allnum',
'br_sl_id_nbank_overdue_allnum',
'br_sl_id_nbank_lost_allnum',
'br_sl_cell_bank_bad_allnum',
'br_sl_cell_bank_overdue_allnum',
'br_sl_cell_bank_lost_allnum',
'br_sl_cell_nbank_bad_allnum',
'br_sl_cell_nbank_overdue_allnum',
'br_sl_cell_nbank_lost_allnum',
'br_bad_count',
'br_execut_count',
'br_execut_count_1y',
'br_in_execut_count',
'br_in_execut_money',
'br_frg_list_level',
'br_frg_group_num',
'br_debtrepaystress',
'br_ae_d3_id_allnum',
'br_ae_d3_id_orgnum_d',
'br_ae_d3_id_bank_allnum',
'br_ae_d3_id_nbank_allnum',
'br_ae_d3_id_bank_orgnum_d',
'br_ae_d3_id_nbank_orgnum_d',
'br_ae_d7_id_allnum',
'br_ae_d7_id_orgnum_d',
'br_ae_d7_id_bank_allnum',
'br_ae_d7_id_nbank_allnum',
'br_ae_d7_id_bank_orgnum_d',
'br_ae_d7_id_nbank_orgnum_d',
'br_ae_d15_id_allnum',
'br_ae_d15_id_orgnum_d',
'br_ae_d15_id_bank_allnum',
'br_ae_d15_id_nbank_allnum',
'br_ae_d15_id_bank_orgnum_d',
'br_ae_d15_id_nbank_orgnum_d',
'br_ae_m1_id_allnum',
'br_ae_m1_id_orgnum_d',
'br_ae_m1_id_bank_allnum',
'br_ae_m1_id_nbank_allnum',
'br_ae_m1_id_bank_orgnum_d',
'br_ae_m1_id_nbank_orgnum_d',
'br_ae_m3_id_allnum',
'br_ae_m3_id_orgnum_d',
'br_ae_m3_id_bank_allnum',
'br_ae_m3_id_nbank_allnum',
'br_ae_m3_id_bank_orgnum_d',
'br_ae_m3_id_nbank_orgnum_d',
'br_ae_m6_id_allnum',
'br_ae_m6_id_orgnum_d',
'br_ae_m6_id_bank_allnum',
'br_ae_m6_id_nbank_allnum',
'br_ae_m6_id_bank_orgnum_d',
'br_ae_m6_id_nbank_orgnum_d',
'br_ae_m12_id_allnum',
'br_ae_m12_id_orgnum_d',
'br_ae_m12_id_bank_allnum',
'br_ae_m12_id_nbank_allnum',
'br_ae_m12_id_bank_orgnum_d',
'br_ae_m12_id_nbank_orgnum_d']

col_model=[
    'br_telperiod',
'br_telstatus',
'br_link1_telstatus',
'br_link2_telstatus',
'br_frg_list_level',
'br_frg_group_num',
'br_sl_id_court_bad_allnum',
'br_sl_id_court_executed_allnum',
'br_sl_id_bank_bad_allnum',
'br_sl_id_bank_overdue_allnum',
'br_sl_id_bank_lost_allnum',
'br_sl_id_nbank_bad_allnum',
'br_sl_id_nbank_overdue_allnum',
'br_sl_id_nbank_lost_allnum',
'br_sl_cell_bank_bad_allnum',
'br_sl_cell_bank_overdue_allnum',
'br_sl_cell_bank_lost_allnum',
'br_sl_cell_nbank_bad_allnum',
'br_sl_cell_nbank_overdue_allnum',
'br_sl_cell_nbank_lost_allnum',
'br_bad_count',
'br_execut_count',
'br_execut_count_1y',
'br_in_execut_count',
'br_in_execut_money'
    ]

# Step1 好像数据库都是默认var类型，所以先剔除score类去看变量的枚举值，枚举值大的再修改为int类型去跑describe

# 参照数据统计数据
result=pd.DataFrame()
for (i,col) in enumerate(col_model):
    crossfreq=bairong_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
   
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)

# Step2 数值型，多头类
duotou_list1=[attr for attr in totallist if  attr not in col_model]
for var in duotou_list1:
    bairong_result_info[var]=bairong_result_info[var].astype(float)

def bin_variable( var_df,var_name, n_bins):#根据当前最大最小值来等距分箱
    # 计算变量的最大值和最小值
    var_min = var_df[var_name].min()
    var_max = var_df[var_name].max()
    # 根据最大值和最小值以及分箱数量计算分箱区间
    bins = np.linspace(var_min, var_max, n_bins + 1)
    bins=np.insert(bins,0, float('-inf'))
    # 对变量进行分箱，并将分箱后的新变量添加到 DataFrame 中
    var_df[var_name] = pd.cut(var_df[var_name], bins)
   
for var_name in duotou_list1:
   bin_variable( bairong_result_info,var_name, n_bins=10)

# Step2 多头
# 参照数据统计数据
result_dt=pd.DataFrame()
for (i,col) in enumerate(duotou_list1):
    crossfreq=bairong_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    
    if i==0:
        result_dt=crossfreq
    else:
        result_dt=pd.concat([result_dt,crossfreq])
        
result_dt['变量名']=result_dt.变量名.apply(lambda x:x.replace("_",""))
result_dt=pd.merge(result_dt,df_base,how='left',left_on='变量名',right_on='COL2')
result_dt.drop(columns=['变量名','COL2'],inplace=True)
result_dt.rename(columns={'COL1':'变量名'},inplace=True)

br_result=pd.concat([result,result_dt],axis=0)
br_result['三方']='百融'

# Step3 score类
# result.to_excel(r'D:\Work\out_data\分析类\特殊名单\百融.xlsx',index=False)

#-------------------------------------------------------------------------------------------------------------------------------------------------同盾
sql = """ 	select  * from juin_loan_core_prd.tongdun_result_info  where process_node='CREDIT'  """
tongdun_result_info=pd.read_sql(sql,cnx)
tongdun_result_info.drop(columns=['id','flow_id','create_time',
                                 'update_time','is_delete','failure_time',
                                 'process_node','td_apply_code','td_ol_d_time_interval_app_1hour','td_ol_ip_time_interval_app_1hour'],inplace=True)
tongdun_result_info.drop_duplicates(subset='user_id',keep='last',inplace=True)

tongdun_result_info=pd.merge(tongdun_result_info,stx,how='inner',on='user_id')

# score类
socre_list=[
'td_de_score',
'td_dt_score',
'td_final_score',  
'td_gl_score',
'td_xw_score'
]

# 语雀里部分变量名含有大写字母，但是开发的时候又统一变成小写
var_list=[
    'td_i1d_cnt_isemulator_all_all_90day',
'td_i1d_cnt_detedebugger_all_all_90day',
'td_i1d_cnt_detehook_all_all_90day',
'td_i1d_cnt_devfirstsee_all_all_90day',
'td_i1d_cnt_devtamper_all_all_90day',
'td_i1d_cnt_incognitomode_all_all_90day',
'td_i1d_cnt_shortuptime_all_all_90day',
'td_i1d_cnt_susfakingloca_all_all_90day',
'td_i1m_cnt_alimobile_all_all_90day',
'td_i1m_cnt_datamobile_all_all_90day',
'td_i1m_cnt_fakemobile_all_all_90day',
'td_i1m_cnt_signup_all_all_90day',
    'td_i_is_enterprise_self_all_taxowing_all_all',
'td_i_is_per_self_all_taxowing_all_all',
'td_i_is_per_all_cease_courtcease_all_all',
'td_i_is_per_all_cease_courtdefault_all_all',
'td_i_is_per_all_close_courtcease_all_all',
'td_i_is_per_all_close_courtdefault_all_all',
'td_i_is_per_all_exect_courtcease_all_all',
'td_i_is_per_all_exect_vag1n_courtcease_all_all',
'td_i_is_per_all_tddefault_courtdefault_all_all',
'td_i_is_per_all_tddefault_vag1n_courtdefault_all_all',
'td_i_is_per_self_all_carleasedefault_all_all',
'td_i_is_per_self_all_exc2m_riskchain_all_all',
'td_i_is_per_self_all_violtraffic_all_all',
'td_m_is_per_self_all_exc2i_riskchain_all_all',
'td_ol_d_freq_record_Loan_all_1hour',
'td_ol_d_freq_record_Loan_app_1day',
'td_ol_ip_freq_record_Loan_all_1hour',
'td_ol_ip_freq_record_Loan_app_1day'
    ]

# Step1
# 参照数据统计数据
result=pd.DataFrame()
for (i,col) in enumerate(var_list):
    crossfreq=tongdun_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result['变量名']=result.变量名.apply(lambda x:x.lower())
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)


# Step2 数值型 基本都是X天内申请次数，但最大值都在20以内
col_model=[attr for attr in list(tongdun_result_info.columns) if  attr not in socre_list]
col_model=[attr for attr in col_model if  attr not in var_list]
col_model=[attr for attr in col_model if  attr not in list(stx.columns)]

# 参照数据统计数据
for var in col_model:
    tongdun_result_info[var]=tongdun_result_info[var].apply(lambda x:x.replace('-999','-1'))
    tongdun_result_info[var]=tongdun_result_info[var].astype(float)
 
    
def bin_variable( var_df,var_name, n_bins):#根据当前最大最小值来等距分箱
    # 计算变量的最大值和最小值
    var_min = var_df[var_name].min()
    var_max = var_df[var_name].max()
    # 根据最大值和最小值以及分箱数量计算分箱区间
    bins = np.linspace(var_min, var_max, n_bins + 1)
    bins=np.insert(bins,0, float('-inf'))
    # 对变量进行分箱，并将分箱后的新变量添加到 DataFrame 中
    var_df[var_name] = pd.cut(var_df[var_name], bins)
   
for var_name in col_model:
    bin_variable(tongdun_result_info ,var_name, n_bins=10)
    
# 参照数据统计数据
result_dt=pd.DataFrame()
for (i,col) in enumerate(col_model):
    crossfreq=tongdun_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    
    if i==0:
        result_dt=crossfreq
    else:
        result_dt=pd.concat([result_dt,crossfreq])
   
# result_dt['变量名']=result_dt.变量名.apply(lambda x:x.replace("_binned",""))       
result_dt['变量名']=result_dt.变量名.apply(lambda x:x.replace("_",""))
result_dt['变量名']=result_dt.变量名.apply(lambda x:x.lower())
result_dt=pd.merge(result_dt,df_base,how='left',left_on='变量名',right_on='COL2')
result_dt.drop(columns=['变量名','COL2'],inplace=True)
result_dt.rename(columns={'COL1':'变量名'},inplace=True)


# Step3 score类
for var in socre_list:
    tongdun_result_info[var]=tongdun_result_info[var].astype(int)

bins1 = [i for i in range(400, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 121, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=[
'td_de_score',
'td_dt_score',
'td_gl_score',
'td_xw_score']

#bins1
for var in socre_list1:
    tongdun_result_info[var]=pd.cut(tongdun_result_info[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    tongdun_result_info[var]=pd.cut(tongdun_result_info[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=tongdun_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
        
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)

tongdun_result=pd.concat([result_score,result_dt],axis=0)
tongdun_result=pd.concat([tongdun_result,result],axis=0)
tongdun_result['三方']='同盾'

# tongdun_result.to_excel(r'D:\Work\out_data\分析类\特殊名单\同盾.xlsx',index=False)

#-------------------------------------------------------------------------------------------------------------------------------------------------电话邦
sql = """ 	select  user_id,type_br_dhb,type_lx1_dhb,type_lx2_dhb from juin_loan_core_prd.dianhua_result_info where process_node='CREDIT'   """
dianhua_result_info=pd.read_sql(sql,cnx)
dianhua_result_info=pd.merge(dianhua_result_info,stx,how='inner',on='user_id')

duotou_list1=[
'type_br_dhb','type_lx1_dhb','type_lx2_dhb'
    ]


# 参照数据统计数据
result_dt=pd.DataFrame()
for (i,col) in enumerate(duotou_list1):
    crossfreq=dianhua_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    
  
    if i==0:
        result_dt=crossfreq
    else:
        result_dt=pd.concat([result_dt,crossfreq])
        
result_dt.变量名.replace({'type_br_dhb':'本人号码标记',
                             'type_lx1_dhb':'联系人1号码标记',
                             'type_lx2_dhb':'联系人2号码标记'},inplace=True)

dhb_result=result_dt.copy()
dhb_result['三方']='电话邦'

#-------------------------------------------------------------------------------------------------------------------------------------------------银联智策
sql = """ 	select  * from juin_loan_core_prd.unionpay_result_info where process_node='CREDIT'  """
unionpay_result_info=pd.read_sql(sql,cnx)
unionpay_result_info.drop(columns=['id','flow_id','create_time','update_time','is_delete','failure_time','process_node','yl_yc_code','yl_od_score'],inplace=True)
unionpay_result_info=pd.merge(unionpay_result_info,stx,how='inner',on='user_id')
unionpay_result_info['yl_success_deduct_cnt_rate_6m']=unionpay_result_info.yl_success_deduct_cnt_rate_6m.apply(lambda x: -1 if x=='"null"' else float(x)*100 )

col_model=[attr for attr in list(unionpay_result_info.columns) if  attr not in list(stx.columns)]
col_model=[attr for attr in col_model if  attr not in ['user_id']]

# step1
for var in col_model:
    unionpay_result_info[var]=unionpay_result_info[var].astype(float)

# 历史最大交易力度
bins_temp = [-float('inf'), 0,0.1,20000.1,30000.1,40000.1,50000.1,60000.1,70000.1,80000.1,150000.1,200000.1, float('inf')]
unionpay_result_info['yl_trade_max_amt'] = pd.cut(unionpay_result_info['yl_trade_max_amt'], bins=bins_temp, right=False)


# 近12个月失败交易力度
bins_temp = [-float('inf'), 0,0.1,5000.1,10000.1,20000.1,30000.1,40000.1,50000.1,60000.1,70000.1,80000.1,150000.1,200000.1, float('inf')]
unionpay_result_info['yl_fail_trade_amt_12m'] = pd.cut(unionpay_result_info['yl_fail_trade_amt_12m'], bins=bins_temp, right=False)

# 近6个月失败交易频度
bins_temp = [-float('inf'), 0,0.1,5.1,10.1,20.1,30.1,40.1,50.1,60.1,100.1, float('inf')]
unionpay_result_info['yl_fail_trade_cnt_6m'] = pd.cut(unionpay_result_info['yl_fail_trade_cnt_6m'], bins=bins_temp, right=False)

# 近12个月资金不足交易频度
bins_temp = [-float('inf'), 0,0.1,5.1,10.1,20.1,30.1,40.1,50.1,60.1,100.1, float('inf')]
unionpay_result_info['yl_underfund_trade_cnt_12'] = pd.cut(unionpay_result_info['yl_underfund_trade_cnt_12'], bins=bins_temp, right=False)

# 近1个月失败交易天数
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,3.1,4.1,5.1,10.1,float('inf')]
unionpay_result_info['yl_fail_trade_days_1m'] = pd.cut(unionpay_result_info['yl_fail_trade_days_1m'], bins=bins_temp, right=False)

# 近12个月失败交易天数
bins_temp = [-float('inf'), 0,0.1,5.1,10.1,20.1,30.1,40.1,50.1,60.1,100.1, float('inf')]
unionpay_result_info['yl_fail_trade_days_12m'] = pd.cut(unionpay_result_info['yl_fail_trade_days_12m'], bins=bins_temp, right=False)

# 近6个月失败出向类的交易频度
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,3.1,4.1,5.1,10.1,20.1, float('inf')]
unionpay_result_info['yl_fail_out_cnt_6m'] = pd.cut(unionpay_result_info['yl_fail_out_cnt_6m'], bins=bins_temp, right=False)


# 近6个月成功扣款频度占比
bins_temp = [-float('inf'),-1, 0, 50, 60,70,80,90,float('inf')]
unionpay_result_info['yl_success_deduct_cnt_rate_6m'] = pd.cut(unionpay_result_info['yl_success_deduct_cnt_rate_6m'], bins=bins_temp, right=True)

# 近6个月同商户最长代扣失败天数
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,3.1,float('inf')]
unionpay_result_info['yl_mhg_fail_deduct_days_6m'] = pd.cut(unionpay_result_info['yl_mhg_fail_deduct_days_6m'], bins=bins_temp, right=False)


# 近12个月二维码失败交易频度
bins_temp = [-float('inf'), 0, 0.1,5.1,10.1,20.1,float('inf')]
unionpay_result_info['yl_code_fail_trade_cnt_12m'] = pd.cut(unionpay_result_info['yl_code_fail_trade_cnt_12m'], bins=bins_temp, right=False)

# 近3个月三方支付失败交易天数
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,3.1,4.1,5.1, float('inf')]
unionpay_result_info['yl_pay_fail_days_3m'] = pd.cut(unionpay_result_info['yl_pay_fail_days_3m'], bins=bins_temp, right=False)


# 近2周最大单日失败交易总频度
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,3.1,4.1, float('inf')]
unionpay_result_info['yl_day_fail_trade_cnt_2w'] = pd.cut(unionpay_result_info['yl_day_fail_trade_cnt_2w'], bins=bins_temp, right=False)

# 近4周代扣失败频度
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,4.1,10.1, float('inf')]
unionpay_result_info['yl_fail_deduct_cnt_4w'] = pd.cut(unionpay_result_info['yl_fail_deduct_cnt_4w'], bins=bins_temp, right=False)

# 近4周失败交易天数
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,4.1,10.1, float('inf')]
unionpay_result_info['yl_fail_trade_day_4w'] = pd.cut(unionpay_result_info['yl_fail_trade_day_4w'], bins=bins_temp, right=False)

# 近6周代扣失败天数 
bins_temp = [-float('inf'), 0, 0.1, 1.1,2.1,3.1,4.1,5.1, float('inf')]
unionpay_result_info['yl_fail_deduct_day_6w'] = pd.cut(unionpay_result_info['yl_fail_deduct_day_6w'], bins=bins_temp, right=False)

# 近6周失败交易频度
bins_temp = [-float('inf'), 0, 0.1, 5, 10, 20, float('inf')]
unionpay_result_info['yl_fail_trade_cnt_6w'] = pd.cut(unionpay_result_info['yl_fail_trade_cnt_6w'], bins=bins_temp, right=False)



result=pd.DataFrame()
for (i,col) in enumerate(col_model):
    crossfreq=unionpay_result_info.groupby(col).agg(授信申请=pd.NamedAgg(column="sx", aggfunc="count"),
                                  授信申请件均=pd.NamedAgg(column="jfCreditLimit", aggfunc="mean"),
                                  授信通过=pd.NamedAgg(column="credit_result", aggfunc="sum"),
                                  提现申请=pd.NamedAgg(column="tx", aggfunc="sum"),
                                  提现申请件均=pd.NamedAgg(column="withdrawAmount", aggfunc="mean")
                                  ).reset_index().rename(columns={col:'取值'})
    crossfreq.sort_values(by='取值',inplace=True)
    
    crossfreq['授信通过率']=crossfreq['授信通过']/crossfreq['授信申请']
    crossfreq['变量名']=col
    crossfreq.fillna(0,inplace=True)
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)


unionpay_result=result.copy()
unionpay_result['三方']='银联智策'



#---------------------------------------------------------------------------------------------------------------整理上面的数据成一个宽表给涛哥
table_list=['apply_result','tongdun_result','br_result','r360_result',
            'bj_result','td_result','wy_result','unionpay_result','dhb_result']

big_table=pd.DataFrame()
for tnmae in table_list:
    if tnmae=='apply_result':
        exec(f"big_table={tnmae}.copy()")
    else:
        exec(f"big_table=pd.concat([big_table,{tnmae}],axis=0)")

big_table=big_table[['三方','变量名','取值','授信申请','授信申请件均','授信通过','授信通过率','提现申请','提现申请件均']]


big_table.to_excel(r'D:\Work\out_data\分析类\juin客户分布\用户画像.xlsx')
# 将申请、反欺诈的存入风控库，技术已经将其他三方的表存在对应的result_info表
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate'
engine = create_engine(db_url)

# inputApplyParameter.drop(columns=['withdrawCnt0D','withdrawCnt3M','withdrawFailCnt3M'],inplace=True)#'withdrawCnt0D','withdrawCnt3M','withdrawFailCnt3M'都是缺失
big_table.to_sql(name='User_Persona', con=engine, if_exists='replace', index=False)







#%%-----------钜银贷urule入参，备份库1122名单
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 借用百融来给思汝解析user_id
sql = """ 	select  user_id,request_param,api_name from juin_cord_prd_backup_1113.third_bairong_request_info   """
third_bairong_request_info=pd.read_sql(sql,cnx)
third_bairong_request_info=third_bairong_request_info[third_bairong_request_info.api_name=='IdTwo_z']
third_bairong_request_info=third_bairong_request_info.reset_index()
# 看每个user_id命中的规则
for i in range(len(third_bairong_request_info)):
    # try:
    temp_json=json.loads(third_bairong_request_info['request_param'][i])
    temp=pd.DataFrame(temp_json['reqData'],index=[i])
    temp['user_id']=third_bairong_request_info['user_id'][i]

    if i==0:
        end=temp
    else:
        end=pd.concat([end,temp])
    # except Exception as e:
    #     continue
db_url = 'mysql+pymysql://test:kytest@192.168.20.213:3306/ky_loan'
engine = create_engine(db_url)
end.to_sql(name='user1500_md5', con=engine, if_exists='replace', index=False)


# 字段的中文解释
df_base = pd.read_excel(r"D:\Work\out_data\urule\提现专项1\df.xlsx",sheet_name=r'Sheet1')
df_base=df_base[df_base.COL2.notna()]
df_base['COL2']=df_base.COL2.apply(lambda x:x.lower())
# a1_df=pd.DataFrame({'COL':a.columns})
# a1_df['COL']=a1_df.COL.apply(lambda x:x.replace("_",""))
# a1_df=pd.merge(a1_df,df_base,how='left',left_on='COL',right_on='COL2')

# 特殊名单样本情况
sql = """ 	select  user_id2 as user_id,case when loan_status='当前逾期' then 1 else 0 end as target
from juin_risk_operate.yx_fk_test1113_v2 where type=1  """
list_sr=pd.read_sql(sql,cnx)
list_sr.user_id=list_sr.user_id.astype('int64')


#--------------------------------------------------------------------------------------------------------------------------------------------------------------微言
sql = """ 	select  * from juin_cord_prd_backup_1113.weiyan_result_info   """
weiyan_result_info=pd.read_sql(sql,cnx)
weiyan_result_info.drop(columns=['id','flow_id','create_time',
                                 'update_time','is_delete','failure_time',
                                 'process_node'],inplace=True)
weiyan_result_info=pd.merge(weiyan_result_info,list_sr,how='inner',on='user_id')
# score类
socre_list=['fxpty_wy',
'fxp_cj_hjfv4_wy',
'zzpf_lspg_v2_wy',
'qt_ndf_v1_wy',
'qt3_ndf_v1_wy',
'dtsq_score_wy',
'dzqdtgz_score_wy',
'czqdtgz_score_wy',
'fyhdtgz_score_wy',
'yhdtgz_score_wy']

# Step1 好像数据库都是默认var类型，所以先剔除score类去看变量的枚举值，枚举值大的再修改为int类型去跑describe
col_model=[attr for attr in list(weiyan_result_info.columns) if  attr not in socre_list]
var_df=weiyan_result_info[col_model]
var_df.drop(columns=['user_id','target'],inplace=True)
# 参照数据统计数据
columns_m=var_df.columns
result=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    crossfreq=pd.crosstab(weiyan_result_info[col],weiyan_result_info['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.COL1=result.apply(lambda x: '15天内总申请平台数' if x.变量名=='apply15dplatcountwy' else x.COL1,axis=1)
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)
# Step2 数值型 基本都是X天内申请次数，但最大值都在20以内
# Step3 score类
score_df=weiyan_result_info[socre_list]
score_df=pd.concat([score_df,weiyan_result_info.target],axis=1)
for var in socre_list:
    score_df[var]=score_df[var].astype(int)
    
bins1 = [i for i in range(300, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 101, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=['fxpty_wy','fxp_cj_hjfv4_wy','zzpf_lspg_v2_wy']

#bins1
for var in socre_list1:
    score_df[var]=pd.cut(score_df[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    score_df[var]=pd.cut(score_df[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=pd.crosstab(score_df[col],score_df['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    # 累计量
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    # ks
    crossdens_ks=crossfreq[['good','bad']].cumsum(axis=0)/crossfreq[['good','bad']].sum()
    crossdens_ks.rename(columns={'good':'good_累计_占比','bad':'bad_累计_占比'},inplace=True)
    crossdens_ks['diff_rate']=abs(crossdens_ks['good_累计_占比']-crossdens_ks['bad_累计_占比'])
    max_value = crossdens_ks['diff_rate'].max()
    crossdens_ks['KS'] = [1 if x == max_value else 0 for x in crossdens_ks['diff_rate']]
    crossfreq=pd.concat([crossfreq,crossdens_ks],axis=1)
    # 坏客户占比
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)
wy_result=pd.concat([result_score,result],axis=0)
wy_result.to_excel(r'D:\Work\out_data\分析类\特殊名单\微言.xlsx',index=False)


#--------------------------------------------------------------------------------------------------------------------------------------------------------------talkingdata
sql = """ 	select  user_id,tengyun108_td,tengyun239_td from juin_cord_prd_backup_1113.td_result_info   """
td_result_info=pd.read_sql(sql,cnx)
td_result_info=pd.merge(td_result_info,list_sr,how='inner',on='user_id')

socre_list=['tengyun108_td','tengyun239_td']

# Step3 score类
score_df=td_result_info[['tengyun108_td','tengyun239_td','target']]
for var in socre_list:
    score_df[var]=score_df[var].astype(float)

# a=score_df.describe()
bins1 = [i for i in range(300, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 101, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=['tengyun108_td']

#bins1
for var in socre_list1:
    score_df[var]=pd.cut(score_df[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    score_df[var]=pd.cut(score_df[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=pd.crosstab(score_df[col],score_df['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    # 累计量
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    # ks
    crossdens_ks=crossfreq[['good','bad']].cumsum(axis=0)/crossfreq[['good','bad']].sum()
    crossdens_ks.rename(columns={'good':'good_累计_占比','bad':'bad_累计_占比'},inplace=True)
    crossdens_ks['diff_rate']=abs(crossdens_ks['good_累计_占比']-crossdens_ks['bad_累计_占比'])
    max_value = crossdens_ks['diff_rate'].max()
    crossdens_ks['KS'] = [1 if x == max_value else 0 for x in crossdens_ks['diff_rate']]
    crossfreq=pd.concat([crossfreq,crossdens_ks],axis=1)
    # 坏客户占比
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)
result_score.to_excel(r'D:\Work\out_data\分析类\特殊名单\TD.xlsx',index=False)


#--------------------------------------------------------------------------------------------------------------------------------------------------------------冰鉴
sql = """ 	select * from juin_cord_prd_backup_1113.icekredit_result_info   """
icekredit_result_info=pd.read_sql(sql,cnx)
icekredit_result_info=icekredit_result_info[['user_id','hui_yan_score','qing_yun_score','huo_mou_score','xing_yu_score','hao_yue_score']]
icekredit_result_info=pd.merge(icekredit_result_info,list_sr,how='inner',on='user_id')

socre_list=['hui_yan_score',
'qing_yun_score',
'huo_mou_score',
'xing_yu_score',
'hao_yue_score']

# Step3 score类
score_df=icekredit_result_info[['hui_yan_score','qing_yun_score','huo_mou_score','xing_yu_score','hao_yue_score','target']]

# 填补缺失项：-1,转换数据类型为int或者float
for var in socre_list:
    score_df[var].fillna('-1',inplace=True)
    score_df[var]=score_df[var].astype(int)

a=score_df.describe()
bins1 = [i for i in range(300, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

#bins1
for var in socre_list:
    score_df[var]=pd.cut(score_df[var], bins=bins1, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=pd.crosstab(score_df[col],score_df['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    # 累计量
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    # ks
    crossdens_ks=crossfreq[['good','bad']].cumsum(axis=0)/crossfreq[['good','bad']].sum()
    crossdens_ks.rename(columns={'good':'good_累计_占比','bad':'bad_累计_占比'},inplace=True)
    crossdens_ks['diff_rate']=abs(crossdens_ks['good_累计_占比']-crossdens_ks['bad_累计_占比'])
    max_value = crossdens_ks['diff_rate'].max()
    crossdens_ks['KS'] = [1 if x == max_value else 0 for x in crossdens_ks['diff_rate']]
    crossfreq=pd.concat([crossfreq,crossdens_ks],axis=1)
    # 坏客户占比
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])

result_score.变量名.replace({'qing_yun_score':'青云分',
                             'hao_yue_score':'皓月分',
                             'hui_yan_score':'慧眼分',
                             'huo_mou_score':'火眸分',
                             'xing_yu_score':'星宇分'},inplace=True)

result_score.to_excel(r'D:\Work\out_data\分析类\特殊名单\冰鉴.xlsx',index=False)


#--------------------------------------------------------------------------------------------------------------------------------------------------------------融360
sql = """ 	select  * from juin_cord_prd_backup_1113.rong360_result_info   """
rong360_result_info=pd.read_sql(sql,cnx)
rong360_result_info['ssxq']=rong360_result_info.apply(lambda x: "1" if len(x.flssxq_r360)>5 else "0",axis=1)
rong360_result_info.drop(columns=['id','flow_id','create_time','update_time','is_delete',
                                  'failure_time','process_node','flsshmd_r360','flssxq_r360',
                                  'pjhf_j3m_lx1_r360',
                                  'pjhf_j3m_lx2_r360',
                                  'pjhf_j3m_dw_r360',
                                  'pjll_j3m_lx1_r360',
                                  'pjll_j3m_lx2_r360',
                                  'pjll_j3m_dw_r360',
                                  'pjthsc_j3m_lx1_r360',
                                  'pjthsc_j3m_lx2_r360',
                                  'pjthsc_j3m_dw_r360'],inplace=True)

rong360_result_info=pd.merge(rong360_result_info,list_sr,how='inner',on='user_id')

# score类
socre_list=['zy_score_r360',
'zxqs_score_r360',
'zxsjwl_score_r360',
'zxjc_r360',
'zxxdfxyz_r360',
'zxxwfx_r360',
'wddt_score_r360']

# Step1 好像数据库都是默认var类型，所以先剔除score类去看变量的枚举值，枚举值大的再修改为int类型去跑describe
col_model=[attr for attr in list(rong360_result_info.columns) if  attr not in socre_list]
var_df=rong360_result_info[col_model]
var_df.drop(columns=['user_id','target'],inplace=True)

# 参照数据统计数据
columns_m=var_df.columns
result=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    crossfreq=pd.crosstab(rong360_result_info[col],rong360_result_info['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)

    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))

result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.COL1=result.apply(lambda x: '涉诉有内容' if x.变量名=='ssxq' else x.COL1,axis=1)
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)
# Step2 数值型 基本都是X天内申请次数，但最大值都在20以内

# Step3 score类
score_df=rong360_result_info[socre_list]
score_df=pd.concat([score_df,rong360_result_info.target],axis=1)

for var in socre_list:
    score_df[var].fillna("-1",inplace=True)
    score_df[var]=score_df[var].astype(int)

bins1 = [i for i in range(300, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 101, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=[
'zxqs_score_r360',
'zxsjwl_score_r360',
'zxjc_r360',
'zxxdfxyz_r360',
'zxxwfx_r360']

#bins1
for var in socre_list1:
    score_df[var]=pd.cut(score_df[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    score_df[var]=pd.cut(score_df[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=pd.crosstab(score_df[col],score_df['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    # 累计量
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    # ks
    crossdens_ks=crossfreq[['good','bad']].cumsum(axis=0)/crossfreq[['good','bad']].sum()
    crossdens_ks.rename(columns={'good':'good_累计_占比','bad':'bad_累计_占比'},inplace=True)
    crossdens_ks['diff_rate']=abs(crossdens_ks['good_累计_占比']-crossdens_ks['bad_累计_占比'])
    max_value = crossdens_ks['diff_rate'].max()
    crossdens_ks['KS'] = [1 if x == max_value else 0 for x in crossdens_ks['diff_rate']]
    crossfreq=pd.concat([crossfreq,crossdens_ks],axis=1)
    # 坏客户占比
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)
wy_result=pd.concat([result_score,result],axis=0)
wy_result.to_excel(r'D:\Work\out_data\分析类\特殊名单\r360.xlsx',index=False)



#--------------------------------------------------------------------------------------------------------------------------------------------------------------百融
sql = """ 	select  * from juin_cord_prd_backup_1113.bairong_result_info   """
bairong_result_info=pd.read_sql(sql,cnx)
duotou_list1=['user_id',
    'br_debtrepaystress',
    'br_ae_d3_id_allnum',
    'br_ae_d3_id_orgnum_d',
    'br_ae_d3_id_bank_allnum',
    'br_ae_d3_id_nbank_allnum',
    'br_ae_d3_id_bank_orgnum_d',
    'br_ae_d3_id_nbank_orgnum_d',
    'br_ae_d7_id_allnum',
    'br_ae_d7_id_orgnum_d',
    'br_ae_d7_id_bank_allnum',
    'br_ae_d7_id_nbank_allnum',
    'br_ae_d7_id_bank_orgnum_d',
    'br_ae_d7_id_nbank_orgnum_d',
    'br_ae_d15_id_allnum',
    'br_ae_d15_id_orgnum_d',
    'br_ae_d15_id_bank_allnum',
    'br_ae_d15_id_nbank_allnum',
    'br_ae_d15_id_bank_orgnum_d',
    'br_ae_d15_id_nbank_orgnum_d',
    'br_ae_m1_id_allnum',
    'br_ae_m1_id_orgnum_d',
    'br_ae_m1_id_bank_allnum',
    'br_ae_m1_id_nbank_allnum',
    'br_ae_m1_id_bank_orgnum_d',
    'br_ae_m1_id_nbank_orgnum_d',
    'br_ae_m3_id_allnum',
    'br_ae_m3_id_orgnum_d',
    'br_ae_m3_id_bank_allnum',
    'br_ae_m3_id_nbank_allnum',
    'br_ae_m3_id_bank_orgnum_d',
    'br_ae_m3_id_nbank_orgnum_d',
    'br_ae_m6_id_allnum',
    'br_ae_m6_id_orgnum_d',
    'br_ae_m6_id_bank_allnum',
    'br_ae_m6_id_nbank_allnum',
    'br_ae_m6_id_bank_orgnum_d',
    'br_ae_m6_id_nbank_orgnum_d',
    'br_ae_m12_id_allnum',
    'br_ae_m12_id_orgnum_d',
    'br_ae_m12_id_bank_allnum',
    'br_ae_m12_id_nbank_allnum',
    'br_ae_m12_id_bank_orgnum_d',
    'br_ae_m12_id_nbank_orgnum_d']

var_df=bairong_result_info[duotou_list1]
var_df=pd.merge(var_df,list_sr,how='inner',on='user_id')

for var in [attr for attr in duotou_list1 if attr not in ['user_id'] ]:
    # var_df[var].fillna("-1",inplace=True)
    var_df[var]=var_df[var].astype(int)

def bin_variable( var_name, n_bins):#根据当前最大最小值来等距分箱
    # 计算变量的最大值和最小值
    var_min = var_df[var_name].min()
    var_max = var_df[var_name].max()
    # 根据最大值和最小值以及分箱数量计算分箱区间
    bins = np.linspace(var_min, var_max, n_bins + 1)
    bins=np.insert(bins,0, float('-inf'))
    # 对变量进行分箱，并将分箱后的新变量添加到 DataFrame 中
    var_df[f'{var_name}_binned'] = pd.cut(var_df[var_name], bins)
   
for var_name in [attr for attr in duotou_list1 if attr not in ['user_id'] ]:
    bin_variable( var_name, n_bins=10)
col_binned=[attr for attr in list(var_df.columns) if  attr.endswith('_binned')]
a=var_df.columns
# Step1 好像数据库都是默认var类型，所以先剔除score类去看变量的枚举值，枚举值大的再修改为int类型去跑describe

# 参照数据统计数据
result=pd.DataFrame()
for (i,col) in enumerate(col_binned):
    crossfreq=pd.crosstab(var_df[col],var_df['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.apply(lambda x:x.replace("_binned",""))       
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)
# Step2 数值型 基本都是X天内申请次数，但最大值都在20以内

# Step3 score类
result.to_excel(r'D:\Work\out_data\分析类\特殊名单\百融.xlsx',index=False)



#--------------------------------------------------------------------------------------------------------------------------------------------------------------同盾
sql = """ 	select  user_id,
td_de_score,
td_dt_score,
td_final_score  ,
td_gl_score,
td_xw_score,
td_i_cnt_deviceid_v3_all_all_365day,
td_i_cnt_mobile_v3_all_all_365day,
td_i_cnt_node_dist1_card_Loan_all_all,
td_i_cnt_node_dist1_mobile_Loan_all_all,
td_i_incr_set_recent90daypartner_v3_Loan_all_270day,
td_i_max_cnt_partner_daily_v3_Loan_Imbank_365day,
td_i_std_cnt_30daypartner_v3_Loan_Imbank_180day,
td_i_wcInterestLevel_partner_Loan_all_180day,
td_i2m_cnt_partner_v3_Loan_all_180day,
td_i2m_cnt_partner_v3_Loan_all_30day,
td_i2m_cnt_partner_v3_Loan_all_365day,
td_i2m_cnt_partner_v3_Loan_all_7day,
td_i2m_cnt_partner_v3_Loan_all_90day,
td_i2m_cnt_partner_v3_Loan_Imbank_180day,
td_i2m_cnt_partner_v3_Loan_Imbank_30day,
td_i2m_cnt_partner_v3_Loan_Imbank_365day,
td_i2m_cnt_partner_v3_Loan_Imbank_7day,
td_i2m_cnt_partner_v3_Loan_Imbank_90day,
td_m_max_cnt_partner_daily_v3_Loan_Imbank_365day 
from juin_cord_prd_backup_1113.tongdun_result_info   """
tongdun_result_info=pd.read_sql(sql,cnx)
tongdun_result_info.drop_duplicates(subset='user_id',keep='last',inplace=True)

tongdun_result_info=pd.merge(tongdun_result_info,list_sr,how='inner',on='user_id')


# score类
socre_list=[
'td_de_score',
'td_dt_score',
'td_final_score',  
'td_gl_score',
'td_xw_score'
]

# Step1 好像数据库都是默认var类型，所以先剔除score类去看变量的枚举值，枚举值大的再修改为int类型去跑describe
col_model=[attr for attr in list(tongdun_result_info.columns) if  attr not in socre_list]
var_df=tongdun_result_info[col_model]
# 参照数据统计数据
for var in [attr for attr in list(var_df.columns) if attr not in ['user_id','target'] ]:
    var_df[var]=var_df[var].astype(float)

def bin_variable( var_name, n_bins):#根据当前最大最小值来等距分箱
    # 计算变量的最大值和最小值
    var_min = var_df[var_name].min()
    var_max = var_df[var_name].max()
    # 根据最大值和最小值以及分箱数量计算分箱区间
    bins = np.linspace(var_min, var_max, n_bins + 1)
    bins=np.insert(bins,0, float('-inf'))
    # 对变量进行分箱，并将分箱后的新变量添加到 DataFrame 中
    var_df[f'{var_name}_binned'] = pd.cut(var_df[var_name], bins)
   
list_not_bin=['user_id','target','td_i_cnt_node_dist1_card_Loan_all_all','td_i_cnt_node_dist1_mobile_Loan_all_all','td_m_max_cnt_partner_daily_v3_Loan_Imbank_365day']
for var_name in [attr for attr in list(var_df.columns) if attr not in list_not_bin ]:
    bin_variable( var_name, n_bins=10)
    
col_binned=[attr for attr in list(var_df.columns) if  attr.endswith('_binned')]
col_binned.append('td_i_cnt_node_dist1_card_Loan_all_all')
col_binned.append('td_i_cnt_node_dist1_mobile_Loan_all_all')
col_binned.append('td_m_max_cnt_partner_daily_v3_Loan_Imbank_365day')

result=pd.DataFrame()
for (i,col) in enumerate(col_binned):
    crossfreq=pd.crosstab(var_df[col],var_df['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result=crossfreq
    else:
        result=pd.concat([result,crossfreq])
        
result['变量名']=result.变量名.apply(lambda x:x.replace("_binned",""))       
result['变量名']=result.变量名.apply(lambda x:x.replace("_",""))
result['变量名']=result.变量名.apply(lambda x:x.lower())
result=pd.merge(result,df_base,how='left',left_on='变量名',right_on='COL2')
result.drop(columns=['变量名','COL2'],inplace=True)
result.rename(columns={'COL1':'变量名'},inplace=True)
# Step2 数值型 基本都是X天内申请次数，但最大值都在20以内

# Step3 score类
score_df=tongdun_result_info[socre_list]
score_df=pd.concat([score_df,tongdun_result_info.target],axis=1)

for var in socre_list:
    score_df[var]=score_df[var].astype(int)

bins1 = [i for i in range(400, 900, 50)]
bins1[-1] = bins1[-1]+1
bins1.insert(0, float('-inf'))

bins2 = [i for i in range(0, 121, 10)]
bins2[-1] = bins2[-1]+1
bins2.insert(0, float('-inf'))

socre_list1=[
'td_de_score',
'td_dt_score',
'td_gl_score',
'td_xw_score']

#bins1
for var in socre_list1:
    score_df[var]=pd.cut(score_df[var], bins=bins1, include_lowest=True,right=False)
#bins2
for var in [attr for attr in socre_list if  attr not in socre_list1]:
    score_df[var]=pd.cut(score_df[var], bins=bins2, include_lowest=True,right=False)

result_score=pd.DataFrame()
for (i,col) in enumerate(socre_list):
    crossfreq=pd.crosstab(score_df[col],score_df['target']).reset_index()
    crossfreq.rename(columns={0:'good',1:'bad',col:'取值'},inplace=True)
    crossfreq.sort_values(by='取值',inplace=True)
    # 累计量
    crossdens=crossfreq[['good','bad']].cumsum(axis=0).rename(columns={'good':'good_累计','bad':'bad_累计'})
    crossfreq=pd.concat([crossfreq,crossdens],axis=1)
    # ks
    crossdens_ks=crossfreq[['good','bad']].cumsum(axis=0)/crossfreq[['good','bad']].sum()
    crossdens_ks.rename(columns={'good':'good_累计_占比','bad':'bad_累计_占比'},inplace=True)
    crossdens_ks['diff_rate']=abs(crossdens_ks['good_累计_占比']-crossdens_ks['bad_累计_占比'])
    max_value = crossdens_ks['diff_rate'].max()
    crossdens_ks['KS'] = [1 if x == max_value else 0 for x in crossdens_ks['diff_rate']]
    crossfreq=pd.concat([crossfreq,crossdens_ks],axis=1)
    # 坏客户占比
    crossfreq['bad_rate']=crossfreq['bad']/(crossfreq['good']+crossfreq['bad'])
    crossfreq['bad_累计_rate']=crossfreq['bad_累计']/(crossfreq['good_累计']+crossfreq['bad_累计'])
    crossfreq['变量名']=col
    if i==0:
        result_score=crossfreq
    else:
        result_score=pd.concat([result_score,crossfreq])
result_score['变量名']=result_score.变量名.apply(lambda x:x.replace("_",""))
result_score=pd.merge(result_score,df_base,how='left',left_on='变量名',right_on='COL2')
result_score.drop(columns=['变量名','COL2'],inplace=True)
result_score.rename(columns={'COL1':'变量名'},inplace=True)
tongdun_result=pd.concat([result_score,result],axis=0)
tongdun_result.to_excel(r'D:\Work\out_data\分析类\特殊名单\同盾.xlsx',index=False)


#------------------------------------------------------------------------------------------------------------------------------------------整理上面的数据成一个宽表给涛哥


table_list=['weiyan_result_info','icekredit_result_info','td_result_info',
            'rong360_result_info','var_df','tongdun_result_info']

for tnmae in table_list:
    if tnmae=='weiyan_result_info':
        exec(f"big_table={tnmae}.copy()")
    else:
        exec(f"big_table=pd.merge(big_table,{tnmae},how='inner',on='user_id')")

big_table=pd.merge(big_table,list_sr,how='inner',on='user_id')

big_table_COLname=pd.DataFrame({"COLname":big_table.columns})
big_table_COLname['COLname']=big_table_COLname.COLname.apply(lambda x:x.replace("_",""))
big_table_COLname['COLname']=big_table_COLname.COLname.apply(lambda x:x.lower())
big_table_COLname=pd.merge(big_table_COLname,df_base,how='left',left_on='COLname',right_on='COL2')

def rename_kb(x):
    if x.COLname=='userid':
        return 'userid'
    elif x.COLname=='apply15dplatcountwy':
        return '15天内总申请平台数'
    elif x.COLname=='huiyanscore':
        return '慧眼分'
    elif x.COLname=='qingyunscore':
        return '青云分'
    elif x.COLname=='huomouscore':
        return '火眸分'
    elif x.COLname=='xingyuscore':
        return '星宇分'
    elif x.COLname=='haoyuescore':
        return '皓月分'
    elif x.COLname=='ssxq':
        return '有涉诉记录'
    elif x.COLname=='target':
        return '逾期客户'
    else :
        return x.COL1

big_table_COLname.COL1=big_table_COLname.apply(rename_kb,axis=1)
# d= dict(zip(big_table_COLname['COLname'],big_table_COLname['COL1']))

newcolname=list(big_table_COLname.COL1)
big_table.columns=newcolname




specific_string = 'NA'
for col in big_table.columns:
    if big_table[col].astype(str).str.contains(specific_string).any():
        print(col)
nalist=[
      '近三个月平均话费_本人',
      '近三个月平均流量_本人',
      '近三个月平均通话时长_本人',
      '风险范围',
      '负面记录类型'  
        ]

for var in [attr for attr in list(big_table.columns) if attr not in nalist ]:
    big_table[var]=big_table[var].astype(float)
    
big_table.to_excel(r'D:\Work\out_data\分析类\特殊名单\宽表.xlsx',index=False)

from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text
db_url = 'mysql+pymysql://test:kytest@192.168.20.213:3306/juin_risk_operate'
engine = create_engine(db_url)
big_table.to_sql(name='table1122', con=engine, if_exists='replace', index=False)



#%%----------单变量分布然后concat
columns_m=var_df.columns
result=pd.DataFrame()
for (i,col) in enumerate(columns_m):
    temp = var_df[col].value_counts().to_frame('计数').reset_index().rename(columns={'index':'取值'})
    temp.sort_values(by='取值',inplace=True)
    temp['ColName']=col
    temp['计数占比']=temp['计数']/var_df.shape[0]
    if i==0:
        result=temp
    else:
        result=pd.concat([result,temp])

#%%-----------（生产）钜银贷urule入参，出参解析（决策模拟、验证、优化）
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ 	 select a.user_id,a.flow_id, a.request_body,c.channel_source from 
                    (
                     select user_id,flow_id,request_body,
                     ROW_NUMBER() over (PARTITION by flow_id order by create_time desc,update_time  desc)  as countid
                     FROM juin_loan_core_prd.risk_request_apply 
                     WHERE process_node IN (0,4) 
                     ) as a
										 left join juin_loan_core_prd.user_info as b on a.user_id=b.id
left join juin_loan_core_prd.white_list as c on b.id_number_md5=c.id_number
where   a.countid=1   """   #注意客户授信失败过了冻结期后的再授信问题，还有客户过了缓存期的提现调用三方问题
risk_request_apply=pd.read_sql(sql,cnx)


sql = """ 	select  规则名称,规则码 from juin_risk_operate.reject_code     """
reject_code_pzb=pd.read_sql(sql,cnx)


# risk_request_apply=risk_request_apply[risk_request_apply['request_body'].str.contains('inputAntiFraudParameter', case=False, na=False)].reset_index()
# risk_request_apply=risk_request_apply[risk_request_apply.index!=293].reset_index()

table_list1=[
'inputAntiFraudParameter',
'inputApplyParameter',
'inputThirdParameter']

table_list2=[
'inputBaiRongParameter',
'inputDianHuaParameter',
'inputIcekreditParameter',
'inputOcrParameter',
'inputRong360Parameter',
'inputTdParameter',
'inputTongDunParameter',
'inputUnionPayParameter',
'inputWeiYanParameter']


# 解析urule入参第一代
def for_jiexi(risk_request_apply,table_list1,table_list2):
    for i in range(len(risk_request_apply)):
        print(i)
        try:
            temp_json=json.loads(risk_request_apply['request_body'][i])
            for table_name1 in table_list1:
                if table_name1!='inputThirdParameter':
                    exec(f"{table_name1+'_temp'} = pd.DataFrame(temp_json[table_name1],index=[i])")
                    exec(f"{table_name1+'_temp'}['user_id'] = risk_request_apply['user_id'][i]")
                    exec(f"{table_name1+'_temp'}['flow_id'] = risk_request_apply['flow_id'][i]")
                    exec(f"{table_name1+'_temp'}['channel_id'] = risk_request_apply['channel_id'][i]")
             
                    if i==0:
                        exec(f"{table_name1} = {table_name1+'_temp'}")
                        
                    else:
                        exec(f"{table_name1} = pd.concat([{table_name1},{table_name1+'_temp'}])")
                else:
                    for table_name2 in table_list2:
                        exec(f"{table_name2+'_temp'} =pd.DataFrame(temp_json['inputThirdParameter'][table_name2],index=[i])")
                        exec(f"{table_name2+'_temp'}['user_id'] = risk_request_apply['user_id'][i]")
                        exec(f"{table_name2+'_temp'}['flow_id'] = risk_request_apply['flow_id'][i]")
                        exec(f"{table_name2+'_temp'}['channel_id'] = risk_request_apply['channel_id'][i]")
                     
                        if i==0:
                            exec(f"{table_name2} = {table_name2+'_temp'}")
                        else:
                            exec(f"{table_name2} = pd.concat([{table_name2},{table_name2+'_temp'}])")
        except:
            continue
    return tables

# 解析urule入参第二代
tables = {}
def for_jiexi(risk_request_apply):
    table_list1=[
    'inputAntiFraudParameter',
    'inputApplyParameter',
    'inputThirdParameter']

    table_list2=[
    'inputBaiRongParameter',
    'inputDianHuaParameter',
    'inputIcekreditParameter',
    'inputOcrParameter',
    'inputRong360Parameter',
    'inputTdParameter',
    'inputTongDunParameter',
    'inputUnionPayParameter',
    'inputWeiYanParameter']


    for i in range(len(risk_request_apply)):
        print(i)
        try:
            temp_json = json.loads(risk_request_apply['request_body'][i])

            for table_list, table_prefix in [(table_list1, ''), (table_list2, 'n')]:
                for table_name in table_list:
                    if table_name in table_list1:
                        table_data = pd.DataFrame(temp_json[table_name], index=[i])
                    else:
                        if table_name in table_list2:
                            table_data = pd.DataFrame(temp_json['inputThirdParameter'][table_name], index=[i])

                    table_data['user_id'] = risk_request_apply['user_id'][i]
                    table_data['flow_id'] = risk_request_apply['flow_id'][i]
               

                    if table_name not in tables:
                        tables[table_name] = table_data
                    else:
                        tables[table_name] = pd.concat([tables[table_name], table_data])
        except Exception as e:
            print(f"An error occurred in iteration {i}: {e}")
            continue

    return tables
for_jiexi(risk_request_apply)


   
# 拼入参宽表
drop_list=table_list1+table_list2
drop_list.remove('inputThirdParameter')

for lis in drop_list:
    exec(f"{lis}.drop_duplicates(subset='flow_id',keep='first',inplace=True)")

table_list=table_list1+table_list2
table_list.remove('inputThirdParameter')

for lis in table_list:
    if lis=="inputAntiFraudParameter":
        big_table=inputAntiFraudParameter.copy()
    else:
        exec(f"big_table=pd.merge(big_table,{lis},how='inner',on=['user_id','flow_id','channel_id'])")
                                         

var_numeric=[
'tengyun108Td',#腾云108分
'tengyun239Td',#腾云239分   
'huiyan23Bj',#慧眼分 
'huomou22Bj',#火眸分 
'qingyun22Bj',#青云分 
'xingyu22Bj',#星宇分 
'tdI2mCntPartnerV3LoanImbank7day',#同盾近7天
'tdI2mCntPartnerV3LoanImbank30day',#同盾近1月
'tdI2mCntPartnerV3LoanImbank90day',#同盾近3月
'tdI2mCntPartnerV3LoanAll90day',#同盾近3月
'tdIIsPerAllExectCourtceaseAllAll',#身份证命中法院执行(案件状态为执行中)
'tdIIsPerAllExectVag1nCourtceaseAllAll',#身份证&姓名命中法院执行模糊名单(案件状态为执行中) 
'tdIIsPerAllTddefaultCourtdefaultAllAll',#身份证命中法院失信(案件状态为失信) 
'tdIIsPerAllTddefaultVag1nCourtdefaultAllAll',#身份证&姓名命中法院失信模糊名单(案件状态为失信) 
'tdIIsPerSelfAllCarleasedefaultAllAll',#身份证命中汽车租赁违
'tdICntDeviceidV3AllAll365day',#同盾身份证号关联设备数
'tdICntMobileV3AllAll365day',#同盾身份证号关联手机号数
'fxptyWy',#微言风险谱通用分
'zzpfLspgV2Wy',#微言资质评分流⽔评估
'fxpCjHjfv4Wy',#微言信用贷类评分
'overdue1mCountWy',#微言1个月内申请人逾期次数
'overdue1mPlatCountWy', #1个月内申请人逾期平台数
'apply7dPlatCountWy', #微言7天内总申请平台数
'apply30dPlatCountWy', #微言30天内总申请平台数
'brAeD7IdNbankOrgnumD' ,#百融近7天多头
'brAeM1IdNbankOrgnumD' ,#百融近1月非银申请机构数
'brAeM3IdNbankOrgnumD' ,#百融近3个月极高
'brAeM3IdOrgnumD' ,#近3个月极低多头
'brIdIdentity' ,#二要素认证
'brTelIdentity' ,#手机实名认证
'brTelstatus' ,#本人手机在网状态
'brLink1Telstatus' ,#联系人1手机在网状态
'brLink2Telstatus' ,#联系人2手机在网状态
'brTelperiod',#手机在网时长
'brSlIdCourtBadAllnum',#id法院失信人
'brSlIdCourtExecutedAllnum',#id法院被执行人
'brSlIdBankBadAllnum',#id银行(含信用卡)中风险
'brSlIdBankLostAllnum',#id银行(含信用卡)高风险
'brSlIdBankOverdueAllnum',#id银行(含信用卡)一般风险
'brSlIdNbankBadAllnum' ,#id非银(含全部非银类型)中风险
'brSlIdNbankLostAllnum' ,#id非银(含全部非银类型)高风险
'brSlIdNbankOverdueAllnum' ,#id非银(含全部非银类型)一般风险
'brSlCellBankBadAllnum',#cell银行(含信用卡)中风险
'brSlCellBankLostAllnum',#cell银行(含信用卡)高风险
'brSlCellBankOverdueAllnum',#cell银行(含信用卡)一般风险
'brSlCellNbankBadAllnum',#cell非银(含全部非银类型)中风险
'brSlCellNbankLostAllnum',#cell非银(含全部非银类型)高风险
'brSlCellNbankOverdueAllnum',#cell非银(含全部非银类型)一般风险
'brBadCount',#百融命中失信人
'brExecutCount1y',#百融命中被执行人
'hmddjR360',#融360有90+以上逾期或欺诈
'jynzyzhmddjR360',#融360近一年有30+逾期
'zxxdfxyzR360',#占信分极低
'zyScoreR360' ,#占御分极高
'ylFailTradeCnt6w',#银联近6周失败交易次数极多
'ylLoanMhgFailDeductMaxDay4w',#银联近4周互联网小贷商户上最长代扣失败天数过长
'ylFailDeductCnt4w' ,#银联近4周代扣失败次数过多
'ylDayFailTradeCnt2w' #银联近2周最大单日失败交易次数过多
    ]

for var in var_numeric:
    big_table[var]=big_table[var].astype(float)

big_table.typeR360.fillna('NA',inplace=True)

def func(x):
    reject_code=[]
    
    if x.tengyun108Td >=0 and x.tengyun108Td <340 :#腾云108分低
        reject_code.append('D-TD2-001-30')
    if x.tengyun239Td >95 :#腾云239分低
        reject_code.append('D-TD2-002-30')
    
    if x.huiyan23Bj >=0 and x.huiyan23Bj <440 :#慧眼分低
        reject_code.append('D-BJ-001-30')
    if x.huomou22Bj >=0 and x.huomou22Bj <440 :#火眸分低
        reject_code.append('D-BJ-002-30')
    if x.qingyun22Bj >=0 and x.qingyun22Bj <440 :#青云分低
        reject_code.append('D-BJ-003-30')
    if x.xingyu22Bj >=0 and x.xingyu22Bj <440 :#星宇分低
        reject_code.append('D-BJ-004-30')
    if x.xingyu22Bj >=0 and x.xingyu22Bj <490  and  \
       x.qingyun22Bj >=0 and x.qingyun22Bj <490 and  \
       x.huomou22Bj >=0 and x.huomou22Bj <490   and \
       x.huiyan23Bj >=0 and x.huiyan23Bj <490 :#冰鉴多评分分低
        reject_code.append('D-BJ-005-30')
    
    if x.tdI2mCntPartnerV3LoanImbank7day >25 :#同盾近7天极高多头
        reject_code.append('D-TD1-001-30')
    if x.tdI2mCntPartnerV3LoanImbank30day >35 :#同盾近1月多头极高
        reject_code.append('D-TD1-002-30')
    if x.tdI2mCntPartnerV3LoanImbank90day >55 :#同盾近3个月极高多头
        reject_code.append('D-TD1-003-30')
    if x.tdI2mCntPartnerV3LoanAll90day ==0 :#同盾近3个月极低多头
        reject_code.append('D-TD1-004-30')
    if x.tdI2mCntPartnerV3LoanImbank90day> 35 and \
       x.qingyun22Bj >=0 and x.qingyun22Bj <470 :#同盾近3个月极高多头
        reject_code.append('D-TD1-005-30')
    if x.tdIIsPerAllExectCourtceaseAllAll ==1 :#身份证命中法院执行(案件状态为执行中) 
        reject_code.append('F-TD1-001-30')
    if x.tdIIsPerAllExectVag1nCourtceaseAllAll ==1 :#身份证&姓名命中法院执行模糊名单(案件状态为执行中) 
        reject_code.append('F-TD1-002-30')
    if x.tdIIsPerAllTddefaultCourtdefaultAllAll ==1 :#身份证命中法院失信(案件状态为失信) 
        reject_code.append('F-TD1-003-30')
    if x.tdIIsPerAllTddefaultVag1nCourtdefaultAllAll ==1 :#身份证&姓名命中法院失信模糊名单(案件状态为失信)  
        reject_code.append('F-TD1-004-30')
    if x.tdIIsPerSelfAllCarleasedefaultAllAll ==1 :#身份证命中汽车租赁违
        reject_code.append('F-TD1-005-30')
    if x.tdICntDeviceidV3AllAll365day >15 :#同盾身份证号关联设备数过多
        reject_code.append('F-TD1-006-30')
    if x.tdICntMobileV3AllAll365day >10 :#同盾身份证号关联手机号数过多
        reject_code.append('F-TD1-007-30')
    
    if x.fxptyWy >=0 and x.fxptyWy <390 :#微言风险谱通用分极低
        reject_code.append('D-WY-001-30')
    if x.zzpfLspgV2Wy >=0 and x.zzpfLspgV2Wy <340 :#微言资质评分流⽔评估极低
        reject_code.append('D-WY-002-30')
    if x.zzpfLspgV2Wy >=0 and x.zzpfLspgV2Wy <440 and \
       x.fxptyWy >=0 and x.fxptyWy <490 and \
       x.fxpCjHjfv4Wy >=0 and x.fxpCjHjfv4Wy <710 :#微言多维度评分低
        reject_code.append('D-WY-003-30')
    if x.overdue1mCountWy >=5 :#微言1个月内申请人逾期次数较多
        reject_code.append('D-WY-004-30')
    if x.overdue1mPlatCountWy >=3 :#1个月内申请人逾期平台数较多
        reject_code.append('D-WY-005-30')
    if x.apply7dPlatCountWy >25 :#微言7天内总申请平台数较多
        reject_code.append('D-WY-006-30')
    if x.apply30dPlatCountWy >35 :#微言30天内总申请平台数较多
        reject_code.append('D-WY-007-30')
    
    if x.brAeD7IdNbankOrgnumD >25 :#百融近7天多头极高
        reject_code.append('D-BR-001-30')
    if x.brAeM1IdNbankOrgnumD >35 :#百融近1月非银申请机构数极高
        reject_code.append('D-BR-002-30')
    if x.brAeM3IdNbankOrgnumD >55 :#百融近3个月极高多头
        reject_code.append('D-BR-003-30')
    if x.brAeM3IdOrgnumD ==0 :#近3个月极低多头
        reject_code.append('D-BR-004-30')
    if x.brAeM3IdOrgnumD >=0 and x.brAeM3IdOrgnumD <=3 and \
       x.tdI2mCntPartnerV3LoanAll90day >=0 and x.tdI2mCntPartnerV3LoanAll90day <=3 :#近3个月多平台极低多头
        reject_code.append('D-BR-005-30')
    if x.brAeM3IdNbankOrgnumD >35 and  \
       x.qingyun22Bj >=0 and x.qingyun22Bj <470 :#近3个月高多头且青云分低
        reject_code.append('D-BR-006-30')
    if x.brAeM3IdNbankOrgnumD >0 and  x.brAeM1IdNbankOrgnumD>18 and \
       x.brAeM1IdNbankOrgnumD/x.brAeM3IdNbankOrgnumD > 0.75:#近期频繁借贷
        reject_code.append('D-BR-007-30')
    if x.brIdIdentity ==2 :#二要素认证不一致
        reject_code.append('ZR-BR-001-30')
    if x.brTelIdentity ==2 :#手机实名认证不一致
        reject_code.append('ZR-BR-002-30')
    if x.brTelstatus in [2,3,4]:#本人手机在网状态为停机销号异常
        reject_code.append('ZR-BR-003-30')
    if x.brLink1Telstatus in [2,3,4]:#联系人1手机在网状态为停机销号异常
        reject_code.append('ZR-BR-004-30')
    if x.brLink2Telstatus in [2,3,4]:#联系人2手机在网状态为停机销号异常
        reject_code.append('ZR-BR-005-30')
    if x.brTelperiod ==1:#手机在网时长较短
        reject_code.append('ZR-BR-006-30')
    if x.brSlIdCourtBadAllnum >0 or \
       x.brSlIdCourtExecutedAllnum >0 or \
       x.brSlIdBankBadAllnum >0 or \
       x.brSlIdBankLostAllnum >0 or \
       x.brSlIdBankOverdueAllnum >0 or \
       x.brSlIdNbankBadAllnum >0 or \
       x.brSlIdNbankLostAllnum >0 or \
       x.brSlIdNbankOverdueAllnum >2 or \
       (x.brSlIdNbankOverdueAllnum in [1,2] and x.huiyan23Bj >0 and x.huiyan23Bj <480):#百融身份证特殊名单拦截
        reject_code.append('F-BR-001-30')
    if x.brSlCellBankBadAllnum >0 or \
       x.brSlCellBankLostAllnum >0 or \
       x.brSlCellBankOverdueAllnum >0 or \
       x.brSlCellNbankBadAllnum >0 or \
       x.brSlCellNbankLostAllnum >0 or \
       x.brSlCellNbankOverdueAllnum >2 or \
       (x.brSlCellNbankOverdueAllnum in [1,2] and x.huiyan23Bj >0 and x.huiyan23Bj <480):#百融手机号特殊名单拦截
        reject_code.append('F-BR-002-30')
    if x.brBadCount >0:#百融命中失信人
        reject_code.append('F-BR-003-30')
    if x.brExecutCount1y >0:#百融命中被执行人
        reject_code.append('F-BR-004-30')
    
    if x.hmddjR360 in [2,3]:#融360有90+以上逾期或欺诈
        reject_code.append('D-R360-001-30')
    if x.hmddjR360 ==1 and x.jynzyzhmddjR360 >1:#融360近一年有30+逾期
        reject_code.append('D-R360-002-30')
    if x.zxxdfxyzR360 >=0 and x.zxxdfxyzR360 <=390:#占信分极低
        reject_code.append('D-R360-003-30')
    if x.zyScoreR360 >95 :#占御分极高
        reject_code.append('D-R360-004-30')
    if x.typeR360 !='NA' :#融360特殊名单
        reject_code.append('F-R360-001-9999')
    
    if x.ylFailTradeCnt6w >=30 :#银联近6周失败交易次数极多
        reject_code.append('D-YL-001-30')
    if x.ylLoanMhgFailDeductMaxDay4w >=5 :#银联近4周互联网小贷商户上最长代扣失败天数过长
        reject_code.append('D-YL-002-30')
    if x.ylFailDeductCnt4w >=20 :#银联近4周代扣失败次数过多
        reject_code.append('D-YL-003-30')
    if x.ylDayFailTradeCnt2w >=10 :#银联近2周最大单日失败交易次数过多
        reject_code.append('D-YL-004-30')
        
    return ', '.join(reject_code)


# 照着urule决策流写规则,想实现的效果是2个函数：每个覆盖所有if规则，apply后生成拒绝码的list格式。区别在于1个是实际规则函数，1个是参考规则函数，（可以部署psi了)
# 1.0 单变量规则：最后通过循环迭代来修改其中的某条规则阈值看两个函数的拒绝效果差异、贷后表现差异。差异达到阈值则提炼出来。
# 2.0 双变量规则检视效果
# x.0 以此类推

# 至此，已成艺术，这里的拒绝码抽查基本没问题，但因为实际生产有反欺诈、申请、腾讯、观测码，没办法做一致性校验
big_table['reject_code']=big_table.apply(func,axis=1)

big_table_exp = big_table.assign(reject_code=big_table['reject_code'].str.split(', ')).explode('reject_code')
result = big_table_exp.groupby('reject_code')['user_id'].agg(['count', 'nunique']).reset_index()
result.columns = ['reject_code', '总数', '去重后的个数']
result.sort_values(by='去重后的个数',ascending=False,inplace=True)

# 目标客群的画像
target_user = [u for u, reject_code in zip(list(big_table_exp.user_id), list(big_table_exp.reject_code)) if reject_code == 'F-R360-001-9999']

result_temp = big_table_exp[big_table_exp.user_id.isin(target_user)].groupby('reject_code')['user_id'].agg(['count', 'nunique']).reset_index()
result_temp.columns = ['规则码', '总数', '去重后的个数']
result_temp.sort_values(by='去重后的个数',ascending=False,inplace=True)
result_temp=pd.merge(result_temp,reject_code_pzb,how='left',on='规则码')


js1 = big_table.loc[big_table.user_id.isin(target_user),['user_id','reject_code','typeR360']]
js1_temp = js1.groupby(['reject_code','typeR360'])['user_id'].size()



# sql = """ 	select  user_id,reason_code,credit_result from juin_loan_core_prd.risk_credit_result where process_node=0    """
# risk_credit_result=pd.read_sql(sql,cnx)
# js=big_table[['user_id','reject_code']]
# js=pd.merge(js,risk_credit_result,how='left',on='user_id')

big_table.duplicated(subset='user_id').any()#判断pd是否有重复值


# 与数据库中的risk表匹配一下



#20240122看步步高、H5、玖富的分布
df_base = pd.read_excel(r"D:\Work\out_data\分析类\步步高urule变量分布\urule变量.xlsx",sheet_name=r'Sheet1')

list850=[
'huiyan23Bj', 'huomou22Bj', 'qingyun22Bj', 'xingyu22Bj',
'zxjcR360', 'zxqsScoreR360', 'zxsjwlScoreR360',
'zxxdfxyzR360', 'zxxwfxR360',  'tengyun108Td',
'tdDeScore', 'tdDtScore', 'tdGlScore',
'tdXwScore', 'fxpCjHjfv4Wy', 'fxptyWy',   'zzpfLspgV2Wy'
    ]

for var in list850:
    big_table[var]=big_table[var].astype(float)
    
list100=[
'wddtScoreR360','zyScoreR360','tengyun239Td',
'czqdtgzScoreWy', 'dtsqScoreWy', 'dzqdtgzScoreWy','fyhdtgzScoreWy','yhdtgzScoreWy', 'qt3NdfV1Wy', 'qtNdfV1Wy'
   ]
for var in list100:
    big_table[var]=big_table[var].astype(float)
    
big_table.rename(columns={'user_id':'id'},inplace=True)

jf_inside850=big_table.loc[big_table.channel_id==7,list850+['id']]

for var in [col for col in jf_inside850.columns if col!='id' ]:
    bins =[float('-inf'),350,450,550,650,750,850,float('inf')]
    temp_intervals = pd.cut(jf_inside850[var], bins=bins, ordered=True)
    jf_inside850[var+'_intervals_pf'] = temp_intervals
    # 添加 'QS' 到类别中
    jf_inside850[var+'_intervals_pf'] = jf_inside850[var+'_intervals_pf'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    jf_inside850[var+'_intervals_pf'].fillna('DEFAULT', inplace=True)

for index,var in enumerate([col for col in jf_inside850.columns if '_intervals_pf' in col]):
    temp = jf_inside850.groupby(var)[var].count().rename_axis("bins").reset_index()
    temp['百分比'+var] = temp[var] / temp[var].sum() 
    if index==0:
        temp_df=temp
    else:
        temp_df=pd.merge(temp_df,temp,how='outer',on='bins')
temp_df4=temp_df.copy()
temp_df5=temp_df.copy()
temp_df7=temp_df.copy()

temp_df7.to_excel(r'D:\Work\out_data\分析类\步步高urule变量分布\8507.xlsx')
   
# 100系列
jf_inside100=big_table.loc[big_table.channel_id==7,list100+['id']]

for var in [col for col in jf_inside100.columns if col!='id' ]:
    bins =[float('-inf'),20,30,40,50,60,70,80,float('inf')]
    temp_intervals = pd.cut(jf_inside100[var], bins=bins, ordered=True)
    jf_inside100[var+'_intervals_pf'] = temp_intervals
    # 添加 'QS' 到类别中
    jf_inside100[var+'_intervals_pf'] = jf_inside100[var+'_intervals_pf'].cat.add_categories('DEFAULT')
    # 使用 fillna 填充缺失值
    jf_inside100[var+'_intervals_pf'].fillna('DEFAULT', inplace=True)
    
for index,var in enumerate([col for col in jf_inside100.columns if '_intervals_pf' in col]):
    temp = jf_inside100.groupby(var)[var].count().rename_axis("bins").reset_index()
    temp['百分比'+var] = temp[var] / temp[var].sum() 
    if index==0:
        temp_df=temp
    else:
        temp_df=pd.merge(temp_df,temp,how='outer',on='bins')
        
        
# temp_df4100=temp_df.copy()
# temp_df5100=temp_df.copy()
temp_df7100=temp_df.copy()

temp_df4100.to_excel(r'D:\Work\out_data\分析类\步步高urule变量分布\1004.xlsx')

# 年龄
jf_inside100=big_table.loc[big_table.channel_id==7, ['id','age']]
bins =[float('-inf'),20,25,30,35,40,45,50,55,float('inf')]
temp_intervals = pd.cut(jf_inside100['age'], bins=bins, ordered=True)
jf_inside100['age'+'_intervals_pf'] = temp_intervals
# 添加 'QS' 到类别中
jf_inside100['age'+'_intervals_pf'] = jf_inside100['age'+'_intervals_pf'].cat.add_categories('DEFAULT')
# 使用 fillna 填充缺失值
jf_inside100['age'+'_intervals_pf'].fillna('DEFAULT', inplace=True)
 
temp = jf_inside100.groupby('age'+'_intervals_pf')['age'].count().rename_axis("bins").reset_index()
temp['百分比'+'age'] = temp['age'] / temp['age'].sum() 
temp_age7=temp.copy()
 
temp_age7.to_excel(r'D:\Work\out_data\分析类\步步高urule变量分布\年龄7.xlsx')


# 单独看某个评分近3个月的数据
alone_pf=big_table[['id','fxpCjHjfv4Wy']]





#%%-----------白名单配置:第一批，第二批

import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate'
engine = create_engine(db_url)

first_kyf=pd.read_excel(r'D:\Work\out_data\快银付\白名单\邮件存档\20240129第一批白名单.xlsx',sheet_name=r'129邮件_数据库版本')
first_kyf=first_kyf[(first_kyf.备注.isna()) & (first_kyf.渠道来源=='kyf')]
first_kyf.rename(columns={'渠道来源':'channel_source','id_number_md5':'id_number'},inplace=True)
first_kyf.to_sql(name='white_list_kyf1', con=engine, if_exists='replace', index=False)


first_yxj=pd.read_excel(r'D:\Work\out_data\快银付\白名单\邮件存档\20240129第一批白名单.xlsx',sheet_name=r'129邮件_数据库版本')
first_yxj=first_yxj[(first_yxj.备注.isna()) & (first_yxj.渠道来源=='yxj')]
first_yxj.rename(columns={'渠道来源':'channel_source','id_number_md5':'id_number'},inplace=True)
first_yxj.to_sql(name='white_list_yxj1', con=engine, if_exists='replace', index=False)



second_kyf=pd.read_excel(r'D:\Work\out_data\快银付\白名单\邮件存档\快银付第二批白名单.xlsx',sheet_name=r'Sheet1')
second_kyf.rename(columns={'id_card_md5':'id_number'},inplace=True)
second_kyf.to_sql(name='white_list_kyf2', con=engine, if_exists='replace', index=False)

second_kyf.columns

#%%-----------urule知识包：基于场景的批量测试,思汝经常要的apply_result_info_juin
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ 	select  * from juin_loan_core_prd.risk_request_apply where process_node=0 and channel_id=7  """
risk_request_apply=pd.read_sql(sql,cnx)

# risk_request_apply=risk_request_apply[risk_request_apply['request_body'].str.contains('inputAntiFraudParameter', case=False, na=False)].reset_index()
# risk_request_apply=risk_request_apply[risk_request_apply.index!=293].reset_index()

table_list1=[
'inputAntiFraudParameter',
'inputApplyParameter',
'inputThirdParameter']

table_list2=[
'inputBaiRongParameter',
'inputDianHuaParameter',
'inputIcekreditParameter',
'inputOcrParameter',
'inputRong360Parameter',
'inputTdParameter',
'inputTongDunParameter',
'inputUnionPayParameter',
'inputWeiYanParameter']


for i in range(len(risk_request_apply)):
    try:
        temp_json=json.loads(risk_request_apply['request_body'][i])
        for table_name1 in table_list1:
            if table_name1!='inputThirdParameter':
                exec(f"{table_name1+'_temp'} = pd.DataFrame(temp_json[table_name1],index=[i])")
                exec(f"{table_name1+'_temp'}['user_id'] = risk_request_apply['user_id'][i]")
                exec(f"{table_name1+'_temp'}['flow_id'] = risk_request_apply['flow_id'][i]")
                exec(f"{table_name1+'_temp'}['order_id'] = risk_request_apply['order_id'][i]")
                if i==0:
                    exec(f"{table_name1} = {table_name1+'_temp'}")
                    
                else:
                    exec(f"{table_name1} = pd.concat([{table_name1},{table_name1+'_temp'}])")
            else:
                for table_name2 in table_list2:
                    exec(f"{table_name2+'_temp'} =pd.DataFrame(temp_json['inputThirdParameter'][table_name2],index=[i])")
                    exec(f"{table_name2+'_temp'}['user_id'] = risk_request_apply['user_id'][i]")
                    exec(f"{table_name2+'_temp'}['flow_id'] = risk_request_apply['flow_id'][i]")
                    exec(f"{table_name2+'_temp'}['order_id'] = risk_request_apply['order_id'][i]")
                    if i==0:
                        exec(f"{table_name2} = {table_name2+'_temp'}")
                    else:
                        exec(f"{table_name2} = pd.concat([{table_name2},{table_name2+'_temp'}])")
    except:
        continue

# 看每个user_id命中的规则,这里直接将拒绝码打平，虽然BI那里已经实现sql打平
# for i in range(len(risk_request_apply)):
#     try:
#         temp_json=json.loads(risk_request_apply['response_body'][i])
#         temp=pd.DataFrame(temp_json)
#         temp['user_id']=risk_request_apply['user_id'][i]
#         temp['flow_id']=risk_request_apply['flow_id'][i]
#         temp['order_id']=risk_request_apply['order_id'][i]
        
#         if i==0:
#             end=temp
#         else:
#             end=pd.concat([end,temp])
#     except Exception as e:
#         continue


# 将申请、反欺诈的存入风控库，技术已经将其他三方的表存在对应的result_info表
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate'
engine = create_engine(db_url)

# inputApplyParameter.drop(columns=['withdrawCnt0D','withdrawCnt3M','withdrawFailCnt3M'],inplace=True)#'withdrawCnt0D','withdrawCnt3M','withdrawFailCnt3M'都是缺失
inputApplyParameter.to_sql(name='apply_result_info_juin', con=engine, if_exists='replace', index=False)
# a1=inputApplyParameter.columns
inputAntiFraudParameter.to_sql(name='fraud_result_info_juin', con=engine, if_exists='replace', index=False)

js_loan_order.to_sql(name='js_loan_order', con=engine, if_exists='replace', index=False)

sql = """ 	select  * from juin_loan_core_prd.risk_request_apply where process_node=1   """
risk_request_apply=pd.read_sql(sql,cnx)

# risk_request_apply=risk_request_apply[risk_request_apply['request_body'].str.contains('inputAntiFraudParameter', case=False, na=False)].reset_index()
# risk_request_apply=risk_request_apply[risk_request_apply.index!=293].reset_index()

table_list1=[
'inputAntiFraudParameter',
'inputApplyParameter',
'inputThirdParameter']

table_list2=[
'inputBaiRongParameter',
'inputDianHuaParameter',
'inputIcekreditParameter',
'inputOcrParameter',
'inputRong360Parameter',
'inputTdParameter',
'inputTongDunParameter',
'inputUnionPayParameter',
'inputWeiYanParameter']


    
for i in range(len(risk_request_apply)):
    try:
        temp_json=json.loads(risk_request_apply['request_body'][i])
        for table_name1 in table_list1:
            if table_name1!='inputThirdParameter':
                exec(f"{table_name1+'_temp'} = pd.DataFrame(temp_json[table_name1],index=[i])")
                exec(f"{table_name1+'_temp'}['user_id'] = risk_request_apply['user_id'][i]")
                exec(f"{table_name1+'_temp'}['flow_id'] = risk_request_apply['flow_id'][i]")
                exec(f"{table_name1+'_temp'}['order_id'] = risk_request_apply['order_id'][i]")
                if i==0:
                    exec(f"{table_name1} = {table_name1+'_temp'}")
                    
                else:
                    exec(f"{table_name1} = pd.concat([{table_name1},{table_name1+'_temp'}])")
            else:
                for table_name2 in table_list2:
                    exec(f"{table_name2+'_temp'} =pd.DataFrame(temp_json['inputThirdParameter'][table_name2],index=[i])")
                    exec(f"{table_name2+'_temp'}['user_id'] = risk_request_apply['user_id'][i]")
                    exec(f"{table_name2+'_temp'}['flow_id'] = risk_request_apply['flow_id'][i]")
                    exec(f"{table_name2+'_temp'}['order_id'] = risk_request_apply['order_id'][i]")
                    if i==0:
                        exec(f"{table_name2} = {table_name2+'_temp'}")
                    else:
                        exec(f"{table_name2} = pd.concat([{table_name2},{table_name2+'_temp'}])")
    except:
        continue
    
    




# 专项测试

df_base = pd.read_excel(r"D:\Work\out_data\urule\提现专项1\df.xlsx",sheet_name=r'Sheet1')



a1=inputAntiFraudParameter[(inputAntiFraudParameter.user_id == 3702) ].drop(columns=['user_id','flow_id','order_id'])
# a1=inputAntiFraudParameter[inputAntiFraudParameter.order_id == 760].drop(columns=['user_id','flow_id','order_id'])
# a1.to_excel(r'D:\Work\out_data\urule\专项1\Fraud.xlsx')
a1.index=['value']
a11=a1.T.reset_index()

a2=inputApplyParameter[(inputApplyParameter.user_id == 3702) ].drop(columns=['user_id','flow_id','order_id'])
# a2=inputApplyParameter[inputApplyParameter.order_id == 760].drop(columns=['user_id','flow_id','order_id'])
# a2.to_excel(r'D:\Work\out_data\urule\专项1\Apply.xlsx')
a2.index=['value']
a21=a2.T.reset_index()

base_out=pd.concat([a11,a21])

for i in range(len(table_list2)):
    exec(f"temp={table_list2[i]}[{table_list2[i]}.user_id == 3702].drop(columns=['user_id','flow_id','order_id'])")
    # exec(f"temp={table_list2[i]}[{table_list2[i]}.order_id == 760].drop(columns=['user_id','flow_id','order_id'])")
    temp.index=['value'] 
    temp=temp.T.reset_index()
    base_out=pd.concat([base_out,temp])

df_base_out=pd.merge(df_base,base_out,how='left',left_on='COL2',right_on='index' )  
df_base_out.to_excel(r'D:\Work\out_data\urule\授信专项1.xlsx')
    # exec(f"a2.to_excel(r'D:\\Work\\out_data\\urule\\专项1\\{table_list2[i]}.xlsx')")
    
 
for tablename in table_list2:
    exec(f"a2={tablename}[{tablename}.user_id == 3702]")
    exec(f"a2.to_excel(r'D:\\Work\\out_data\\urule\\专项1\\{tablename}.xlsx')")



import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

df = pd.read_excel(r"D:\Work\out_data\parse\语雀底表.xlsx",sheet_name=r'申请-变量')

# 创建元数据对象
metadata = MetaData()


# table = Table(
#    'new_table', metadata, 
#    Column('column1', Integer, nullable=True, default=0, comment='This is column1'),
#    Column('column2', String(255), nullable=True, default='default', comment='This is column2'),
# )



# 定义表结构
columns = []
for i in range(df.shape[0]):
    
    column_name = df.iloc[i, 0]
    
    # column_type = df.iloc[i, 3]  # 假设这是列的数据类型
    
    
    # column_default = df.iloc[i, 2]  # 假设这是列的默认值
    
    column_comment = df.iloc[i, 1]  # 假设这是列的备注
    
    columns.append(Column(column_name))

table = Table('new_table', metadata, *columns)

dir(columns)
# 创建表



# 请根据你的数据库信息进行替换
# 将申请、反欺诈的存入风控库，技术已经将其他三方的表存在对应的result_info表
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

db_url = 'mysql+pymysql://songzhipei:Songzhipei_20231106@jyloan-public.rwlb.rds.aliyuncs.com:3306/juin_risk_operate'
engine = create_engine(db_url)

# inputApplyParameter.drop(columns=['withdrawCnt0D','withdrawCnt3M','withdrawFailCnt3M'],inplace=True)#'withdrawCnt0D','withdrawCnt3M','withdrawFailCnt3M'都是缺失
inputApplyParameter.to_sql(name='apply_result_info', con=engine, if_exists='replace', index=False)
# a1=inputApplyParameter.columns
inputAntiFraudParameter.to_sql(name='fraud_result_info', con=engine, if_exists='replace', index=False)


metadata = MetaData(bind=engine)

# 获取对应的表对象
table = Table('fraud_result_info', metadata, autoload_with=engine)

# 添加COMMENT到列上
with engine.connect() as connection:
    # connection.execute(f"COMMENT ON COLUMN fraud_result_info.blBankTel IS '高端的食材'")
    connection.execute(f"ALTER TABLE fraud_result_info MODIFY COLUMN samHomaddr bigint(20) COMMENT '高端的食材'")


connection.close()
    
    


# inputApplyParameter.to_sql(name='test_Fraud', con=engine, if_exists='replace', index=False)
# a=inputApplyParameter.columns
# a1=inputApplyParameter[['withdrawCnt0D','withdrawCnt3M','withdrawFailCnt3M','withdrawCnt0d','withdrawCnt3m','withdrawFailCnt3m']]



#%%-----------钜银贷规则思维导图
import pandas as pd
from pyecharts import options as opts
from pyecharts.charts import Tree
import pyecharts
import os



os.chdir(r"D:\Work\Information\策略_思汝")
df = pd.read_excel(r"D:\Work\Information\策略_思汝\聚银贷策略_钜银贷_zipper.xlsx",sheet_name=r'Sheet1')
df=df[['数据来源','规则大类','规则细类','规则逻辑']]
# df.rename(columns={'数据来源':'col1','规则大类':'col2','规则细类':'col3','规则逻辑':'col4'},inplace=True)#灵活调整顺序方便下面输出
df.rename(columns={'数据来源':'col2','规则大类':'col1','规则细类':'col3','规则逻辑':'col4'},inplace=True)#灵活调整顺序方便下面输出


def build_nested_dict(data_frame, columns, current_level=0):
    result = []
    current_col = columns[current_level]

    for name, group in data_frame.groupby(current_col):
        if current_level == len(columns) - 2:
            children_list = [{"name": row[columns[current_level + 1]]} for _, row in group.iterrows()]
            result.append({"name": name, "children": children_list})
        else:
            child = {"name": name, "children": build_nested_dict(group, columns, current_level + 1)}
            result.append(child)

    return result
# col
columns = ['col1', 'col2', 'col3', 'col4']#需要在这里添加新的列名，最好是col+数字
data_all = build_nested_dict(df, columns)
data_all_list_dg = [{'name': '风控', 'children': data_all}]

# layout="radial"可以完美适配
c = (
    Tree(init_opts=opts.InitOpts(width="100%", height="100vh"))
    .add("", data_all_list_dg,is_expand_and_collapse=True,collapse_interval=0)
    .set_global_opts(datazoom_opts=opts.DataZoomOpts(orient="vertical"))
    # .add("", data_all_list_dg,pos_top=0.1,pos_bottom=0.9,pos_right=0.1)
)
# c.render_notebook()
c.render("fk_tree7.html")

# 二阶



visual_map_opts = opts.VisualMapOpts(
    min_=0,  # 最小值
    max_=10,  # 最大值
    dimension=0,  # 针对第一个维度（X轴）
    is_piecewise=True,  # 分段显示
    pos_right="5%",  # 位置
    pos_bottom="10%",  # 位置
    range_color=["white", "blue"],  # 范围颜色
)

c = (
     Tree(init_opts=opts.InitOpts(width="100%", height="100vh"))
     .set_global_opts(
         visualmap_opts=visual_map_opts,  # 设置初始视图范围
     )
    # Tree(init_opts=opts.InitOpts(width="100%", height="100vh"))
    .add("", data_all_list_dg,is_expand_and_collapse=True)
    # .add("", data_all_list_dg,pos_top=0.1,pos_bottom=0.9,pos_right=0.1)
)
c.render("fk_tree4.html")


#%%-----------转pdf
import pdfkit
def html_to_pdf(html, to_file):
    # 将wkhtmltopdf.exe程序绝对路径传入config对象
    path_wkthmltopdf = r'D:\\WorkSoft\\wkhtmltopdf\\bin\\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=path_wkthmltopdf)
    # 生成pdf文件，to_file为文件路径
    pdfkit.from_file(html, to_file, configuration=config)
    print('完成')

html_to_pdf('fk_tree4.html','fk_tree4.pdf')#这里的html其实是pyecharts，所以转pdf后是空的


(init_opts=opts.InitOpts(width="100%", height="100vh"))



help(pyecharts)
#%%-----------递归法之思维导图
import pandas as pd
from pyecharts import options as opts
from pyecharts.charts import Tree
import os
os.chdir(r"D:\Work\out_data\语雀")
data = [
    {'col1': 'A', 'col2': 'X', 'col3': '1', 'col4': 'a', 'col5': 'I','col6': '你好'},
    {'col1': 'A', 'col2': 'X', 'col3': '1', 'col4': 'a', 'col5': 'I','col6': '你不好'},
    {'col1': 'A', 'col2': 'X', 'col3': '2', 'col4': 'b', 'col5': 'II','col6': '不知道'},
    {'col1': 'A', 'col2': 'Y', 'col3': '1', 'col4': 'c', 'col5': 'I','col6': '不知道'},
    {'col1': 'A', 'col2': 'Y', 'col3': '2', 'col4': 'd', 'col5': 'II','col6': '不知道'},
    {'col1': 'B', 'col2': 'X', 'col3': '1', 'col4': 'e', 'col5': 'I','col6': '不知道'},
    {'col1': 'B', 'col2': 'X', 'col3': '2', 'col4': 'f', 'col5': 'II','col6': '不知道'},
]

df = pd.DataFrame(data)

def build_nested_dict(data_frame, columns, current_level=0):
    result = []
    current_col = columns[current_level]

    for name, group in data_frame.groupby(current_col):
        if current_level == len(columns) - 2:
            children_list = [{"name": row[columns[current_level + 1]]} for _, row in group.iterrows()]
            result.append({"name": name, "children": children_list})
        else:
            child = {"name": name, "children": build_nested_dict(group, columns, current_level + 1)}
            result.append(child)

    return result
# col
columns = ['col1', 'col2', 'col3', 'col4', 'col5','col6']#需要在这里添加新的列名，最好是col+数字
data_all = build_nested_dict(df, columns)
data_all_list_dg = [{'name': '风控', 'children': data_all}]

c = (
    Tree()
    .add("", data_all_list_dg)
)
# c.render_notebook()
c.render("my_tree5.html")





#%%-----------银信金rds4
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds2.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ select order_no,user_id,approval_status,apply_time,finish_time,approval_type,apply_id  from loan_approval.approval_order where deleted=0   """
approval_order=pd.read_sql(sql,cnx)
approval_shouxin=approval_order[approval_order.approval_type==1].rename(columns={'apply_time':'apply_time_shouxin','finish_time':'finish_time_shouxin','approval_status':'approval_status_shouxin'}).drop(columns='approval_type')
approval_shouxin['授信']=1
approval_shouxin['授信通过']=approval_shouxin.apply(lambda x: 1 if x.approval_status_shouxin==2 else 0,axis=1)

approval_tixian=approval_order[approval_order.approval_type==2].rename(columns={'apply_time':'apply_time_tixian','finish_time':'finish_time_tixian','approval_status':'approval_status_tixian'}).drop(columns='approval_type')
approval_tixian['提现']=1
approval_tixian['提现通过']=approval_tixian.apply(lambda x: 1 if x.approval_status_tixian==2 else 0,axis=1)


os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds4.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 存在初筛记录，初筛只有手机号，身份证姓名3要素，这里我们先只匹配授信/提现阶段客户，即存在caseNo
#这个表存在测试用户，导致一个user_id存在多条apply_id,无解，案例：288215
#存在无法在crd_trade_decisioncontent匹配的SERIAL_NUMBER，可能也是测试流程导致，案例340497
# 存在单个授信在审批流程内(人工流程？)多次SERIAL_NUMBER，目前查到的差异就是crd_trade_decisioncontent里查询人行日期#queryPBCDate不同，案例272184
sql = """ select APP_DATA,SERIAL_NUMBER,PRODUCT_CODE,CREATE_TIME  from creditengine_v2.crd_trade_application   """
crd_trade_application=pd.read_sql(sql,cnx)
crd_trade_application['APP_DATA'] = crd_trade_application['APP_DATA'].apply(ast.literal_eval)

# 选择要提取的键
keys_to_extract = ["appIdcard", "customerPhone", "gender", "appCell", "appName", "marriage", "job", "caseNo", "age", "customerName", "customerIdentityNo"]
# 展开字典并创建新的DataFrame列
for key in keys_to_extract:
    crd_trade_application[key] = crd_trade_application['APP_DATA'].apply(lambda x: x.get(key))
    
shouxin_no=crd_trade_application[crd_trade_application['caseNo'].notna() ]# 存在初筛记录，初筛只有手机号，身份证姓名3要素，这里我们先只匹配授信/提现阶段客户，即存在caseNo
# shouxin_no=shouxin_no[shouxin_no['customerIdentityNo'].notna()]#初筛成功后会补上caseNo，还得需要用其他字段如customerIdentityNo是否缺失来锁定caseNo
shouxin_no['caseNo'] = shouxin_no['caseNo'].apply(lambda x: Decimal(x))

#暂时这么处理，如果有问题再过来检视
shouxin_no = shouxin_no.sort_values(by=['caseNo', 'CREATE_TIME'], ascending=[True, False])
shouxin_no.drop_duplicates(subset='caseNo',keep='first',inplace=True)

# 删除原始的app_data列
crd_trade_application = crd_trade_application.drop(columns=['APP_DATA'])
risk_sx = shouxin_no[shouxin_no['caseNo'].isin(approval_shouxin['apply_id'])]
# a_sample_js=crd_trade_application[crd_trade_application.caseNo=='349496']
risk_tx = shouxin_no[shouxin_no['caseNo'].isin(approval_tixian['apply_id'])]

# ------------------------------------------------------------------------------------授信节点
# 将列表中的值转换为逗号分隔的字符串
# serial_numbers = ', '.join(['"' + str(val) + '"' for val in risk_sx.SERIAL_NUMBER])
# # 构建 SQL 查询
# sql_query = f"SELECT SERIAL_NUMBER,APP_RESPONSE_DATA FROM creditengine_v2.crd_trade_decisioncontent  WHERE SERIAL_NUMBER IN ({serial_numbers})"
# crd_trade_decisioncontent = pd.read_sql(sql_query, cnx)
# 下载一趟不容易，赶紧存
# crd_trade_decisioncontent.to_csv(r"D:\Work\out_data\big_table\crd_trade_decisioncontent_sx.csv",index=False)
# sx_deci=crd_trade_decisioncontent.copy()
sx_deci=pd.read_csv(r"D:\Work\out_data\big_table\crd_trade_decisioncontent_sx.csv")
sx_deci['APP_RESPONSE_DATA']=sx_deci.apply(lambda x: x.APP_RESPONSE_DATA.replace('false','0')  ,axis=1)
sx_deci['APP_RESPONSE_DATA']=sx_deci.apply(lambda x: x.APP_RESPONSE_DATA.replace('true','1')  ,axis=1)

sx_st_list1=[]
sx_st_list2=[]
for i in range(sx_deci.shape[0]):
# for i in range(100):
    print(i)
    try:
        a1 = '→'.join(f"{json.loads(my_dict)['nodeName']}" for my_dict in json.loads(sx_deci['APP_RESPONSE_DATA'][i]).get("matchNodes"))
        a2 = '→'.join(f"{json.loads(my_dict)['runFile']}" for my_dict in json.loads(sx_deci['APP_RESPONSE_DATA'][i]).get("matchNodes"))
        sx_st_list1.append(a1)
        sx_st_list2.append(a2)
    except Exception as e:
        continue
sx_st_list1_df=pd.DataFrame(sx_st_list1)
sx_st_list2_df=pd.DataFrame(sx_st_list2)
sx_deci_end=pd.concat([sx_deci,sx_st_list1_df],ignore_index=True,axis=1)
sx_deci_end=pd.concat([sx_deci_end,sx_st_list2_df],ignore_index=True,axis=1)
new_column_names = {0: 'SERIAL_NUMBER', 1: 'APP_RESPONSE_DATA', 2: 'stra_1', 3: 'stra_2'}
sx_deci_end.rename(columns=new_column_names,inplace=True)
sx_deci_end.to_csv(r"D:\Work\out_data\big_table\crd_trade_decisioncontent_sx_str.csv",index=False)

risk_sx_temp=risk_sx[['caseNo','SERIAL_NUMBER']]
risk_sx_temp=risk_sx_temp[risk_sx_temp.SERIAL_NUMBER.isin(sx_deci_end.SERIAL_NUMBER)]
approval_shouxin_temp=pd.merge(approval_shouxin,risk_sx_temp,how='left',left_on='apply_id',right_on='caseNo')
approval_shouxin_temp=pd.merge(approval_shouxin_temp,sx_deci_end,how='left',on='SERIAL_NUMBER')
approval_shouxin_temp['月份']=pd.to_datetime(approval_shouxin_temp['apply_time_shouxin']).dt.strftime('%Y%m')
approval_shouxin_temp.to_csv(r'D:\Work\out_data\big_table\crd_trade_decisioncontent_sx_str_merge.csv')
a_group1=approval_shouxin_temp.groupby('月份').agg({'SERIAL_NUMBER':{np.size,'count'}})
a_group2=approval_shouxin_temp.groupby(['月份','stra_1']).agg({'SERIAL_NUMBER':'count'}).reset_index()



# a_group_temp=approval_shouxin_temp[(approval_shouxin_temp.月份=='202309') & (approval_shouxin_temp.SERIAL_NUMBER.isna())]
# a_sample=approval_shouxin_temp.sample(40)
# a_sample1=crd_trade_application.sample(40)
# ------------------------------------------------------------------------------------提现节点
serial_numbers = ', '.join(['"' + str(val) + '"' for val in risk_tx.SERIAL_NUMBER])
# 构建 SQL 查询
sql_query = f"SELECT SERIAL_NUMBER,APP_RESPONSE_DATA FROM creditengine_v2.crd_trade_decisioncontent  WHERE SERIAL_NUMBER IN ({serial_numbers})"
tx_deci = pd.read_sql(sql_query, cnx)
# 下载一趟不容易，赶紧存
# tx_deci.to_csv(r"D:\Work\out_data\big_table\crd_trade_decisioncontent_tx.csv",index=False)
tx_deci['APP_RESPONSE_DATA']=tx_deci.apply(lambda x: x.APP_RESPONSE_DATA.replace('false','0')  ,axis=1)
tx_deci['APP_RESPONSE_DATA']=tx_deci.apply(lambda x: x.APP_RESPONSE_DATA.replace('true','1')  ,axis=1)


tx_st_list1=[]
tx_st_list2=[]
for i in range(tx_deci.shape[0]):
# for i in range(100):
    print(i)
    try:
        a1 = '→'.join(f"{json.loads(my_dict)['nodeName']}" for my_dict in json.loads(tx_deci['APP_RESPONSE_DATA'][i]).get("matchNodes"))
        a2 = '→'.join(f"{json.loads(my_dict)['runFile']}" for my_dict in json.loads(tx_deci['APP_RESPONSE_DATA'][i]).get("matchNodes"))
        tx_st_list1.append(a1)
        tx_st_list2.append(a2)
    except Exception as e:
        continue
tx_st_list1_df=pd.DataFrame(tx_st_list1)
tx_st_list2_df=pd.DataFrame(tx_st_list2)
tx_deci_end=pd.concat([tx_deci,tx_st_list1_df],ignore_index=True,axis=1)
tx_deci_end=pd.concat([tx_deci_end,tx_st_list2_df],ignore_index=True,axis=1)
new_column_names = {0: 'SERIAL_NUMBER', 1: 'APP_RESPONSE_DATA', 2: 'stra_1', 3: 'stra_2'}
tx_deci_end.rename(columns=new_column_names,inplace=True)

tx_deci_end.to_csv(r"D:\Work\out_data\big_table\crd_trade_decisioncontent_tx_str.csv",index=False)


risk_tx_temp=risk_tx[['caseNo','SERIAL_NUMBER']]
risk_tx_temp=risk_tx_temp[risk_tx_temp.SERIAL_NUMBER.isin(tx_deci_end.SERIAL_NUMBER)]
approval_tixian_temp=pd.merge(approval_tixian,risk_tx_temp,how='left',left_on='apply_id',right_on='caseNo')
approval_tixian_temp=pd.merge(approval_tixian_temp,tx_deci_end,how='left',on='SERIAL_NUMBER')
approval_tixian_temp['月份']=pd.to_datetime(approval_tixian_temp['apply_time_tixian']).dt.strftime('%Y%m')
approval_tixian_temp.to_csv(r'D:\Work\out_data\big_table\crd_trade_decisioncontent_tx_str_merge.csv')
a_group1_tx=approval_tixian_temp.groupby('月份').agg({'SERIAL_NUMBER':{np.size,'count'}})
a_group2_tx=approval_tixian_temp.groupby(['月份','stra_1','stra_2']).agg({'SERIAL_NUMBER':'count'}).reset_index().sort_values(by=['月份', 'SERIAL_NUMBER'], ascending=[True, False])


#%%-----------银信金验证之前白名单为何不在当前逻辑里
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook



first_white_yx=pd.read_excel(r'D:\Work\out_data\银信系列\银信金\第一批银信金白名单异常名单.xlsx',sheet_name=r'Sheet1')







# 第二次完整取数
# 因为当前在库的名单都是结清，结清的异常检视在上面已经结案：除了3个在user_info表查不到(据说是注销行为),其他都是结清时间超过了90天(结清的取数逻辑是90天以内),所以下面的取数逻辑不需要异常检视20240221
# 此次取数含了在贷，下次取数要检视异常客户 20240221

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ select * from juin_loan_core_prd.white_list   """
white_list=pd.read_sql(sql,cnx)
# white_list.channel_source.value_counts()

sql="""
SELECT
			a1.*,a2.credit_result,a3.id_number_md5
			from (
                    select a.user_id,a.flow_id, a.审批结果 from 
                    (
                     select user_id,flow_id,
                     JSON_UNQUOTE(JSON_EXTRACT(response_body, '$.授信审批结果')) AS 审批结果,
                     ROW_NUMBER() over (PARTITION by flow_id order by create_time desc,update_time  desc)  as countid
                     FROM juin_loan_core_prd.risk_request_apply 
                     WHERE process_node IN (0,4) 
                     ) as a
                    where a.countid=1
                    ) as a1 
			left join juin_loan_core_prd.risk_credit_result as a2 on a1.flow_id=a2.flow_id and a2.process_node in (0,4)
			left join juin_loan_core_prd.user_info as a3 on a1.user_id=a3.id
"""
sx=pd.read_sql(sql,cnx)




os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql="""

SELECT

l1.user_id,
real_name_md5 as name_md5 ,
identity_no_md5 as id_number_md5,
phone_md5,
'yxj' as 渠道来源,
total_credit_amount,
FLOOR(total_credit_amount*1.3/100)*100  as credit_limit,
case when max_cycle>=12 then 6 else max_cycle end as loan_term,
31.92 as loan_rate,
case when his_od_day<=0 and 未结清订单数>0  and 最大还款期数>=6 and DATEDIFF(CURRENT_DATE,最近提现时间)>90 then "在贷" else "结清" end  as 白名单类型,
剩余未还本金,
outstanding_credit_amount as 	已用额度,
available_credit_amount	 as 可用额度,
frozen_credit_amount as 冻结额度,
max_cycle	as   最大贷款期数,
最大还款期数 

from (
			SELECT
			a.user_id,
			count(a.order_no) 放款订单数,
			max(withdraw_amount) as 最大提现金额,
			max(withdraw_time) as 最近提现时间,
			count(case when settle_status in (1,2,3) then a.order_no end ) as 结清订单数,
			count(case when settle_status in (2,3) then a.order_no end) as 提前结清笔数,
			count(a.order_no)-count(case when settle_status in (1,2,3) then a.order_no end ) as 未结清订单数,
			max(case when settle_status in (1,3,2) then settle_date end )  as max_settle_date	
			from loan_core.core_loan_order  as a 
			left join loan_core.core_loan_order_extend  as a2 on a.order_no=a2.order_no 
			where loan_status=2 and a.is_del=0 and a.loan_channel_code not in ('ceshi') and product_name not in ('灰度测试')
			group by 1
			order by 1
	    ) as l1 
left join (	
          SELECT
					a.user_id,
					max(total_cycle) as max_cycle,
					max(case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time) end )  as his_od_day,
					max(case when repayment_status in (2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end )  as now_od_day,
					count(a.order_no) AS receivable_term,
					count(case when repayment_status in (3) then  a.order_no end ) AS repay_term,
					count(case when repayment_status in (3) and  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)<=0 then  a.order_no end ) AS intime_repay_term,
					max(case when repayment_status in (3) then  receivable_time end )  as max_receivable_time,
				  max(case when repayment_status in (3) then a.cycle end) as 最大还款期数,
					count(case when repayment_status in (1,2) then a.order_no  end) as 剩余未还期数,
				  sum(case when repayment_status in (1,2) then a.receivable_principal  end)剩余未还本金
			
					from loan_core.core_repayment_plan as a 
					left join loan_core.core_loan_order  as a2 on a.order_no=a2.order_no 
					where a.is_del=0 and loan_status=2 and loan_channel_code not in ('ceshi') and product_name not in ('灰度测试')
					group by 1
					order by 1
          ) as l2 on l1.user_id=l2.user_id
					
left join loan_core.core_user_info  as l3 on l1.user_id=l3.id  	
left join loan_core.core_credit_account as l4 on l1.user_id=l4.user_id 
				
where


 l4.quota_status=1 and 
	(
		 (his_od_day<=0 and 未结清订单数=0 and 提前结清笔数>0 and intime_repay_term>=3 and DATEDIFF(CURRENT_DATE,max_settle_date)<=90 ) 
	or (his_od_day<=0 and 未结清订单数=0 and 提前结清笔数=0 and DATEDIFF(CURRENT_DATE,max_settle_date)<=90  ) 
	or 
	(his_od_day<=0 and 未结清订单数>0  and 最大还款期数>=6 and DATEDIFF(CURRENT_DATE,最近提现时间)>90 ) 
	) 
"""
new_white_yx=pd.read_sql(sql,cnx)






sql="""
SELECT
  l2.now_od_day,
	l1.未结清订单数,
	DATEDIFF(CURRENT_DATE,l1.max_settle_date) as  上笔结清距今日期,
	l3.identity_no_md5


from (
			SELECT
			a.user_id,
			count(a.order_no) 放款订单数,
			max(withdraw_amount) as 最大提现金额,
			max(withdraw_time) as 最近提现时间,
			count(case when settle_status in (1,2,3) then a.order_no end ) as 结清订单数,
			count(case when settle_status in (2,3) then a.order_no end) as 提前结清笔数,
			count(a.order_no)-count(case when settle_status in (1,2,3) then a.order_no end ) as 未结清订单数,
			max(case when settle_status in (1,3,2) then settle_date end )  as max_settle_date	
			from loan_core.core_loan_order  as a 
			left join loan_core.core_loan_order_extend  as a2 on a.order_no=a2.order_no 
			where loan_status=2 and a.is_del=0 and a.loan_channel_code not in ('ceshi') and product_name not in ('灰度测试')
			group by 1
			order by 1
	    ) as l1 
left join (	
          SELECT
					a.user_id,
					max(total_cycle) as max_cycle,
					max(case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time) end )  as his_od_day,
					max(case when repayment_status in (2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end )  as now_od_day,
					count(a.order_no) AS receivable_term,
					count(case when repayment_status in (3) then  a.order_no end ) AS repay_term,
					count(case when repayment_status in (3) and  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)<=0 then  a.order_no end ) AS intime_repay_term,
					max(case when repayment_status in (3) then  receivable_time end )  as max_receivable_time,
				  max(case when repayment_status in (3) then a.cycle end) as 最大还款期数,
					count(case when repayment_status in (1,2) then a.order_no  end) as 剩余未还期数,
				  sum(case when repayment_status in (1,2) then a.receivable_principal  end)剩余未还本金
			
					from loan_core.core_repayment_plan as a 
					left join loan_core.core_loan_order  as a2 on a.order_no=a2.order_no 
					where a.is_del=0 and loan_status=2 and loan_channel_code not in ('ceshi') and product_name not in ('灰度测试')
					group by 1
					order by 1
          ) as l2 on l1.user_id=l2.user_id
					
left join loan_core.core_user_info  as l3 on l1.user_id=l3.id  	
left join loan_core.core_credit_account as l4 on l1.user_id=l4.user_id 
""" 
yx_list_base=pd.read_sql(sql,cnx) 

# white_list:钜银库中的白名单
# new_white_yx：当前规则跑出来的新白名单
# yx_list_base：查看钜银库的客户信息
old_white_yx=white_list[white_list.channel_source=='yxj']
old_white_yichang=old_white_yx[~old_white_yx.id_number.isin(new_white_yx.id_number_md5)]

old_white_yichang1=yx_list_base[yx_list_base.identity_no_md5.isin(old_white_yichang.id_number)]#看这里的字段
# old_white_yichang1.to_excel(r'D:\Work\out_data\银信系列\银信金\白名单草稿\当前白名单中不符合新跑批的在银信库表现20240318.xlsx',index=False)
yx2=pd.read_excel(r'D:\Work\out_data\快银付\白名单\邮件存档\银信金第二批白名单含在贷.xlsx',sheet_name=r'Sheet1')
yx2=yx2[['id_number','type']]
old_white_yx=pd.merge(old_white_yx,yx2,how='left',on='id_number')
old_white_yx.type.fillna('结清',inplace=True)
old_white_yx_temp=old_white_yx[['id_number','type']]
old_white_yichang1=pd.merge(old_white_yichang1,old_white_yx_temp,how='left',left_on='identity_no_md5',right_on='id_number')
old_white_yichang1=old_white_yichang1[(old_white_yichang1.上笔结清距今日期<90) | ( old_white_yichang1.上笔结清距今日期.isna())]
# 看一下异常名单是否已授信
old_white_yichang1=pd.merge(old_white_yichang1,sx,how='left',left_on='identity_no_md5',right_on='id_number_md5')


# old_white_yichang12=first_white_yx[~first_white_yx.id_number_md5.isin(new_white_yx.identity_no_md5)] #这批如果取loan_core.core_user_info找不到，大概率是注销了


# 新发白名单要剔除已存在的所有名单，没办法与业务注册或者业务授信名单去重，这里靠系统（产品）设计好一点
# 新跑出来的数据要去excel里手动设计调整，其实也可以代码化，但感觉标准不稳定，可能这次代码写好了，下次得重构
new_white_yx=new_white_yx[~new_white_yx.id_number_md5.isin(white_list.id_number)]
new_white_yx.to_excel(r'D:\Work\out_data\银信系列\银信金\银信金第三批白名单含在贷2_底表20240318.xlsx',index=False)




second_white_yx=pd.read_excel(r'D:\Work\out_data\银信系列\银信金\银信金第二批白名单含在贷.xlsx',sheet_name=r'下发白名单')

second_white_yx1=second_white_yx[second_white_yx.id_number.isin(white_list.id_number)]



#%%-----------银信金宽表 
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook


-----------------------------------------------贷前指标（从订单角度）
申请量:申请时间
批核量:批核时间（通过与拒绝：拒绝的话就是拒绝时间+拒绝原因）
件均:申请件均，批核件均
取消：一般是触发时的create_time
放弃：一般是触发时的create_time
签约：时间，金额
放款：时间，金额






-----------------------------------------------贷中指标（从订单角度）
合同状态：C(未违约正常还款),M1,M2,M3...,S(正常结清),ES(提前结清),WO(核销)
贷款余额：剩余本金+实际发生的利息(按天计息)
剩余本金
os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 想看各个期限的利率

sql = """ select order_no,withdraw_amount,loan_time,loan_channel_code,cycle,is_first_loan  from loan_core.core_loan_order where contract_status=2 and is_del=0 and not (loan_channel_code='ceshi' or product_name='灰度测试')     """
core_loan_order=pd.read_sql(sql,cnx)
cycle_loan=core_loan_order[['order_no','cycle']]


# 还款计划表

sql = """ select order_no,receivable_principal,receivable_composite_fee,repayment_sub_status,total_cycle  from loan_core.core_repayment_plan where is_del=0     """
core_repayment_plan=pd.read_sql(sql,cnx)
core_repayment_plan = core_repayment_plan[core_repayment_plan['order_no'].isin(core_loan_order['order_no'])]

cycle_amount=core_repayment_plan.groupby(['order_no']).agg({'receivable_principal':sum,'receivable_composite_fee':sum}).reset_index()
cycle_amount=pd.merge(cycle_amount,cycle_loan,how='left',on='order_no')

cycle_amount.to_excel(r'D:\Work\out_data\big_table\客户利息.xlsx')



os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds2.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# approval_order按user_id记录授信记录及后续的提现记录，两个节点的order_no不会一样，只能用user_id
# approval_quota.approval_order_id=approval_order.id得到user_id的初次授信额度
# withdrawal_approval_order记录user_id的提额申请及结果(提额,降额,原额度),apply_id一直套用approval_order首次授信的apply_id,即首次授信额度，额度改变都会存储在这里，后续可以考虑用回溯来匹配每个提现时的授信额度
# order_no无重复,approval_sub_status跟allocation_status不可用
sql = """ select order_no,approval_status,apply_time,finish_time,approval_type  from loan_approval.approval_order where deleted=0   """
approval_order=pd.read_sql(sql,cnx)
approval_shouxin=approval_order[approval_order.approval_type==1].rename(columns={'apply_time':'apply_time_shouxin','finish_time':'finish_time_shouxin','approval_status':'approval_status_shouxin'}).drop(columns='approval_type')
approval_tixian=approval_order[approval_order.approval_type==2].rename(columns={'apply_time':'apply_time_tixian','finish_time':'finish_time_tixian','approval_status':'approval_status_tixian'}).drop(columns='approval_type')
approval_tixian['提现']=1
approval_tixian['提现通过']=approval_tixian.apply(lambda x: 1 if x.approval_status_tixian==2 else 0,axis=1)
# approval_tixian['提现月份']=pd.to_datetime(approval_tixian['finish_time_tixian']).dt.strftime('%Y%m')
# kan=approval_tixian.groupby('提现月份').agg({'提现':sum,'提现通过':sum})


os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 用loan_order表反推，order_no无重复

sql = """ select order_no,contract_status,withdraw_amount,loan_time  from loan_core.core_loan_order where  is_del=0 and not (loan_channel_code='ceshi' or product_name='灰度测试')     """
core_loan_order_all=pd.read_sql(sql,cnx)

big_table=pd.merge(approval_tixian,core_loan_order_all,how='left',on='order_no')
big_table['approval_status_tixian'].replace({1:'审批中',2:'审批通过',3:'审批拒绝'},inplace=True)
big_table['contract_status'].replace({1:'待签约',2:'签约成功',3:'签约失败',4:'撤销签约成功',5:'撤销签约失败'},inplace=True)

big_table.to_excel(r'D:\Work\out_data\big_table\银信金宽表.xlsx',index=False)




temp_all=pd.DataFrame()
for var in ['approval_status_shouxin','approval_status_tixian','contract_status']:
    temp=big_table[var].value_counts().to_frame().reset_index()
    temp['varname']=var
    temp.rename(columns={var:'values'},inplace=True)
    if var=='approval_status_shouxin':
        temp_all=temp
    else:
        temp_all= pd.concat([temp_all,temp])

big_table['approval_status_tixian'].value_counts()
# approval_order.duplicated(subset='order_no').sum()

approval_order_loan=approval_order[approval_order.order_no.isin(core_loan_order.order_no)]
approval_order_noloan=approval_order[~approval_order.order_no.isin(core_loan_order.order_no)]

approval_order_noloan_explore= pd.DataFrame()
for var in ['approval_status','approval_sub_status','allocation_status','approval_type','manual']:
    temp=approval_order_noloan[var].value_counts().to_frame().reset_index()
    temp['varname']=var
    temp.rename(columns={var:'values'},inplace=True)
    if var=='approval_status':
        approval_order_noloan_explore=temp
    else:
        approval_order_noloan_explore= pd.concat([approval_order_noloan_explore,temp])

a=pd.crosstab(approval_order_noloan['approval_status'], approval_order_noloan['approval_type'])
a1=approval_order_noloan[(approval_order_noloan.approval_type==2) & (approval_order_noloan.approval_status==3)  ]

a1a=pd.merge(a1,core_loan_order_all,how='left',on='order_no')
a1a.contract_status.value_counts()


sql = """ select *  from loan_server.apply_order   """
apply_order=pd.read_sql(sql,cnx)

apply_order['apply_time']=pd.to_datetime(apply_order['apply_time']).dt.date
apply_order['yes']=apply_order.apply(lambda x: 1 if x.else 0,axis=1)
test1=apply_order.groupby('apply_time').

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds3.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


# 用loan_order表反推，order_no无重复
sql = """ select order_no  from kqj_loan_core.core_loan_order where contract_status=2 and is_del=0 and not (loan_channel_code='ceshi' or product_name='灰度测试')     """
kqj_core_loan_order=pd.read_sql(sql,cnx)

a1a=a1[a1.order_no.isin(kqj_core_loan_order.order_no)]

#%%-----------银信金vintages 

# rds1
# loan_core：核心库
# rds2
# loan_admin：后端网关
# loan_approval：审批库
# loan_server：服务端
# rds3
# loan_collect：催收库
# rds4
# creditengine_v2：风控数据库
# tinyid：风控分布式ID

import os
import pandas as pd

import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook


import pymysql
os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)


# 放款表
sql = """ select order_no,withdraw_amount,loan_time,loan_channel_code,cycle,is_first_loan  from loan_core.core_loan_order where loan_status=2 and is_del=0 and not (loan_channel_code='ceshi' or product_name='灰度测试')     """
core_loan_order=pd.read_sql(sql,cnx)
core_loan_order['loan_time']=pd.to_datetime(core_loan_order['loan_time']).dt.date
a=core_loan_order['loan_time'].min()#2022年7月19日
a1=core_loan_order['loan_time'].max()#2022年9月9日

# 还款明细表 repay_record_type=2是减免
sql = """ select order_no,create_time,repayment_amount from loan_core.core_repay_record_detail where subject_type='principal' and repay_record_type=1 and is_del=0 """
core_repay_record_detail=pd.read_sql(sql,cnx)
core_repay_record_detail['create_time']=pd.to_datetime(core_repay_record_detail['create_time']).dt.date
core_repay_record_detail = core_repay_record_detail[core_repay_record_detail['order_no'].isin(core_loan_order['order_no'])]

# 还款计划表,提前结清的客户会在其中一期还完，然后 后面期数的repayment_status=1(待还款),巧妙的是is_del=1（不存在）
sql = """ select order_no,receivable_time,repayment_time,repayment_sub_status   from loan_core.core_repayment_plan where is_del=0     """
core_repayment_plan=pd.read_sql(sql,cnx)
core_repayment_plan.loc[core_repayment_plan['repayment_sub_status']==3, 'repayment_time'] = '' #将部分还款的repayment_time置换为空
core_repayment_plan['repayment_time']=pd.to_datetime(core_repayment_plan['repayment_time']).dt.date
core_repayment_plan['receivable_time']=pd.to_datetime(core_repayment_plan['receivable_time']).dt.date
core_repayment_plan = core_repayment_plan[core_repayment_plan['order_no'].isin(core_loan_order['order_no'])]
# 计算fpd spd tpd qpd
sql='''
select
		a.order_no,
		fpd.receivable_time as fpd_receivable_time,
		spd.receivable_time as spd_receivable_time,
		tpd.receivable_time as tpd_receivable_time,
		qpd.receivable_time as qpd_receivable_time,
		case when f_od>=1  then 1 else 0 end as fpd1_fz,
		case when f_od>=7  then 1 else 0 end as fpd7_fz,
    	case when f_od>=15 then 1 else 0 end as fpd15_fz,
		case when f_od>=30  then 1 else 0 end as fpd30_fz,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=1 then 1 else 0 end fpd1_fm,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=7 then 1 else 0 end fpd7_fm,
        case when DATEDIFF(NOW(),fpd.receivable_time)>=15 then 1 else 0 end fpd15_fm,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=30 then 1 else 0 end fpd30_fm,
		
		case when f_od>=1  then withdraw_amount else 0 end as fpd1_amt_fz,
		case when f_od>=7  then withdraw_amount else 0 end as fpd7_amt_fz,
       
		case when f_od>=30  then withdraw_amount else 0 end as fpd30_amt_fz,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=1 then withdraw_amount else 0 end fpd1_amt_fm,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=7 then withdraw_amount else 0 end fpd7_amt_fm,
		case when DATEDIFF(NOW(),fpd.receivable_time)>=30 then withdraw_amount else 0 end fpd30_amt_fm,
		

		case when f_od<=0 and s_od>=1  then 1 else 0 end as spd1_fz,
		case when f_od<=0 and s_od>=7  then 1 else 0 end as spd7_fz,
		case when f_od<=0 and s_od>=30  then 1 else 0 end as spd30_fz,
		case when DATEDIFF(NOW(),spd.receivable_time)>=1 then 1 else 0 end spd1_fm,
		case when DATEDIFF(NOW(),spd.receivable_time)>=7 then 1 else 0 end spd7_fm,
		case when DATEDIFF(NOW(),spd.receivable_time)>=30 then 1 else 0 end spd30_fm,
		
		case when f_od<=0 and s_od>=1  then withdraw_amount else 0 end as spd1_amt_fz,
		case when f_od<=0 and s_od>=7  then withdraw_amount else 0 end as spd7_amt_fz,
		case when f_od<=0 and s_od>=30  then withdraw_amount else 0 end as spd30_amt_fz,
		case when DATEDIFF(NOW(),spd.receivable_time)>=1 then withdraw_amount else 0 end spd1_amt_fm,
		case when DATEDIFF(NOW(),spd.receivable_time)>=7 then withdraw_amount else 0 end spd7_amt_fm,
		case when DATEDIFF(NOW(),spd.receivable_time)>=30 then withdraw_amount else 0 end spd30_amt_fm,
		
		
		case when f_od<=0 and s_od<=0 and t_od>=1  then 1 else 0 end as tpd1_fz,
		case when f_od<=0 and s_od<=0 and t_od>=7  then 1 else 0 end as tpd7_fz,
		case when f_od<=0 and s_od<=0 and t_od>=30  then 1 else 0 end as tpd30_fz,
		case when DATEDIFF(NOW(),tpd.receivable_time)>=1 then 1 else 0 end tpd1_fm,
		case when DATEDIFF(NOW(),tpd.receivable_time)>=7 then 1 else 0 end tpd7_fm,
		case when DATEDIFF(NOW(),tpd.receivable_time)>=30 then 1 else 0 end tpd30_fm,
		
		case when f_od<=0 and s_od<=0 and t_od>=1  then withdraw_amount else 0 end as tpd1_amt_fz,
		case when f_od<=0 and s_od<=0 and t_od>=7  then withdraw_amount else 0 end as tpd7_amt_fz,
		case when f_od<=0 and s_od<=0 and t_od>=30  then withdraw_amount else 0 end as tpd30_amt_fz,
		case when DATEDIFF(NOW(),tpd.receivable_time)>=1 then withdraw_amount else 0 end tpd1_amt_fm,
		case when DATEDIFF(NOW(),tpd.receivable_time)>=7 then withdraw_amount else 0 end tpd7_amt_fm,
		case when DATEDIFF(NOW(),tpd.receivable_time)>=30 then withdraw_amount else 0 end tpd30_amt_fm,

		case when f_od<=0 and s_od<=0 and t_od<=0  and q_od>=1  then 1 else 0 end as qpd1_fz,
		case when f_od<=0 and s_od<=0 and t_od<=0  and q_od>=7  then 1 else 0 end as qpd7_fz,
		case when f_od<=0 and s_od<=0 and t_od<=0  and q_od>=30  then 1 else 0 end as qpd30_fz,
		case when DATEDIFF(NOW(),qpd.receivable_time)>=1 then 1 else 0 end qpd1_fm,
		case when DATEDIFF(NOW(),qpd.receivable_time)>=7 then 1 else 0 end qpd7_fm,
		case when DATEDIFF(NOW(),qpd.receivable_time)>=30 then 1 else 0 end qpd30_fm,
		
		case when f_od<=0 and s_od<=0 and t_od<=0  and q_od>=1  then withdraw_amount else 0 end as qpd1_amt_fz,
		case when f_od<=0 and s_od<=0 and t_od<=0  and q_od>=7  then withdraw_amount else 0 end as qpd7_amt_fz,
		case when f_od<=0 and s_od<=0 and t_od<=0  and q_od>=30  then withdraw_amount else 0 end as qpd30_amt_fz,
		case when DATEDIFF(NOW(),qpd.receivable_time)>=1 then withdraw_amount else 0 end qpd1_amt_fm,
		case when DATEDIFF(NOW(),qpd.receivable_time)>=7 then withdraw_amount else 0 end qpd7_amt_fm,
		case when DATEDIFF(NOW(),qpd.receivable_time)>=30 then withdraw_amount else 0 end qpd30_amt_fm


		from loan_core.core_loan_order as a
		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as f_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=1 ) as fpd on a.order_no=fpd.order_no
										
		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as s_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=2 ) as spd on a.order_no=spd.order_no

		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as t_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=3 ) as tpd on a.order_no=tpd.order_no

		left join (SELECT *,case when repayment_status in (3,2) then  DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time)  end  as q_od
							from loan_core.core_repayment_plan where is_del=0 and cycle=4 ) as qpd on a.order_no=qpd.order_no
							
		left join (SELECT order_no,
							 max( DATEDIFF(IFNULL(repayment_time,NOW()),receivable_time) )  as max_od_day           
							from loan_core.core_repayment_plan
							where is_del=0 and repayment_status in (3,2) group by 1 order by 1) as od on a.order_no=od.order_no

		where loan_status=2 and a.is_del=0 and a.loan_channel_code not in ('ceshi') and product_name not in ('灰度测试')  

'''

fstq=pd.read_sql(sql,cnx)







channel_product_info=pd.read_excel(r'D:\Work\out_data\vintage\银信线上线下渠道标识_z.xlsx')
core_loan_order=pd.merge(core_loan_order,channel_product_info,how='left',left_on='loan_channel_code',right_on='渠道简称')
core_loan_order.columns
core_loan_order.drop(columns=['渠道简称'],inplace=True)

def channel_3(x):
    if '武汉捷得' in x.渠道名称:
        return '武汉捷得'
    elif '江苏名恒' in x.渠道名称:
        return '江苏名恒'
    elif '和淳金服' in x.渠道名称:
        return '和淳金服'
    elif '银信金' in x.渠道名称:
        return '银信金'
    elif '数字魔方' in x.渠道名称:
        return '数字魔方'
    else:
        return '其他'
core_loan_order['temp_channel']=core_loan_order.apply(channel_3,axis=1)
core_loan_order.drop(columns=['渠道名称'],inplace=True)
def od_days(x):#计算每个账单日的逾期天数，如果结清则为0
    if x.repayment_time>end_of_month or pd.isna(x.repayment_time):
        return  (end_of_month-x.receivable_time).days
    else:
        return 0
    
def mob(x):
    return (x['cut_date'].year - x['loan_time'].year) * 12 + x['cut_date'].month - x['loan_time'].month

def bin_overdue_days(x):
    if x.od_days==0 or pd.isna(x.od_days):
        return "C"
    elif 0 < x.od_days <= 30:
        return "M1"
    elif 30 < x.od_days <= 60:
        return "M2"
    elif 60 < x.od_days <= 90:
        return "M3"
    elif 90 < x.od_days <= 120:
        return "M4"
    elif 120 < x.od_days <= 150:
        return "M5"
    elif 150 < x.od_days <= 180:
        return "M6"
    else:
        return "WO"
    
# 获取当前日期
current_date = datetime.date.today()

# 银信金第一笔放款日是2022 年 7 月 19 日
target_date = datetime.date(2022, 7, 19)
months_difference = (current_date.year - target_date.year) * 12 + current_date.month - target_date.month
vintage = pd.DataFrame()

with tqdm(total=2) as pbar:
    pbar.set_description('Processing:')
    for m in range(1,months_difference+1):
        time.sleep(0.1)
        pbar.update()
        # while target_date <= current_date:
        year = target_date.year
        month = target_date.month + 1
        if month > 12:
            year += 1
            month = 1
            
        # 计算下个月的第一天
        next_month_first_day =datetime.date(year, month, 1)
        target_date = next_month_first_day
        
        # 计算当前月的最后一天，即下个月第一天的前一天
        end_of_month = next_month_first_day - datetime.timedelta(days=1)
        
        total_principal=core_loan_order[core_loan_order['loan_time']<=end_of_month].copy()
        paid_principal=core_repay_record_detail[core_repay_record_detail['create_time']<=end_of_month].groupby('order_no')['repayment_amount'].sum().reset_index()#已还本金
        temp_user=pd.merge(total_principal,paid_principal,how='left',on='order_no')
        temp_user['cut_date']=end_of_month
        temp_user['mob']=temp_user.apply(mob,axis=1)
        
        temp_od=core_repayment_plan[core_repayment_plan['receivable_time']<=end_of_month].copy()
        temp_od.loc[temp_od['repayment_time']>end_of_month, 'repayment_time'] = pd.NaT
        
        try:#放款当月截止日期会出现空表
            temp_od['od_days']=temp_od.apply(od_days,axis=1)
            temp_od = temp_od.sort_values(by=['order_no', 'receivable_time'], ascending=[True, False])
            result_df = temp_od.groupby('order_no').apply(lambda x: x.iloc[0] if x['od_days'].iloc[0]==0 else x.loc[x['od_days'].idxmax()]).reset_index(drop=True)
            result_df.drop(columns=['repayment_sub_status'],inplace=True)
            temp_user=pd.merge(temp_user,result_df,how='left',on='order_no')
        except Exception as e:
            pass
        
        if m == 1:
            vintage = temp_user
        else:
            vintage = pd.concat([vintage,temp_user])

# 截至昨天部分，先这么算这
end_of_month=datetime.date(2023, 10, 22)
total_principal=core_loan_order[core_loan_order['loan_time']<=end_of_month].copy()
paid_principal=core_repay_record_detail[core_repay_record_detail['create_time']<=end_of_month].groupby('order_no')['repayment_amount'].sum().reset_index()#已还本金
temp_user=pd.merge(total_principal,paid_principal,how='left',on='order_no')
temp_user['cut_date']=end_of_month
temp_user['mob']=temp_user.apply(mob,axis=1)

temp_od=core_repayment_plan[core_repayment_plan['receivable_time']<=end_of_month].copy()
temp_od.loc[temp_od['repayment_time']>end_of_month, 'repayment_time'] = pd.NaT

temp_od['od_days']=temp_od.apply(od_days,axis=1)
temp_od = temp_od.sort_values(by=['order_no', 'receivable_time'], ascending=[True, False])
result_df = temp_od.groupby('order_no').apply(lambda x: x.iloc[0] if x['od_days'].iloc[0]==0 else x.loc[x['od_days'].idxmax()]).reset_index(drop=True)
result_df.drop(columns=['repayment_sub_status'],inplace=True)
temp_user=pd.merge(temp_user,result_df,how='left',on='order_no')
vintage = temp_user
# vintage = pd.concat([vintage,temp_user])


vintage['放款月份']=pd.to_datetime(vintage['loan_time']).dt.strftime('%Y%m')
vintage['repayment_amount'].fillna(0,inplace=True)
vintage['剩余本金']=vintage.apply(lambda x: x['withdraw_amount']-x['repayment_amount'] ,axis=1)
vintage['剩余本金_30']=vintage.apply(lambda x:x['withdraw_amount']-x['repayment_amount'] if x.od_days>=30 else 0,axis=1)
vintage['剩余本金_1']=vintage.apply(lambda x:x['withdraw_amount']-x['repayment_amount'] if x.od_days>=1 else 0,axis=1)
vintage['status']=vintage.apply(bin_overdue_days,axis=1)
# 按放款月份和mob进行排序
vintage = vintage.sort_values(by=['放款月份', 'mob'])

# file_path = r'D:\Work\out_data\vintage\vintage.xlsx'
# # with模块执行结束后自动关闭文件,有点慢
# with pd.ExcelWriter(file_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
#     vintage.to_excel(writer, sheet_name='底表', index=False)

# 这个快，但会覆盖文件
vintage.to_excel(r'D:\Work\out_data\vintage\1.xlsx',index=False)

#%%------------短信提醒文字模板




import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)



sql = """ 	 select * from juin_loan_core_prd.repayment_plan_period  """
repayment_plan_period=pd.read_sql(sql,cnx)

user_info=inputApplyParameter[['user_id','name','tel']]

repayment_plan_period=pd.merge(repayment_plan_period,user_info,on='user_id',how='left')

repayment_plan_period['repayment_date'] = pd.to_datetime(repayment_plan_period['repayment_date']).dt.date


result_strings = []
for index, row in repayment_plan_period.iterrows():
    result_string = f"尊敬的{row['name']}，您本期应还款金额为{row['period_total_amount']}元，还款日为{row['repayment_date'].month}月{row['repayment_date'].day}日，请确保还款银行卡余额充足。"
    result_strings.append(result_string)
repayment_plan_period['短信提示内容']=result_strings
end=repayment_plan_period[['repayment_date','name','period_total_amount','短信提示内容','user_id','tel']]
end.to_excel(r'D:\Work\out_data\需求历史\知识付费\短信提醒1.xlsx',index=False)
    


#%%------------类似sas的DDE
import pandas as pd
import datetime

# 假设你已经有了 atest 数据框
# 创建一个示例的 DataFrame
data = {
    'cut_date': [datetime.date(2023, 8, 31)] * 30,
    '放款月份': ['202306'] * 30,
    'loan_time': pd.date_range(start='2023-01-01', periods=30, freq='D'),
    '剩余本金': range(30)
}

vintage = pd.DataFrame(data)

# 假设你已经有了一个名为 sheet 的 ExcelWriter 对象

# 从 vintage 中选择符合条件的数据
atest = vintage.loc[(vintage.cut_date == datetime.date(2023, 8, 31)) & (vintage.放款月份 == '202306'), ['loan_time', '剩余本金']]


import xlwings as xw

# 连接到已打开的Excel实例或打开一个新的Excel实例
app = xw.App(visible=True)

# 打开工作簿（替换为你的Excel文件路径）
wb = app.books.open(r'D:\Work\out_data\vintage\test.xlsx')

# 选择要更新的工作表
sheet = wb.sheets['Sheet1']  # 替换为你的工作表名称

# 更新单元格的内容
sheet.range('A1:B2000').value = atest.values

# 从Excel获取单元格的值
# cell_value = sheet.range('B2').value
# print(f'Value in cell B2: {cell_value}')

# 关闭工作簿
wb.save()
wb.close()

# 关闭Excel实例
app.quit()
#%%------------测试样本15000
# -------------------------------------------------------------------------------------------------------------钜银玖富样本725
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
import numpy as np
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN_sc.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# risk_credit_result只有通过才有额度，干脆直接解析报文来获取授信额度与提现额度
sql = """ select user_id,DATE_FORMAT(create_time,'%Y-%m-%d') as 回溯日期  from juin_loan_core_prd.risk_credit_result  where process_node=0  """
jyjf=pd.read_sql(sql,cnx)

sql = """ select id,id_number_md5,phone_md5  from juin_loan_core_prd.user_info    """
user_info=pd.read_sql(sql,cnx)

jyjf=pd.merge(jyjf,user_info,left_on='user_id',right_on='id',how='left')
jyjf=jyjf[jyjf.id.notna()]
jyjf.drop(columns=['id'],inplace=True)

jyjf.rename(columns={"id_number_md5":"identity_no_md5"},inplace=True)
jyjf['source']='jf'

jyjf.duplicated(subset='user_id').any()#判断pd是否有重复值

# -------------------------------------------------------------------------------------------------------------聚合吧样本1122
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
from openpyxl import load_workbook
import ast 
from decimal import Decimal
import sqlalchemy
import pandas as pd
from sqlalchemy import create_engine, Table, MetaData, Column, Integer, String, Text

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"JYLOAN.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ 	select  user_id2 as user_id,md5(phone) as phone_md5 ,md5(id) as identity_no_md5
from juin_risk_operate.yx_fk_test1113_v2 where type=1  """
list_sr=pd.read_sql(sql,cnx)

list_sr.user_id=list_sr.user_id.astype('int64')

list_sr['source']='jhb'
list_sr['回溯日期']='2023-11-14'


list_sr.duplicated(subset='user_id').any()#判断pd是否有重复值

# -------------------------------------------------------------------------------------------------------------银信金全体放款样本5666
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
import warnings

# 设置警告过滤器，忽略特定类型的警告
warnings.filterwarnings("ignore", category=FutureWarning)

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ select user_id,order_no,DATE_FORMAT(loan_time,'%Y-%m-%d') as 回溯日期  from loan_core.core_loan_order where contract_status=2 and is_del=0 and not (loan_channel_code='ceshi' or product_name='灰度测试')     """
loan_test=pd.read_sql(sql,cnx)
loan_test.drop_duplicates(subset='user_id',keep='last',inplace=True)

sql = """ select id,identity_no_md5,phone_md5  from loan_core.core_user_info     """
core_user_info=pd.read_sql(sql,cnx)
loan_test=pd.merge(loan_test,core_user_info,left_on='user_id',right_on='id',how='left')
loan_test=loan_test[loan_test.id.notna()]
loan_test.drop(columns=['id'],inplace=True)

loan_test['source']='yxjfk'

loan_test.duplicated(subset='user_id').any()#判断pd是否有重复值
# -------------------------------------------------------------------------------------------------------------银信金拒绝样本7478,记得按拒绝时间降序
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
import warnings

# 设置警告过滤器，忽略特定类型的警告
warnings.filterwarnings("ignore", category=FutureWarning)

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds2.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

sql = """ select user_id,i from loan_server.apply_order where is_del=0 and  apply_type=1 and apply_status=2 order by 1,  2 desc """
apply_order=pd.read_sql(sql,cnx)
apply_order.drop_duplicates(subset=['user_id'],keep='first' ,  inplace=True)

apply_order=pd.merge(apply_order,core_user_info,left_on='user_id',right_on='id',how='left')
apply_order=apply_order[apply_order.id.notna()]

apply_order.sort_values(by='apply_time',inplace=True,ascending=False)
apply_order=apply_order.reset_index().drop(columns='index')
apply_order['回溯日期'] = apply_order['apply_time'].apply(lambda x: x.strftime('%Y-%m-%d'))

apply_order_7487=apply_order[apply_order.index<=7478] 

apply_order_7487.drop(columns=['id','apply_time'],inplace=True)

apply_order_7487['source']='yxjjj'

apply_order_7487.duplicated(subset='user_id').any()#判断pd是否有重复值


data_test=pd.concat([jyjf,loan_test],axis=0)
data_test=pd.concat([data_test,apply_order_7487],axis=0)
data_test=pd.concat([data_test,list_sr],axis=0)
data_test['new_column'] = range(len(data_test))

data_test.to_excel(r"D:\Work\Information\三方对接\银豆\测试样本\测试样本.xlsx",index=False)


#%%------------测试样本2000
import os
import pandas as pd
import pymysql
import json
import datetime
import time
from tqdm import tqdm, trange
import warnings

# 设置警告过滤器，忽略特定类型的警告
warnings.filterwarnings("ignore", category=FutureWarning)

os.chdir(r"D:\Work\out_data\PythonCode")
with open(r"rds1.json") as db_config:
    cnx_args = json.load(db_config)
cnx = pymysql.connect(**cnx_args)

# 还款计划表
sql = """ select order_no,user_id,receivable_time,repayment_time,repayment_sub_status  from loan_core.core_repayment_plan where is_del=0  and user_id>100  and start_date < DATE_SUB( NOW(), INTERVAL 30 DAY )   """
repay_plan_test=pd.read_sql(sql,cnx)
repay_plan_test.loc[repay_plan_test['repayment_sub_status']==3, 'repayment_time'] = ''
repay_plan_test['repayment_time']=pd.to_datetime(repay_plan_test['repayment_time']).dt.date
repay_plan_test['receivable_time']=pd.to_datetime(repay_plan_test['receivable_time']).dt.date
repay_plan_test=repay_plan_test[repay_plan_test.receivable_time<=datetime.date.today()]

# 放款表
sql = """ select user_id,order_no,loan_channel_code,cycle  from loan_core.core_loan_order where contract_status=2 and is_del=0 and not (loan_channel_code='ceshi' or product_name='灰度测试')     """
loan_test=pd.read_sql(sql,cnx)

loan_test['loan_time']=pd.to_datetime(loan_test['loan_time']).dt.date

# 用户表
sql = """ select id,real_name_md5,identity_no_md5,phone_md5  from loan_core.core_user_info     """
core_user_info=pd.read_sql(sql,cnx)

def od_days_now(x):#计算每个账单日的逾期天数，未结清就用今天-账单日：结清就用结清日期-账单日(dpd概念)
    if  pd.isna(x.repayment_time):
        return  (datetime.date.today()-x.receivable_time).days
    else:
        return (x.repayment_time-x.receivable_time).days

repay_plan_test['od_days']=repay_plan_test.apply(od_days_now,axis=1)
user_max_od=repay_plan_test.groupby('order_no')['od_days'].max().reset_index()

# 逾期样本
test=user_max_od[user_max_od.od_days>0]
loan_testa=pd.merge(loan_test,test,how='right',on='order_no')
loan_testa.sort_values(by=['user_id','od_days'], ascending=[True, False],inplace=True)
loan_testa.drop_duplicates(subset='user_id',keep='first',inplace=True)
loan_testa1=loan_testa[loan_testa.cycle.isin([1,3,6])]
loan_testa2=loan_testa[loan_testa.cycle==12].head(523)
loan_test_bad=pd.concat([loan_testa1,loan_testa2])

# 正常样本
test=user_max_od[user_max_od.od_days==0]
loan_testa=pd.merge(loan_test,test,how='right',on='order_no')
loan_testa=loan_testa[~loan_testa.user_id.isin(loan_due_test.user_id)]
loan_testa.sort_values(by=['user_id','cycle'], ascending=[True, True],inplace=True)
loan_testa.drop_duplicates(subset='user_id',keep='first',inplace=True)

# loan_testa.cycle.value_counts()
loan_test_good1=loan_testa[loan_testa.cycle==1].head(67)#只有67个没有重复的
loan_test_good3=loan_testa[loan_testa.cycle==3].head(38)
loan_test_good6=loan_testa[loan_testa.cycle==6].head(362)
loan_test_good12=loan_testa[loan_testa.cycle==12].head(533)#给cycle==1补充
loan_test_good=pd.concat([loan_test_good1,loan_test_good3,loan_test_good6,loan_test_good12])
loan_test_out=pd.concat([loan_test_good,loan_test_bad])

loan_test_out=pd.merge(loan_test_out,core_user_info,how='left',left_on='user_id',right_on='id')

# 修补，因为user表缺失，尴尬
loan_test_out=loan_test_out[~loan_test_out.real_name_md5.isna()]
loan_test_good12_all=loan_testa[loan_testa.cycle==12]
loan_test_good12_all=pd.merge(loan_test_good12_all,core_user_info,how='left',left_on='user_id',right_on='id')
loan_test_good12_all=loan_test_good12_all[~loan_test_good12_all.real_name_md5.isna()]
loan_test_good12_all=loan_test_good12_all[~loan_test_good12_all.user_id.isin(loan_test_out.user_id)]
loan_test_good12_all_14=loan_test_good12_all.head(14)
loan_test_out=pd.concat([loan_test_out,loan_test_good12_all_14])

loan_test_out['real_name_md5'] = loan_test_out['real_name_md5'].str.lower()
loan_test_out['identity_no_md5'] = loan_test_out['identity_no_md5'].str.lower()
loan_test_out['phone_md5'] = loan_test_out['phone_md5'].str.lower()
loan_test_out.to_excel(r'D:\Work\Information\三方对接\测试数据\样本2000_原始.xlsx',index=False)


    期限与渠道
#%%------------给思汝看的海外设备变量
import os
import jsonO
import pandas as pd
import time
from itertools import islice
import datetime
import saspy
import numpy as np



os.chdir(r"C:\Users\zp457\Desktop\新建文件夹\test")


with open('PL201905112017597360066_267434.json','r',encoding='UTF-8') as temp:
    temp_json = json.load(temp)
    temp_df = pd.DataFrame(temp_json['result_desc'])

 
temp_json1=temp_json['result_desc']
temp_df = pd.DataFrame(temp_json1['INFOANALYSIS']['device_info'])

temp_df.to_excel(r'device.xlsx')
 
 
keys_list = list(temp_json1['INFOANALYSIS'].keys())

temp_json.keys


#%%-----------快银.准备
# 用户画像草案
class UserProfile:
    def __init__(self, user_id, age, gender, locsation, interests):
        self.user_id = user_id
        self.age = age
        self.gender = gender
        self.location = location
        self.interests = interests

# 创建用户画像列表
user_profiles = [
    UserProfile(1, 25, 'Male', 'New York', ['Music', 'Sports']),
    UserProfile(2, 30, 'Female', 'Los Angeles', ['Art', 'Travel']),
    UserProfile(3, 22, 'Male', 'Chicago', ['Technology', 'Movies']),
    # 可以继续添加更多用户画像
]

# 打印用户画像信息
for user_profile in user_profiles:
    print(f"User ID: {user_profile.user_id}")
    print(f"Age: {user_profile.age}")
    print(f"Gender: {user_profile.gender}")
    print(f"Location: {user_profile.location}")
    print(f"Interests: {', '.join(user_profile.interests)}")
    print("\n")




散点图：如果您有多个维度的数据，可以使用散点图来展示两个维度之间的关系。例如，您可以将年龄与购买金额进行对比，以查看是否存在某种趋势或相关性。

热力图：热力图可以用来显示两个不同维度之间的相对密度。它可以帮助您发现用户在哪些特定区域或属性上集中。

雷达图：雷达图可以将多个维度的数据展示在同一个图表中，帮助您直观地比较不同用户在多个属性上的表现。

地图：如果用户有地理位置信息，您可以使用地图来显示用户在不同地区的分布情况，以及在各个地区的属性分布。

报告和仪表盘：创建报告或仪表盘可以将多个图表和可视化元素结合起来，以呈现用户画像的全貌。这样的报告可以用来详细阐述不同维度之间的关系和趋势。

在选择展示形式时，考虑您的受众以及他们最容易理解的方式。合适的可视化方式可以帮助您更好地传达数据，并使用户画像更直观且易于理解。


#%%------------针对经常统计的逻辑设定成函数，比如groupby
import pandas as pd

# 示例数据
data = {
    'Date': ['2023-01-01', '2023-01-01', '2023-01-02', '2023-01-02'],
    'Category': ['A', 'B', 'A', 'B'],
    'Value': [10, 20, 15, 25]
}

df = pd.DataFrame(data)

# 创建统计函数
def calculate_statistics(df, group_by_column, value_column):
    result = df.groupby(group_by_column)[value_column].sum()
    return result

# 调用统计函数
result_by_category = calculate_statistics(df, 'Category', 'Value')
result_by_date = calculate_statistics(df, 'Date', 'Value')



import numpy as np

def custom_log_loss(y_true, y_pred, alpha):
    N = len(y_true)
    loss = -np.sum(y_true * np.log(y_pred) + (1 - y_true) * np.log(1 - y_pred)) / N
    regularizer = 0.5 * alpha * np.sum(np.square(y_pred)) / N  # 自定义正则化项
    total_loss = loss + regularizer
    return total_loss

# 示例数据
y_true = np.array([0, 1, 1, 0, 1])
y_pred = np.array([0.2, 0.8, 0.9, 0.3, 0.6])  # 预测概率值

alpha = 0.1  # 正则化系数
loss = custom_log_loss(y_true, y_pred, alpha)

print("Custom Log Loss:", loss)

#%%------------封装groupby的示例




# 导入自定义的groupby函数
from mygroupby import custom_groupby#已放D:\Work\out_data\PythonCode\package


# 创建一个示例DataFrame
data = {'Category': ['A', 'B', 'A', 'B', 'A'],
        'Value': [10, 20, 15, 25, 12]}
df = pd.DataFrame(data)

# 定义聚合函数字典
aggregation_functions = {'Value': 'sum'}

# 使用自定义的groupby函数进行分组和聚合
result = custom_groupby(df, 'Category', aggregation_functions)

# 打印结果
print(result)


#%%------------自动化提送github（初稿）
import schedule
import time
from git import Repo

def commit_and_push():
    repo_path = "/path/to/your/local/repository"  # 本地仓库的路径
    repo = Repo(repo_path)
    repo.git.add(all=True)
    repo.index.commit("Automatic daily commit")
    repo.remotes.origin.push()

# 每天的特定时间执行 commit_and_push 函数
schedule.every().day.at("12:00").do(commit_and_push)  # 这里设置为每天中午 12 点

while True:
    schedule.run_pending()
    time.sleep(1)
#%%------------使用结巴给单位名称分词
import jieba
import pandas as pd
from collections import Counter


# 分词并且放进set里，利用set的元素不重复特点去重关键字
def segment_text(text):
    seg_list = jieba.cut(text)
    return set(" ".join(seg_list).split())

# if __name__ == "__main__":
#     input_text = "深圳市南山区斦蓝咨询服务工作室"
#     segmented_text = segment_text(input_text)
#     print("分词结果：", segmented_text)

# data = {
#     'company': ['深圳市南山区服务斦蓝咨询服务工作室', '深圳市舞蹈斦蓝舞蹈工作室', '深圳市蔚蓝舞蹈工作室']
# }
# df = pd.DataFrame(data)


#****************************************************** 入参：df,   df.company
#****************************************************** 出参:element_percentage

df['out_str_set']=df.company.apply(segment_text)

# 将目标列内的每个set的元素合并在list里，利用list元素可以重复特点保留关键字方便后续的统计占比
set_to_list = []
for s in df['out_str_set']:
    set_to_list.extend(list(s))
    
element_count = Counter(set_to_list)# 使用 Counter 统计list内每个元素的出现次数

total_elements = len(df['out_str_set'])# 计算分母

element_percentage = {element: count / total_elements * 100 for element, count in element_count.items()}# 计算每个元素出现的百分比

#%%------------利用上面步骤生成词云
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from IPython.display import display
import numpy as np
from PIL import Image


# 假设 element_percentage 是一个字典，包含关键词及其占比
element_percentage = {
    '南山区': 10,
    '咨询服务': 10,
    '工作室': 30,
    '斦': 20,
    '服务': 10,
    '深圳市': 30,
    '舞蹈': 20,
    '蓝': 20,
    '蔚蓝': 10
    
    # ...
}

mask = np.array(Image.open(r"D:\Work\out_data\PythonCode\test.png"))  # 自定义形状的图像文件
wordcloud = WordCloud(font_path="D:\Work\out_data\PythonCode\SourceHanSerifSC-Bold.otf",mask=mask,  background_color='white')
# 生成词云
wordcloud.generate_from_frequencies(element_percentage)

wordcloud.to_file("wordcloud.png")#可能是词太少，覆盖不出来轮廓，等以后词多了再研究

#%%------------使用地图（待完善）
import folium
import json

# 创建一个地图对象
m = folium.Map(location=[35, 105], zoom_start=5)  # 使用合适的中心坐标和缩放级别

# 读取包含柱状图数据的 GeoJSON 文件
with open(r'C:\Users\zp457\Desktop\新建文件夹\china.geoJson', 'r', encoding='utf-8') as geojson_file:
    geojson_data = json.load(geojson_file)

# 假设 bar_data 是一个包含柱状图数据的字典，其中每个键代表一个省份或区域，值是柱状图的高度
bar_data = {
    '广东': 50,
    '海南': 30,
    '北京': 20,
    # ...
}

# 添加柱状图数据到地图
for feature in geojson_data['features']:
    province = feature['properties']['name']
    if province in bar_data:
        height = bar_data[province]
        folium.Marker(
            location=[feature['properties']['centroid'][1], feature['properties']['centroid'][0]],
            icon=folium.DivIcon(html=f'<div style="width: 20px; height: {height}px; background-color: blue;"></div>')
        ).add_to(m)

# 显示地图
m.save("china_map_with_bars.html")  # 保存为 HTML 文件


#%%------------使用地图2（待修复）



是的，GitHub 上有许多关于使用地图来可视化数据的项目。您可以在 GitHub 上的搜索栏中输入相关关键词，例如 "map visualization", "geospatial data visualization" 等，来找到适合您需求的项目。以下是一些可能的项目示例：

awesome-geojson： 这个项目收集了许多关于 GeoJSON 数据的资源，包括示例、工具和库，可以用于地图数据的可视化。
项目链接：awesome-geojson

Leaflet： Leaflet 是一个用于创建交互式地图的 JavaScript 库，它在 GitHub 上有自己的仓库。您可以查看这个仓库中的示例和文档，学习如何使用 Leaflet 来可视化地理数据。
仓库链接：Leaflet

d3-geo： D3.js 是一个强大的 JavaScript 数据可视化库，d3-geo 子模块专门用于地理数据可视化。您可以在这个项目中找到关于地图投影、路径生成等方面的示例和文档。
仓库链接：d3-geo

Geospatial Data Abstraction Library (GDAL)： GDAL 是一个用于处理地理空间数据的库，支持多种格式和数据源。该项目的 GitHub 仓库包含了 GDAL 相关的示例和文档。
仓库链接：GDAL

Folium： Folium 是一个基于 Python 的库，用于在地图上创建交互式地理数据可视化。它可以生成 Leaflet 地图，并支持将数据叠加在地图上。
仓库链接：Folium

#%%------------计算不同时间段的变量的PSI
import numpy as np

def calculate_psi(base_probs, new_probs):
    psi_values = (base_probs - new_probs) * np.log(base_probs / new_probs)
    psi = np.sum(psi_values)
    return psi
new_probs
# 基准数据集和新数据集的分布比例
base_probs = np.array([0.1, 0.2, 0.3, 0.4])
new_probs = np.array([0.15, 0.18, 0.28, 0.39])

psi = calculate_psi(base_probs, new_probs)
print("PSI:", psi)

#%%------------模型画图 K-S
import numpy as np
import matplotlib.pyplot as plt
from sklearn.metrics import roc_curve, auc

# 生成随机分类模型的真实标签和预测概率
y_true = np.random.randint(0, 2, size=100)
y_prob = np.random.rand(100)

# 计算 ROC 曲线
fpr, tpr, thresholds = roc_curve(y_true, y_prob)
roc_auc = auc(fpr, tpr)

# 计算 K-S 统计量
ks_values = np.abs(tpr - fpr)
max_ks_index = np.argmax(ks_values)
max_ks = ks_values[max_ks_index]
threshold_for_max_ks = thresholds[max_ks_index]

# 绘制 K-S 图
plt.figure(figsize=(8, 6))
plt.plot(thresholds, ks_values, label=f'Max K-S = {max_ks:.2f} at Threshold = {threshold_for_max_ks:.2f}')
plt.xlabel('Probability Threshold')
plt.ylabel('K-S Value')
plt.title('K-S Curve')
plt.legend(loc='upper left')
plt.grid(True)
plt.show()

#%%------------模型画图 AUC
import numpy as np
import matplotlib.pyplot as plt
from sklearn.metrics import roc_curve, auc

# 生成随机分类模型的真实标签和预测概率
y_true = np.random.randint(0, 2, size=100)
y_prob = np.random.rand(100)

# 计算 ROC 曲线
fpr, tpr, thresholds = roc_curve(y_true, y_prob)
roc_auc = auc(fpr, tpr)

# 绘制 AUC 曲线
plt.figure(figsize=(8, 6))
plt.plot(fpr, tpr, color='darkorange', lw=2, label='ROC curve (area = %0.2f)' % roc_auc)
plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
plt.xlim([0.0, 1.0])
plt.ylim([0.0, 1.05])
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('Receiver Operating Characteristic (ROC) Curve')
plt.legend(loc="lower right")
plt.grid(True)
plt.show()


def calculate_minimum_copper_coins():
    copper_coins = 1
    while True:
        silver_coins = copper_coins // 3
        gold_coins = silver_coins // 3
        if gold_coins >= 1:
            return copper_coins
        copper_coins += 1

minimum_copper_coins = calculate_minimum_copper_coins()
print("小兔至少需要有", minimum_copper_coins, "枚铜币才能兑换出一枚金币。")

#%%------------类别不平衡
# 解决数据不平衡问题的方法包括：

# 重采样（Resampling）：包括过采样（增加少数类别样本数量）和欠采样（减少多数类别样本数量）两种方法。

# 合成样本（Synthetic Samples）：通过合成生成一些类似的样本来平衡类别分布，如SMOTE算法。

# 类别权重调整（Class Weighting）：在模型训练时，为不同类别设置不同的权重，使模型更关注少数类别。

# 集成方法（Ensemble Methods）：如随机森林和XGBoost，在集成中处理不平衡类别更有效。

# 算法选择：选择适合处理不平衡数据的算法，如SVM、决策树等。

# 评估指标调整：使用与问题相关的评估指标，如精确率、召回率、F1值等

# F1值的取值范围在0和1之间，值越接近1表示模型的性能越好。F1值对于处理不平衡类别分布的数据集尤为重要，因为它考虑了模型的预测结果和真实情况之间的平衡。在一些情况下，F1值也可以作为调整阈值的依据，以平衡精确率和召回率之间的关系

#%%------------可视化框架

pyecharts:
    
DataZoomOpts-区域缩放配置项
Timeline-时间轴




# Metabase：Metabase 允许用户使用 SQL 查询来创建报表和可视化，也可以使用 Metabase 的查询编辑器来构建查询和报表。

# Superset：Apache Superset 支持使用 SQL 或 Python 编写查询和报表逻辑，可以在报表中嵌入 Python 代码来实现高级计算和自定义可视化。

# Redash：Redash 支持通过 SQL 查询来创建报表，也可以在报表中使用 Python 脚本来执行计算、处理数据和生成可视化。

#%%------------lightgbm模型_Banking-Case
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
# import pandas_profiling as pp
from sklearn.model_selection import train_test_split
from sklearn.metrics import roc_auc_score, accuracy_score
import lightgbm as lgb
from lightgbm import log_evaluation, early_stopping#新版

bank_data=pd.read_csv("D:/Work/github/Banking-Case-study---LightGBM/bank-additional-full.csv", delimiter=";")
bank_data.sample(3)
bank_data.info()

#Checking whether the data is balanced or unbalanced : distribution of class variable y--jupyter
y_age=bank_data.groupby('y').age.count()
sns.barplot(x=y_age.index, y=y_age[:])


# Identifying categorical and numerical variables
bank_data.columns
# Separating out Categorical and Numerical Variables
categorical_cols = []
num_cols = []
for variable in bank_data.columns:
    if bank_data[variable].dtype.name in ['object']:
        categorical_cols.append(variable)
    else:
        num_cols.append(variable)
        
# Missing Value treatment
for i in categorical_cols:
    a = bank_data[i].fillna('Missing')
    print("\nThe distinct categories of variable '{0}' is :\n ".format(i), a.value_counts(), '\n')
print("The conclusion is that none of the categorical or numerical variable requires missing values treatment \n")
nmiss_df = bank_data.isnull().sum(axis=0)
nmiss_df.name = 'NMiss'
test_song=pd.concat([bank_data.describe().T,nmiss_df], axis=1,join='inner')




# 箱体图-with y jupyter
for num_variable in num_cols:
    fig,axes = plt.subplots(figsize=(10,4))
    sns.boxplot(x='y', y=num_variable, data = bank_data)
    plt.title(str('Box Plot of ') + str(num_variable))

# 箱体图-not with y jupyter
for num_variable in num_cols:
    fig,axes = plt.subplots(figsize=(10,4))
    sns.boxplot(y=num_variable, data = bank_data)
    plt.title(str('Box Plot of ') + str(num_variable))


## only 'cons.conf.idx' requires outlier treatment  箱体图 秀 微调操作
threshold=bank_data['cons.conf.idx'].quantile(0.99)
bank_data['cons.conf.idx'] =bank_data['cons.conf.idx'].clip(upper=threshold)
sns.boxplot(y="cons.conf.idx", data = bank_data)

# 连续变量直方图-jupyter
for num_variable in num_cols:
    plt.figure(figsize=(8,6))
    sns.displot(bank_data[num_variable])
    plt.title(str('Distribution Plot of ') + str(num_variable))
    
#Age: 分箱后的直方图-jupyter
sns.distplot(bank_data['age'], hist=True, kde=True, 
             bins=10, color = 'orange',
             hist_kws={'edgecolor':'black'})


# Univariate analysis of features (categorical variables) --分类变量直方图,还亮了百分比
# Defining a function which takes a feature and dataset as input and shows the countplot of the feature
def countplot(feature, dataset):
    plt.figure(figsize=(13,8))
    total = len(dataset[feature])*1.
    
    #Plotting ditribution of different categories of the feature in the dataset
    ax=sns.countplot(x=feature, data=dataset)
    
    # To show % contribution of each category of the feature
    for p in ax.patches:
        ax.annotate('{:.1f}%'.format(100*p.get_height()/total), (p.get_x()+0.2, p.get_height()+5))
    
    # Rotating X labels for readibility
    ax.set_xticklabels(ax.get_xticklabels(), rotation=40, ha="right")
    
    plt.title(str('Count Plot of ') + str(feature))
    
    plt.show()


# Defining a function which takes a feature and dataset as input and shows the countplot of the feature with class information
%matplotlib inline

def countplot_withY(feature, dataset):
    plt.figure(figsize=(13,8))
    total = len(dataset[feature])*1.
    
    ax=sns.countplot(x=feature, data=dataset, hue="y")
    
    # To show % contribution of each category of the feature
    for p in ax.patches:
        ax.annotate('{:.1f}%'.format(100*p.get_height()/total), (p.get_x()+0.1, p.get_height()+5))

    # Rotating X labels for readibility
    ax.set_xticklabels(ax.get_xticklabels(), rotation=40, ha="right")
    
    plt.title(str('Count Plot of ') + str(feature) + str(' with Y'))
    
    plt.show()

# jupyters
for variable in categorical_cols:
    countplot(variable, bank_data)
    countplot_withY(variable, bank_data)

# 相关系数矩阵-jupyters
plt.subplots(figsize=(12, 12))
sns.heatmap(bank_data.corr(), annot=True)
plt.show()


bank_data.shape#查看行数与列数
bank_data=bank_data.drop_duplicates()#去重
bank_data.shape#查看去重后的行数与列数
data_y= (bank_data['y'] == 'yes')*1#将y变量的yes转化为1 另一个值就是0了
bank_data.drop('y', axis=1, inplace = True)#训练样本去除Y变量
bank_data.shape#查看微调后的行数与列数

# filter categorical columns using dtypes as objects and turn it into a list
categorical_cols = bank_data.columns[bank_data.dtypes==object].tolist()
categorical_cols

# 将类别变量的类别打成数字码，不过还是建议达成one-hot码会好一点
from sklearn.preprocessing import LabelEncoder
bank_data_final=bank_data
for i in categorical_cols:
    le = LabelEncoder()
    bank_data_final[i] = le.fit_transform(bank_data_final[i])
print(bank_data_final.head())

#Drop feature duration，这个通话时长据说是强变量，越长target越可能是1，之所以去掉是因为业务场景认为，通话结束后就确定了0或1，都不用分析预测
bank_data_final.drop("duration", axis=1, inplace=True)
bank_data_final.shape


# 建模
data_train, data_test, y_train, y_test = train_test_split(bank_data_final, data_y, train_size = 0.7, test_size = 0.3)
lgb_train = lgb.Dataset(data=data_train, label=y_train,  free_raw_data=False)
#Creat Evaluation Dataset 
lgb_eval = lgb.Dataset(data=data_test, label=y_test, reference=lgb_train,  free_raw_data=False)

# Categorical index needed bacause lightgbm can handle categorical features very well，这里能看懂，但不知道有啥意义
categorical_index = [1,2,3,4,5,6,7,8,9,13]
print('Categorical parametres: ' + str(data_train.columns[categorical_index].values))


# 这里可能要去找新一点的代码，因为这个bank-case太久远了，用的lgm是旧版，里面的参数也是旧版，导致新版都跑不了旧版的代码里需要的参数
params = {
    'task': 'train',
    'boosting_type': 'gbdt',
    'objective': 'binary',      #Used ‘binary’ as objective --> this is classification problem)
    'metric': 'auc',
    'num_leaves': 31,
    'learning_rate': 0.05,
    'verbose': -1
}

gbm = lgb.train(params,
                lgb_train,
                valid_sets = lgb_eval,
                categorical_feature = categorical_index,
                num_boost_round= 150,
                callbacks = [log_evaluation(period=100), early_stopping(stopping_rounds=30)]
        )


# predict
y_pred = gbm.predict(data_test, num_iteration=gbm.best_iteration)

print('The Best iteration: ', gbm.best_iteration)
print('roc_auc_score:', roc_auc_score(y_test, y_pred))
print('accuracy_score:', accuracy_score(y_test, ( y_pred>= 0.5)*1))


# jupyter ，这里因为少了evals_result表(旧版参数)，画不了图
ax = lgb.plot_metric(evals_result, metric='auc')
ax.set_title('Variation of the Curved Area According to Iteration')
ax.set_xlabel('İteration')
ax.set_ylabel('roc_auc_score')
ax.legend_.remove()


# jupyter
ax = lgb.plot_importance(gbm, max_num_features=10)
ax.set_title('The values of Parametres')
ax.set_xlabel('Values')
ax.set_ylabel('Parametres')


#%%------------pdf转为word,太难了
import PyPDF2
import docx#google搜过import pdf2docx 报错No module named 'exceptions'
from docx import Document

# 打开PDF文件
pdf_file = open(r'D:\Work\Information\涛哥未归类\征信与个人信息授权协议.pdf', 'rb')

# 创建一个PyPDF2的PDF文件阅读器对象
pdf_reader = PyPDF2.PdfReader(pdf_file)

# 创建一个Word文档对象
doc = Document()

# 遍历PDF中的每一页并提取文本，然后写入Word文档
for page_num in range(len(pdf_reader.pages)):
    page = pdf_reader.pages[page_num]
    text = page.extract_text()
    doc.add_paragraph(text)


# 调整Word文档的排版
doc.add_page_break()

# 保存Word文档
doc.save(r'D:\Work\Information\涛哥未归类\test.docx')

# 关闭PDF文件
pdf_file.close()


import PyPDF2
from docx import Document
from docx.shared import Inches

# 打开PDF文件
pdfFileObj = open(r'D:\Work\Information\涛哥未归类\征信与个人信息授权协议.pdf', 'rb')
pdfReader = PyPDF2.PdfReader(pdfFileObj)

# 新建Word文档
doc = Document()

# 将PDF中的文本提取出来并添加到Word中
for page in range(pdfReader.numPages):
    pageObj = pdfReader.getPage(page)
    text = pageObj.extractText()
    doc.add_paragraph(text)

# 调整Word文档的排版
doc.add_page_break()
doc.add_heading('Title', 0)
doc.add_picture('image.jpg', width=Inches(6))

# 保存Word文档
doc.save('example.docx')

import pdf2docx
from pdf2docx import Converter
import re

def pdf_to_word(fileName) :
    pdf_file = fileName
    # 正则获取不含文件类型后缀的部分，用于组成word文档绝对路径
    name = re.findall(r'(.*?)\. ' ,pdf_file)[0]
    docx_file = f'{name}.docx'
    cv = Converter(pdf_file)
    cv.convert( docx_file, start=e, end=None)cv.close()

#%%------------账单案例
import pandas as pd
import numpy as np
import datetime

due_data=pd.read_excel(r"D:\Work\out_data\核心概念\逾期状态.xlsx")
for avar in ['账单日','实际还款日']:
    due_data[avar]=pd.to_datetime(due_data[avar]).dt.date

def od_days_count(x):#用回溯日与应还日比较来计算逾期天数
    if x.实际还款日>cut_date:
        return (cut_date-x.账单日).days
    else:
        return 0
    
for i in range(90):#每次循环生成一个回溯日当天的逾期天数
    cut_date=datetime.date.today()-datetime.timedelta(days=i)#回溯日期
    df_temp=due_data.loc[due_data.账单日<cut_date,:].copy()#每次回溯时保留账单日是回溯日以前的账单
    df_temp['od_days']=df_temp.apply(od_days_count,axis=1)#然后对每个账单计算逾期天数
    od_days=df_temp['od_days'].max()#然后对所有账单的逾期天数求max
    data = {'cut_date': cut_date,
         'od_days': od_days}
    df_tmep1 = pd.DataFrame(data,index=[i])
    if i==0:#结果表的首行
        df_end = df_tmep1
    else:
        df_end = pd.concat([df_end,df_tmep1],axis=0,sort=True)#结果表的拼接



df_end = df_end.rename_axis('往前推的天数')
df_end.to_excel(r'D:\Work\out_data\核心概念\逾期案例.xlsx') 


#%%------------信贷产品定价
需要了解的概念：
年利率、月利率、还款期限、复利等


第1种情况:12期产品，已知年化收益率,等额本息,
第2种情况:

import pandas as pd

def calculate_emi(principal, annual_interest_rate, loan_term):
    # 将年利率转换为月利率
    monthly_interest_rate = annual_interest_rate / 12 / 100
    
    # 计算每月还款额
    emi = principal * monthly_interest_rate * (1 + monthly_interest_rate)**loan_term / ((1 + monthly_interest_rate)**loan_term - 1)
    
    return emi

# 设定参数
principal = 10000  # 贷款本金
annual_interest_rate = 36  # 年化收益率
loan_term = 12  # 贷款期限（月）

# 初始化 DataFrame
columns = ['Month', 'Total Payment', 'Principal Payment', 'Interest Payment', 'Remaining Principal']
df = pd.DataFrame(columns=columns)

# 计算每月还款额
monthly_payment = calculate_emi(principal, annual_interest_rate, loan_term)

# 构建还款计划并添加到 DataFrame
for month in range(1, loan_term + 1):
    interest_payment = principal * annual_interest_rate / 12 / 100
    principal_payment = monthly_payment - interest_payment
    principal -= principal_payment
    
    # 将还款计划添加到 DataFrame
    row = {'Month': month, 'Total Payment': monthly_payment, 'Principal Payment': principal_payment, 'Interest Payment': interest_payment, 'Remaining Principal': principal}
    df = df.append(row, ignore_index=True)

df.to_excel(r'D:\Work\risk_rule\IRR\IRR样例.xlsx')



import pandas as pd

def calculate_emi(principal, annual_interest_rate, loan_term_months):
    # 将年化利率转换为月利率
    monthly_interest_rate = annual_interest_rate / 12 / 100
    
    # 计算每月还款额的公式
    emi = principal * (monthly_interest_rate * (1 + monthly_interest_rate)**loan_term_months) / \
          ((1 + monthly_interest_rate)**loan_term_months - 1)
    
    return emi

def loan_amortization_schedule(principal, annual_interest_rate, loan_term_months):
    monthly_interest_rate = annual_interest_rate / 12 / 100
    
    emi = calculate_emi(principal, annual_interest_rate, loan_term_months)
    
    # 创建一个空的 DataFrame 来存储还款计划
    columns = ["月份", "本月还款额", "本金部分", "利息部分", "剩余本金"]
    schedule_df = pd.DataFrame(columns=columns)
    
    remaining_principal = principal
    
    for month in range(1, loan_term_months + 1):
        interest_payment = remaining_principal * monthly_interest_rate
        principal_payment = emi - interest_payment
        remaining_principal -= principal_payment
        
        # 将每月的还款详情添加到 DataFrame
        row_data = [month, round(emi, 2), round(principal_payment, 2), round(interest_payment, 2), round(remaining_principal, 2)]
        schedule_df = schedule_df.append(pd.Series(row_data, index=columns), ignore_index=True)
    
    return schedule_df

# 例子
principal_loan_amount = 10000  # 贷款本金
annual_interest_rate = 36     # 年化利率
loan_term_months = 24          # 贷款期限（月）

# 获取还款计划的 DataFrame
df = loan_amortization_schedule(principal_loan_amount, annual_interest_rate, loan_term_months)

# 打印 DataFrame
df.to_excel(r'D:\Work\risk_rule\IRR\IRR样例1.xlsx')

拒绝原因









